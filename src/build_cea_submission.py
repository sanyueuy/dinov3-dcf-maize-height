#!/usr/bin/env python3
from __future__ import annotations

import csv
import json
import shutil
import subprocess
import textwrap
import zipfile
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from build_plant_methods_submission_ready import metrics


ROOT = Path(r"D:\cornTrain\DINOV3")
OUT = ROOT / "cea_submission"
FIG_SRC = ROOT / "paper_figures"
FIG_OUT = OUT / "figures"
TABLE_OUT = OUT / "tables"
REPRO_OUT = OUT / "reproducibility_json"
DATA325_IMAGE_DIR = ROOT / "images"
DATA325_TTA8_JSON = ROOT / "data325_zero_shot_attn_aug_tta8" / "data325_zero_shot_comparison_attn_aug_tta8.json"
SOURCE_SAMPLE_DIR = ROOT / "hand_bbox_visualization" / "samples"

BLUE = "4472C4"      # matches common.py COLOR_SOURCE
GREEN = "1D9E75"     # matches common.py COLOR_TARGET
ORANGE = "D85A30"    # matches common.py COLOR_WARN
GRAY = "888780"      # matches common.py COLOR_NEUTRAL
BLACK = "202124"
HEADER = "EAF1F7"
LIGHT_GREEN = "E7F3EE"
LIGHT_ORANGE = "F8EEE9"

TITLE = (
    "Attention-guided DINOv3-DiffCorn-Fusion for zero-shot "
    "cross-greenhouse maize height estimation"
)
AUTHOR_NAME = "Hong Wu"
AUTHOR_EMAIL = "wuhong@cau.edu.cn"
COAUTHOR_NAME = "Jian Chen"
AFFILIATION = "College of Engineering, China Agricultural University, Beijing 100083, China"
CORRESPONDING_EMAIL = "jchen@cau.edu.cn"
OPEN_RELEASE_ARCHIVE = "dinov3_dcf_maize_height_open_release_v0_1_0.zip"
REPOSITORY_URL = "https://github.com/sanyueuy/dinov3-dcf-maize-height"
CAPTURE_TOOL_URL = "https://github.com/sanyueuy/corn-capture"

FIGURES = [
    (
        "fig_real_protocol",
        "Figure_1_Real_image_protocol",
        "Fig. 1. Real-image protocol for zero-shot cross-greenhouse maize height estimation. (a) Source-domain hand-box examples with measured heights. (b) Independent DATA325 target greenhouse images with manual bounding-box annotations. (c) Extracted ROI crops and per-box evaluation records including ground-truth height, DINOv3-DCF prediction, absolute error, camera height, and TTA standard deviation.",
        6.15,
    ),
    (
        "fig3",
        "Figure_2_Dataset_statistics",
        "Fig. 2. Source and DATA325 dataset statistics. The source-domain hand-box set spans taller plants, whereas DATA325 contains more early-stage plants and a different target greenhouse distribution.",
        4.95,
    ),
    (
        "fig1",
        "Figure_3_DINOv3_DCF_workflow",
        "Fig. 3. DINOv3-DiffCorn-Fusion (DINOv3-DCF) workflow for ROI-level maize height estimation. A frozen DINOv3 ViT-L backbone extracts plant-region features, the DCF head maps visual and camera-height inputs to phytomer parameters, and height is obtained from internode-related outputs.",
        6.15,
    ),
    (
        "fig2",
        "Figure_4_Attention_weighted_pooling",
        "Fig. 4. Attention-weighted patch aggregation. Final-layer CLS-to-patch attention weights are averaged across heads, normalized across patch positions, and used to pool patch tokens before DCF regression.",
        6.15,
    ),
    (
        "fig4",
        "Figure_5_Domain_shift_tsne",
        "Fig. 5. DINOv3 feature-space domain shift between source ROIs and DATA325 target ROIs. The t-SNE visualization uses 20 sampled source and 20 sampled target ROI features.",
        4.95,
    ),
    (
        "fig5",
        "Figure_6_Ablation_results",
        "Fig. 6. Zero-shot DATA325 ablation results. Replacing CLS pooling with attention-weighted patch aggregation produced the largest single reduction in external MAE.",
        6.15,
    ),
    (
        "fig6",
        "Figure_7_Height_bin_errors",
        "Fig. 7. DATA325 error distribution by ground-truth plant-height bin for the best Attn+aug+TTA8 model. Early-stage plants below 80 cm remain the dominant source of relative error.",
        6.15,
    ),
    (
        "fig_real_stage_error_gallery",
        "Figure_8_DATA325_stage_error_gallery",
        "Fig. 8. DATA325 real-image stage and error gallery. Full target-greenhouse images with manual boxes and ROI crops show that early plants and cluttered backgrounds are visually overrepresented among the high-error cases.",
        6.15,
    ),
    (
        "fig7",
        "Figure_9_Diagnostic_experiments",
        "Fig. 9. Diagnostic experiments for alternative explanations of the domain gap, including geometric priors, feature-statistic alignment, and domain-adversarial training.",
        6.15,
    ),
    (
        "fig8",
        "Figure_10_Qualitative_DATA325_examples",
        "Fig. 10. Real DATA325 ROI examples and DINOv3 attention overlays. Low-stage examples have less plant-structure evidence and show larger relative errors than taller plants.",
        6.15,
    ),
]
FIG_NAME_BY_STEM = {source_stem: upload_stem for source_stem, upload_stem, _, _ in FIGURES}

REPRO_FILES = [
    ROOT / "data325_zero_shot_attn_aug_tta8" / "data325_zero_shot_comparison_attn_aug_tta8.json",
    ROOT / "data325_zero_shot_corrected_camheight" / "data325_zero_shot_comparison_corrected_camheight.json",
    ROOT / "data325_zero_shot_attn_aug" / "data325_zero_shot_comparison_attn_aug.json",
    ROOT / "data325_zero_shot_attn_aug_featurealign" / "data325_zero_shot_comparison_attn_aug_featurealign.json",
    ROOT / "data325_zero_shot_dann" / "data325_zero_shot_comparison_dann.json",
    ROOT / "tsne_source_vs_data325.json",
    ROOT / "data325_zero_shot_attn_aug_featurealign" / "data325_feature_stats_attn_aug_featurealign.json",
    ROOT / "bbox_geometry_prior_phaseA.json",
    ROOT / "attention_geometry_prior_phaseA.json",
    ROOT / "checkpoints" / "diffcorn_fusion_hand_dann_history.json",
]


def reset_dir(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def fmt(value: float, ndigits: int = 2) -> str:
    try:
        value_f = float(value)
    except (TypeError, ValueError):
        return "NA"
    if not np.isfinite(value_f):
        return "NA"
    return f"{value_f:.{ndigits}f}"


def setup_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    # Match the compact Elsevier manuscript-source feel used in the reference
    # document rather than a report-style A4 cover-page layout.
    sec.page_width = Cm(19.0)
    sec.page_height = Cm(26.0)
    sec.top_margin = Cm(1.35)
    sec.bottom_margin = Cm(2.65)
    sec.left_margin = Cm(1.35)
    sec.right_margin = Cm(2.1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)

    def ensure_style(name: str, base: str = "Normal"):
        if name in styles:
            return styles[name]
        return styles.add_style(name, 1)

    style_specs = [
        ("Title", 13.5, BLACK, True, False),
        ("Subtitle", 10, BLACK, False, False),
        ("Heading 1", 10.5, BLACK, True, False),
        ("Heading 2", 10, BLACK, True, False),
        ("Heading 3", 10, BLACK, True, True),
        ("Els-Title", 13.5, BLACK, True, False),
        ("Els-Affiliation", 9.5, BLACK, False, False),
        ("Els-body-text", 10, BLACK, False, False),
        ("Els-reference", 9, BLACK, False, False),
        ("Els-1storder-head", 10.5, BLACK, True, False),
        ("Els-2ndorder-head", 10, BLACK, True, False),
    ]
    for name, size, color, bold, italic in style_specs:
        style = styles[name] if name in styles else styles.add_style(name, 1)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = bold
        style.font.italic = italic
        style.paragraph_format.line_spacing = 2.0
        style.paragraph_format.space_after = Pt(0)

    if "CaptionClean" not in styles:
        caption = styles.add_style("CaptionClean", 1)
    else:
        caption = styles["CaptionClean"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_after = Pt(6)
    return doc


def add_p(doc: Document, text: str = "", style: str | None = None) -> None:
    p = doc.add_paragraph(style=style or "Els-body-text")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="CaptionClean")
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)


def add_figure(doc: Document, stem: str, caption: str, width_in: float) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(FIG_OUT / f"{FIG_NAME_BY_STEM.get(stem, stem)}.png"), width=Inches(min(width_in, 6.05)))
    add_caption(doc, caption)


def add_named_figure(doc: Document, stem: str) -> None:
    for source_stem, _, caption, width_in in FIGURES:
        if source_stem == stem:
            add_figure(doc, stem, caption, width_in)
            return
    raise KeyError(f"Unknown figure stem: {stem}")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))


def set_table_borders(table, color: str = "808080", size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_no_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "FFFFFF")


def fill_cell(cell, text: str, bold: bool = False, align: int = WD_ALIGN_PARAGRAPH.CENTER, size: float = 8.2) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_layout = table._tbl.tblPr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        table._tbl.tblPr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    set_table_borders(table)

    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "FFFFFF")
        set_cell_width(cell, widths_cm[idx])
        fill_cell(cell, text, bold=True, size=8.0)
    for row_idx, row_values in enumerate(rows, start=1):
        cells = table.add_row().cells
        for col_idx, text in enumerate(row_values):
            cell = cells[col_idx]
            set_cell_width(cell, widths_cm[col_idx])
            shade_cell(cell, "FFFFFF")
            align = WD_ALIGN_PARAGRAPH.LEFT if col_idx in (0, 1, len(headers) - 1) else WD_ALIGN_PARAGRAPH.CENTER
            fill_cell(cell, text, align=align, size=7.8)
    doc.add_paragraph()


def pil_font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    candidates = [
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def draw_wrapped(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, font: ImageFont.ImageFont, fill: tuple[int, int, int], max_width: int, line_gap: int = 4) -> int:
    x, y = xy
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textlength(trial, font=font) <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    line_h = int(font.size * 1.15) if hasattr(font, "size") else 16
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += line_h + line_gap
    return y


def paste_fit(canvas: Image.Image, image: Image.Image, box: tuple[int, int, int, int], bg: tuple[int, int, int] = (255, 255, 255)) -> tuple[int, int, int, int]:
    x0, y0, x1, y1 = box
    target = Image.new("RGB", (x1 - x0, y1 - y0), bg)
    im = image.convert("RGB")
    im.thumbnail((x1 - x0, y1 - y0), Image.Resampling.LANCZOS)
    px = (target.width - im.width) // 2
    py = (target.height - im.height) // 2
    target.paste(im, (px, py))
    canvas.paste(target, (x0, y0))
    return (x0 + px, y0 + py, x0 + px + im.width, y0 + py + im.height)


def draw_bbox_on_image(path: Path, bbox: list[float], color: tuple[int, int, int] = (0, 158, 115)) -> Image.Image:
    image = Image.open(path).convert("RGB")
    draw = ImageDraw.Draw(image)
    x, y, w, h = [int(round(v)) for v in bbox]
    width = max(5, image.width // 260)
    draw.rectangle((x, y, x + w, y + h), outline=color, width=width)
    return image


def crop_roi(path: Path, bbox: list[float], pad: int = 8) -> Image.Image:
    image = Image.open(path).convert("RGB")
    x, y, w, h = [int(round(v)) for v in bbox]
    x0 = max(0, x - pad)
    y0 = max(0, y - pad)
    x1 = min(image.width, x + w + pad)
    y1 = min(image.height, y + h + pad)
    return image.crop((x0, y0, x1, y1))


def save_png_pdf(image: Image.Image, upload_stem: str) -> None:
    png_path = FIG_OUT / f"{upload_stem}.png"
    pdf_path = FIG_OUT / f"{upload_stem}.pdf"
    image.save(png_path)
    image.convert("RGB").save(pdf_path, "PDF", resolution=300.0)


def patch_embedded_figure_titles() -> None:
    fixes = [
        ("Figure_5_Domain_shift_tsne", "t-SNE visualization of source and target domain feature distributions", 82, 38),
    ]
    for upload_stem, title, cover_h, font_size in fixes:
        png_path = FIG_OUT / f"{upload_stem}.png"
        if not png_path.exists():
            continue
        image = Image.open(png_path).convert("RGB")
        draw = ImageDraw.Draw(image)
        draw.rectangle((0, 0, image.width, cover_h), fill=(255, 255, 255))
        font = pil_font(font_size, True)
        text_w = draw.textlength(title, font=font)
        draw.text(((image.width - text_w) / 2, max(6, (cover_h - font_size) // 2)), title, font=font, fill=(32, 33, 36))
        image.save(png_path)
        image.save(FIG_OUT / f"{upload_stem}.pdf", "PDF", resolution=300.0)


def load_data325_records() -> list[dict]:
    data = json.loads(DATA325_TTA8_JSON.read_text(encoding="utf-8"))
    records = data["model_results"][0]["box_results"]
    out = []
    seen = set()
    for record in records:
        image_path = DATA325_IMAGE_DIR / record["file_name"]
        key = (record["file_name"], record["box_id"])
        if image_path.exists() and key not in seen:
            seen.add(key)
            out.append(record)
    return out


def choose_record(records: list[dict], predicate, rank: str = "low") -> dict:
    subset = [r for r in records if predicate(r)]
    if not subset:
        raise RuntimeError("No DATA325 records matched the requested figure filter.")
    subset = sorted(subset, key=lambda r: r["abs_error_cm"])
    if rank == "high":
        return subset[-1]
    if rank == "mid":
        return subset[len(subset) // 2]
    return subset[0]


def source_sample(pattern: str) -> Path:
    matches = sorted(SOURCE_SAMPLE_DIR.glob(pattern))
    if not matches:
        raise RuntimeError(f"Missing source-domain sample for pattern {pattern}")
    return matches[0]


def build_real_image_protocol_figure(records: list[dict]) -> None:
    upload_stem = FIG_NAME_BY_STEM["fig_real_protocol"]
    canvas = Image.new("RGB", (2400, 1810), (255, 255, 255))
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(52, True)
    head_font = pil_font(31, True)
    label_font = pil_font(26, True)
    body_font = pil_font(24)
    small_font = pil_font(21)
    tag_font = pil_font(23, True)
    blue = (68, 114, 196)     # #4472C4
    green = (29, 158, 117)    # #1D9E75
    orange = (216, 90, 48)    # #D85A30
    gray = (136, 135, 128)    # #888780
    black = (32, 33, 36)
    light_blue = (234, 241, 247)
    light_green = (231, 243, 238)
    light_orange = (248, 238, 233)

    # Title
    draw.text((70, 38), "Real-image zero-shot protocol for DINOv3-DCF", font=title_font, fill=black)
    draw.text((70, 105), "Source-domain hand boxes and independent DATA325 target greenhouse ROIs", font=body_font, fill=gray)

    # Column headers with (a)(b)(c) labels
    cols = [(70, 230, 730, 1545), (870, 230, 1530, 1545), (1670, 230, 2330, 1545)]
    headers = [
        ("(a)", "Source hand-box image", blue, light_blue),
        ("(b)", "DATA325 target image + bbox", green, light_green),
        ("(c)", "ROI crop + prediction record", orange, light_orange),
    ]
    for (x0, y0, x1, _), (tag, header, col_color, bg_color) in zip(cols, headers):
        draw.rectangle((x0, 175, x1, 225), fill=bg_color, outline=col_color, width=2)
        draw.text((x0 + 15, 184), tag, font=tag_font, fill=col_color)
        draw.text((x0 + 55, 184), header, font=head_font, fill=col_color)

    # Row definitions with height range labels
    rows = [
        ("Early-stage (0-80 cm)", "106-64-*_bbox.jpg", choose_record(records, lambda r: r["true_height_cm"] < 80, "low"), green),
        ("Mid-stage (80-120 cm)", "140-121-*_bbox.jpg", choose_record(records, lambda r: 80 <= r["true_height_cm"] < 120, "mid"), orange),
        ("Tall (120+ cm)", "140-270-*_bbox.jpg", choose_record(records, lambda r: r["true_height_cm"] >= 120, "low"), blue),
    ]
    y_starts = [255, 695, 1135]
    for row_idx, ((label, source_pattern, record, color), y0) in enumerate(zip(rows, y_starts)):
        y1 = y0 + 385
        # Row label on the left
        draw.text((72, y0 - 32), label, font=label_font, fill=color)

        # Column cell outlines
        for x0, _, x1, _ in cols:
            draw.rectangle((x0, y0, x1, y1), outline=(218, 224, 230), width=2)

        # (a) Source image
        source = Image.open(source_sample(source_pattern)).convert("RGB")
        paste_fit(canvas, source, (95, y0 + 25, 705, y1 - 58), (250, 250, 250))
        draw.text((105, y1 - 43), f"Source example ({source_pattern.split('-')[1]} cm class)", font=small_font, fill=gray)

        # Arrow from (a) to (b) - zero-shot transfer indicator
        arrow_y = y0 + 190
        draw.line((740, arrow_y, 860, arrow_y), fill=gray, width=3)
        draw.polygon([(855, arrow_y - 8), (870, arrow_y), (855, arrow_y + 8)], fill=gray)
        if row_idx == 1:
            draw.text((745, arrow_y - 28), "Zero-shot", font=small_font, fill=gray)

        # (b) Target image with bbox
        target_path = DATA325_IMAGE_DIR / record["file_name"]
        target = draw_bbox_on_image(target_path, record["bbox"], color)
        paste_fit(canvas, target, (895, y0 + 25, 1505, y1 - 58), (250, 250, 250))
        draw.text((905, y1 - 43), f"Manual bbox: {record['box_id']}", font=small_font, fill=gray)

        # Arrow from (b) to (c)
        draw.line((1515, arrow_y, 1635, arrow_y), fill=gray, width=3)
        draw.polygon([(1630, arrow_y - 8), (1645, arrow_y), (1630, arrow_y + 8)], fill=gray)

        # (c) ROI crop + prediction card
        roi = crop_roi(target_path, record["bbox"])
        paste_fit(canvas, roi, (1695, y0 + 25, 1965, y1 - 72), (250, 250, 250))
        card_x = 1990
        card_y = y0 + 28
        draw.rectangle((card_x, card_y, 2308, y1 - 68), fill=(247, 249, 251), outline=color, width=3)
        draw.text((card_x + 20, card_y + 18), "Evaluation record", font=label_font, fill=color)
        lines = [
            f"GT height: {record['true_height_cm']:.1f} cm",
            f"Prediction: {record['pred_height_cm']:.1f} cm",
            f"Abs. error: {record['abs_error_cm']:.1f} cm",
            f"Camera height: {record['camera_height_cm']:.0f} cm",
            f"TTA std.: {record['pred_std_cm']:.2f} cm",
        ]
        yy = card_y + 72
        for line in lines:
            draw.text((card_x + 20, yy), line, font=body_font, fill=black)
            yy += 45

    # Bottom workflow bar
    bar_y = 1620
    bar_h = 50
    bar_steps = [
        ("Source images", blue),
        ("Manual bbox", green),
        ("ROI crop", green),
        ("DINOv3-DCF", blue),
        ("Height prediction", orange),
    ]
    step_w = 400
    total_w = step_w * len(bar_steps) + 40 * (len(bar_steps) - 1)
    bar_x0 = (2400 - total_w) // 2
    for i, (step_label, step_color) in enumerate(bar_steps):
        sx = bar_x0 + i * (step_w + 40)
        draw.rectangle((sx, bar_y, sx + step_w, bar_y + bar_h), fill=step_color, outline=step_color, width=2)
        text_w = draw.textlength(step_label, font=label_font)
        draw.text((sx + (step_w - text_w) // 2, bar_y + 10), step_label, font=label_font, fill=(255, 255, 255))
        if i < len(bar_steps) - 1:
            ax = sx + step_w + 5
            draw.line((ax, bar_y + bar_h // 2, ax + 28, bar_y + bar_h // 2), fill=gray, width=3)
            draw.polygon([(ax + 23, bar_y + bar_h // 2 - 7), (ax + 33, bar_y + bar_h // 2), (ax + 23, bar_y + bar_h // 2 + 7)], fill=gray)

    draw.text((70, 1720), "All images are real captured greenhouse photographs; boxes, ROI crops, and prediction values are from the DATA325 evaluation JSON.", font=small_font, fill=gray)
    save_png_pdf(canvas, upload_stem)


def build_real_stage_error_gallery(records: list[dict]) -> None:
    upload_stem = FIG_NAME_BY_STEM["fig_real_stage_error_gallery"]
    canvas = Image.new("RGB", (2400, 2020), (255, 255, 255))
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(50, True)
    head_font = pil_font(29, True)
    label_font = pil_font(24, True)
    small_font = pil_font(20)
    black = (32, 33, 36)
    gray = (136, 135, 128)    # #888780
    colors = [(29, 158, 117), (216, 147, 36), (216, 90, 48)]  # green, amber, orange

    draw.text((70, 35), "DATA325 stage and error gallery from real target-greenhouse images", font=title_font, fill=black)
    draw.text((70, 100), "Each cell shows a full image with manual bbox, the extracted ROI crop, and the Attn+aug+TTA8 height record.", font=small_font, fill=gray)

    col_labels = ["Low error", "Median error", "High error"]
    rank_names = ["low", "mid", "high"]
    bins = [
        ("0-80 cm early stage", lambda r: r["true_height_cm"] < 80),
        ("80-120 cm jointing stage", lambda r: 80 <= r["true_height_cm"] < 120),
        ("120+ cm tall stage", lambda r: r["true_height_cm"] >= 120),
    ]
    cell_w = 680
    cell_h = 520
    x0s = [260, 955, 1650]
    y0s = [245, 790, 1335]
    for x, label, color in zip(x0s, col_labels, colors):
        draw.text((x + 18, 190), label, font=head_font, fill=color)
    for row_idx, (bin_label, predicate) in enumerate(bins):
        y0 = y0s[row_idx]
        for col_idx, rank in enumerate(rank_names):
            x0 = x0s[col_idx]
            record = choose_record(records, predicate, rank)
            color = colors[col_idx]
            draw.rectangle((x0, y0, x0 + cell_w, y0 + cell_h), fill=(255, 255, 255), outline=(218, 224, 230), width=2)
            target_path = DATA325_IMAGE_DIR / record["file_name"]
            target = draw_bbox_on_image(target_path, record["bbox"], color)
            roi = crop_roi(target_path, record["bbox"])
            paste_fit(canvas, target, (x0 + 20, y0 + 24, x0 + 365, y0 + 315), (248, 249, 250))
            paste_fit(canvas, roi, (x0 + 390, y0 + 24, x0 + 660, y0 + 315), (248, 249, 250))
            draw.rectangle((x0 + 20, y0 + 340, x0 + 660, y0 + 492), fill=(247, 249, 251), outline=color, width=2)
            text = (
                f"GT {record['true_height_cm']:.1f} cm | Pred {record['pred_height_cm']:.1f} cm | "
                f"Err {record['abs_error_cm']:.1f} cm | TTA std. {record['pred_std_cm']:.2f} cm"
            )
            draw_wrapped(draw, (x0 + 42, y0 + 365), text, label_font, black, 580, 5)
            draw.text((x0 + 42, y0 + 447), f"{record['file_name']} / {record['box_id']}", font=small_font, fill=gray)
        draw_wrapped(draw, (58, y0 + 18), bin_label, head_font, (68, 114, 196), 170, 6)

    draw.text((70, 1945), "The gallery is deterministic: examples are selected by ground-truth height bin and absolute-error rank from the same DATA325 result file used for quantitative evaluation.", font=small_font, fill=gray)
    save_png_pdf(canvas, upload_stem)


def build_real_image_figures() -> None:
    records = load_data325_records()
    build_real_image_protocol_figure(records)
    build_real_stage_error_gallery(records)


def build_diagnostic_figure() -> None:
    bbox = json.loads((ROOT / "bbox_geometry_prior_phaseA.json").read_text(encoding="utf-8"))
    attn_aug = json.loads((ROOT / "data325_zero_shot_attn_aug" / "data325_zero_shot_comparison_attn_aug.json").read_text(encoding="utf-8"))
    feature_align = json.loads((ROOT / "data325_zero_shot_attn_aug_featurealign" / "data325_zero_shot_comparison_attn_aug_featurealign.json").read_text(encoding="utf-8"))
    dann = json.loads((ROOT / "checkpoints" / "diffcorn_fusion_hand_dann_history.json").read_text(encoding="utf-8"))

    gt = [s["true_height_cm"] for s in bbox["samples"]]
    est = [s["estimated_height_cm"] for s in bbox["samples"]]
    attn_mae = attn_aug["model_results"][0]["summary"]["mae_cm"]
    align_mae = feature_align["model_results"][0]["summary"]["mae_cm"]
    epochs = dann["history"]["epoch"]
    domain_acc = dann["history"]["domain_acc"]

    # Unified color palette matching common.py
    c_source = "#4472C4"
    c_target = "#1D9E75"
    c_warn = "#D85A30"
    c_neutral = "#888780"

    plt.rcParams.update({"font.family": "DejaVu Sans", "font.size": 10})
    fig, axes = plt.subplots(1, 3, figsize=(13.0, 4.4), dpi=300)
    fig.patch.set_facecolor("white")

    # (a) Bbox geometric prior scatter
    ax = axes[0]
    ax.scatter(gt, est, s=40, color=c_neutral, alpha=0.70, edgecolor="white", linewidth=0.3)
    lo = min(min(gt), min(est)) - 5
    hi = max(max(gt), max(est)) + 5
    ax.plot([lo, hi], [lo, hi], "--", color="black", linewidth=1.2, label="Ideal (r=1.0)")
    coef = np.polyfit(gt, est, 1)
    xfit = [min(gt), max(gt)]
    yfit = [coef[0] * x + coef[1] for x in xfit]
    ax.plot(xfit, yfit, color=c_source, linewidth=2.0)
    ax.text(0.97, 0.13, f"r = {bbox['pearson_corr']:.3f}\nMAE = {bbox['mae_cm']:.2f} cm", transform=ax.transAxes, ha="right", va="bottom")
    ax.set_title("(a) Bbox geometric prior", weight="bold")
    ax.set_xlabel("Ground truth height (cm)")
    ax.set_ylabel("Geometric estimate (cm)")
    ax.legend(loc="upper left", frameon=True)
    ax.grid(True, color="#D0D0D0", alpha=0.35)

    # (b) Feature alignment bar chart
    ax = axes[1]
    bars = ax.bar(["Attn+aug", "Feature-aligned"], [attn_mae, align_mae], color=[c_target, c_warn], width=0.56)
    ax.axhline(attn_mae, color=c_neutral, linestyle="--", linewidth=1.3)
    degradation = (align_mae - attn_mae) / attn_mae * 100.0
    ax.text(1, align_mae + 1.4, f"+{degradation:.1f}% degradation", color=c_warn, ha="center")
    ax.set_title("(b) Feature alignment fails", weight="bold")
    ax.set_ylabel("DATA325 MAE (cm)")
    ax.set_ylim(0, max(align_mae, attn_mae) + 10)
    ax.grid(True, axis="y", color="#D0D0D0", alpha=0.35)
    for bar in bars:
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.8, f"{bar.get_height():.2f}", ha="center", fontsize=9)

    # (c) DANN training curves
    ax = axes[2]
    ax.axhspan(0.60, 0.70, color="#F4A3A3", alpha=0.40)
    ax.plot(epochs, domain_acc, color=c_source, linewidth=2.2)
    ax.text(0.08, 0.65, "Target range\n60-70%", transform=ax.transAxes, color=c_warn)
    ax.text(0.95, 0.96, "~100%", transform=ax.transAxes, color=c_warn, ha="right", va="top")
    ax.set_title("(c) DANN instability", weight="bold")
    ax.set_xlabel("Epoch")
    ax.set_ylabel("Domain classifier accuracy")
    ax.set_ylim(0.45, 1.05)
    ax.grid(True, color="#D0D0D0", alpha=0.35)

    fig.tight_layout()
    upload_stem = FIG_NAME_BY_STEM["fig7"]
    fig.savefig(FIG_OUT / f"{upload_stem}.png", bbox_inches="tight", pad_inches=0.08, dpi=300)
    fig.savefig(FIG_OUT / f"{upload_stem}.pdf", bbox_inches="tight", pad_inches=0.08)
    plt.close(fig)


def copy_assets() -> None:
    reset_dir(OUT)
    FIG_OUT.mkdir(parents=True, exist_ok=True)
    TABLE_OUT.mkdir(parents=True, exist_ok=True)
    REPRO_OUT.mkdir(parents=True, exist_ok=True)

    for source_stem, upload_stem, _, _ in FIGURES:
        for ext in ("png", "pdf"):
            src = FIG_SRC / f"{source_stem}.{ext}"
            if src.exists():
                shutil.copy2(src, FIG_OUT / f"{upload_stem}.{ext}")

    for src in REPRO_FILES:
        if src.exists():
            if src.name.endswith("_local.json"):
                continue
            try:
                obj = json.loads(src.read_text(encoding="utf-8"))
            except json.JSONDecodeError:
                shutil.copy2(src, REPRO_OUT / src.name)
            else:
                (REPRO_OUT / src.name).write_text(json.dumps(sanitize_json_value(obj), indent=2, ensure_ascii=True) + "\n", encoding="utf-8")

    for src in [
        CEA_EXP_DIR / "roi_quality_metrics.csv",
        CEA_EXP_DIR / "error_taxonomy.csv",
        CEA_EXP_DIR / "seed_retraining_summary.csv",
        CEA_EXP_DIR / "source_morphometric_baseline.csv",
    ]:
        if src.exists():
            shutil.copy2(src, TABLE_OUT / src.name)


def write_csvs(m: dict) -> None:
    with (TABLE_OUT / "table1_ablation.csv").open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Variant", "Feature/intervention", "Source MAE (cm)", "DATA325 MAE (cm)", "DATA325 RMSE (cm)", "DATA325 MAPE (%)", "Note"])
        for row in canonical_ablation_rows(m):
            writer.writerow(
                [
                    row["method"],
                    row["feature"],
                    fmt(row["source_mae"]),
                    fmt(row["data325_mae"]),
                    fmt(row["data325_rmse"]),
                    fmt(row["data325_mape"]),
                    row["note"],
                ]
            )
    with (TABLE_OUT / "table2_height_bins.csv").open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Height bin (cm)", "n", "MAE (cm)", "MAPE (%)"])
        for row in m["height_bins"]:
            writer.writerow([row["bucket"], row["n"], fmt(row["mae"]), fmt(row["mape"])])


def build_graphical_abstract(m: dict) -> Path:
    def c(hex_color: str) -> str:
        return f"#{hex_color}"

    plt.rcParams.update({"font.family": "DejaVu Sans", "font.size": 11})
    fig = plt.figure(figsize=(13.28, 5.31), dpi=300)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_axis_off()
    fig.patch.set_facecolor("white")

    boxes = [
        (0.04, 0.56, 0.19, 0.28, "Source greenhouse", "156 hand-box ROIs\n64-270 cm"),
        (0.29, 0.56, 0.19, 0.28, "Frozen DINOv3", "ViT-L features\n224 x 224 ROI"),
        (0.54, 0.56, 0.19, 0.28, "Attention pooling", "CLS-to-patch weights\n1024D descriptor"),
        (0.79, 0.56, 0.17, 0.28, "DCF head", "camera height + features\nplant height"),
    ]
    for idx, (x, y, w, h, title, body) in enumerate(boxes):
        color = c([BLUE, GREEN, ORANGE, BLUE][idx])
        ax.add_patch(plt.Rectangle((x, y), w, h, fc="#F7F9FB", ec=color, lw=2))
        ax.text(x + w / 2, y + h - 0.075, title, ha="center", va="center", weight="bold", color=color, fontsize=13)
        ax.text(x + w / 2, y + 0.105, body, ha="center", va="center", color=c(BLACK), fontsize=11, linespacing=1.35)
        if idx < len(boxes) - 1:
            ax.annotate("", xy=(x + w + 0.045, y + h / 2), xytext=(x + w + 0.01, y + h / 2), arrowprops=dict(arrowstyle="->", lw=2, color=c(GRAY)))

    ax.text(0.04, 0.91, "Zero-shot external-greenhouse maize height estimation", fontsize=21, weight="bold", color=c(BLACK))
    # Compact result panel
    ax.add_patch(plt.Rectangle((0.05, 0.22), 0.42, 0.22, fc=c(LIGHT_GREEN), ec=c(GREEN), lw=1.5))
    ax.text(0.07, 0.39, "Main gain", weight="bold", color=c(GREEN), fontsize=13)
    ax.text(0.07, 0.335, "CLS baseline: 41.76 cm MAE", color=c(BLACK), fontsize=11.5)
    rev = load_revision_results()
    attn_mae = rev["summary"].get("model_summary", {}).get("attn", {}).get("mae_cm", 30.407656506794254)
    ax.text(0.07, 0.285, f"Attention pooling: {fmt(attn_mae)} cm MAE", color=c(BLACK), fontsize=11.5)
    source_morph = load_revision_results().get("source_morphometric", {})
    rf_mae = source_morph.get("models", {}).get("random_forest", {}).get("summary", {}).get("mae_cm", 27.10317886989101)
    ax.text(0.07, 0.235, f"Source morph RF: {fmt(rf_mae)} cm MAE", color=c(BLACK), fontsize=11.5)

    ax.add_patch(plt.Rectangle((0.53, 0.22), 0.42, 0.22, fc=c(LIGHT_ORANGE), ec=c(ORANGE), lw=1.5))
    ax.text(0.55, 0.39, "Remaining bottleneck", weight="bold", color=c(ORANGE), fontsize=13)
    ax.text(0.55, 0.31, f"Feature-domain centroid distance: {fmt(m['centroid_distance'])}", color=c(BLACK), fontsize=12)
    ax.text(0.55, 0.25, "Plants below 80 cm: 55.11% MAPE", color=c(BLACK), fontsize=12)

    out_png = OUT / "graphical_abstract_non_ai.png"
    out_pdf = OUT / "graphical_abstract_non_ai.pdf"
    out_tif = OUT / "graphical_abstract_non_ai.tif"
    fig.savefig(out_png, bbox_inches="tight", pad_inches=0.03, dpi=300)
    fig.savefig(out_tif, bbox_inches="tight", pad_inches=0.03, dpi=300)
    fig.savefig(out_pdf, bbox_inches="tight", pad_inches=0.03)
    plt.close(fig)
    return out_png


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    add_p(doc, f"Authors: {AUTHOR_NAME}^1, {COAUTHOR_NAME}^1,*")
    add_p(doc, f"^1 {AFFILIATION}")
    add_p(doc, f"*Corresponding author: {COAUTHOR_NAME}, {CORRESPONDING_EMAIL}")
    doc.add_page_break()


def add_ablation_table(doc: Document, m: dict) -> None:
    rows = []
    note_map = {
        "": "",
        "best": "best tested",
        "diagnostic": "diagnostic control",
        "failed": "diagnostic negative",
        "source-trained geometry baseline": "source-trained baseline",
    }
    for row in canonical_ablation_rows(m):
        rows.append(
            [
                row["method"],
                row["feature"],
                fmt(row["source_mae"]),
                fmt(row["data325_mae"]),
                fmt(row["data325_mape"]),
                note_map.get(row["note"], row["note"]),
            ]
        )
    add_caption(doc, "Table 1. Ablation study under the DATA325 zero-shot cross-greenhouse protocol.")
    add_table(
        doc,
        ["Variant", "Feature / intervention", "Source MAE", "DATA325 MAE", "DATA325 MAPE", "Interpretation"],
        rows,
        [2.35, 5.05, 1.45, 1.65, 1.65, 2.85],
    )


def add_height_bin_table(doc: Document, m: dict) -> None:
    rows = [[row["bucket"], str(row["n"]), fmt(row["mae"]), fmt(row["mape"])] for row in m["height_bins"]]
    add_caption(doc, "Table 2. Height-bin error profile for the best Attn+aug+TTA8 model on DATA325.")
    add_table(doc, ["Height bin (cm)", "n", "MAE (cm)", "MAPE (%)"], rows, [4.0, 2.0, 3.0, 3.0])


def build_manuscript(m: dict) -> Path:
    doc = setup_doc()
    add_title_page(doc)

    doc.add_heading("Abstract", level=1)
    abstract = (
        "Real greenhouse images contain maize plants together with pots, benches, labels, irrigation hardware, shadows, and facility-specific backgrounds. "
        "These factors make cross-greenhouse plant-height estimation difficult when a model is trained in one imaging environment and deployed in another. "
        "We evaluated DINOv3-DiffCorn-Fusion (DINOv3-DCF), an ROI-level maize height pipeline that connects frozen DINOv3 visual features to a DiffCorn-Fusion regression head, under a strict zero-shot protocol in which the independent DATA325 greenhouse was not used for training, tuning, or model selection. "
        "The study emphasizes real captured images, manual ROI annotation, crop extraction, attention-based feature pooling, and external error analysis. "
        "Replacing CLS-token pooling with attention-weighted patch aggregation reduced DATA325 MAE from 41.76 to 30.41 cm. "
        "Adding visual augmentation and eight-sample test-time augmentation gave the best tested result, with 29.57 cm MAE, 38.98 cm RMSE, and 36.14% MAPE. "
        "Feature analysis confirmed domain separation, with a t-SNE centroid distance of 36.89 and a 36.72% relative mean shift. "
        "Residual error concentrated in plants below 80 cm, where MAPE reached 55.11%. "
        "Per-image camera-height correction, simple geometric priors, feature-statistic alignment, and the tested DANN configuration did not close the gap. "
        "These results show that attention-guided foundation-model feature aggregation improves zero-shot maize height transfer, while early-stage plants require stage-aware and plant-focused adaptation."
    )
    add_p(doc, abstract)
    add_p(doc, "Keywords: maize phenotyping; plant height; DINOv3; foundation model; zero-shot transfer; domain shift; attention pooling")

    doc.add_heading("1. Introduction", level=1)
    for para in [
        "High-throughput crop phenotyping is a core enabling technology for breeding, controlled-environment experimentation, and digital crop management. Plant height is frequently measured because it reflects developmental stage, biomass accumulation, lodging risk, irrigation response, and management status. Image-based crop-height systems have been developed with UAS structure-from-motion products, crop height models, stereo vision, and protected-facility machine vision, showing that automated height measurement can reduce manual labor while preserving biologically useful growth information (Chang et al., 2017; Xie et al., 2021; Kim et al., 2021; Jayasuriya et al., 2024).",
        "Agricultural computer vision has also moved from hand-engineered descriptors toward deep learning, object detection, and phenotyping-specific visual pipelines (Patricio and Rieder, 2018; Kamilaris and Prenafeta-Boldu, 2018; Li et al., 2020; Ariza-Sentis et al., 2024). For maize and other crops, vision systems have been used for crop counting, crop evaluation, growth-stage estimation, and decision-support inputs (Li et al., 2019; Veramendi and Cruvinel, 2024; Che et al., 2024). These studies demonstrate the value of visual phenotyping, but they also show that imaging geometry, field or greenhouse environment, and data-processing assumptions strongly affect downstream trait estimates.",
        "Cross-domain generalization is therefore a persistent obstacle in agricultural computer vision. Multi-environment phenotyping studies explicitly account for spatial heterogeneity and environmental variation, while broader machine-learning benchmarks show that distribution shift can cause large performance drops when training and deployment data differ (Che et al., 2024; Koh et al., 2021; Gulrajani and Lopez-Paz, 2021). Greenhouse images contain plant tissue together with pots, benches, labels, substrate, shadows, and equipment. These context cues vary between facilities and can be strongly encoded by deep representations even when they are unrelated to plant height.",
        "ROI quality and plant-structure representation are additional bottlenecks. Precision-farming detection systems must separate crop targets from cluttered agricultural scenes, and robot-localization studies report failures caused by lighting variation, complex backgrounds, and small target regions (Ariza-Sentis et al., 2024; Xing et al., 2023). Segmentation foundation models such as Segment Anything provide a promising route to plant-mask normalization, but they do not by themselves solve trait regression or cross-greenhouse representation shift (Kirillov et al., 2023). Three-dimensional and hyperspectral phenotyping reviews and datasets further emphasize that structural traits require high-quality annotated data and explicit geometric or part-level reasoning (Liu et al., 2020; Reena et al., 2025).",
        "Recent self-supervised vision transformers provide robust general image features and have become attractive backbones for phenotyping pipelines (Dosovitskiy et al., 2021; Caron et al., 2021; Oquab et al., 2024; Simeoni et al., 2025). However, the standard CLS token is a global summary representation. In a plant ROI, it can integrate background and acquisition context together with plant structure. For plant-height estimation, patch-level aggregation may be more appropriate because plant tissue occupies only part of the crop, especially at early growth stages.",
        "This study evaluates DINOv3-DiffCorn-Fusion (DINOv3-DCF), a pipeline that connects frozen DINOv3 features to a DCF head that predicts a phytomer-parameter representation and plant height. The central question is whether attention-weighted patch aggregation improves zero-shot cross-greenhouse maize height estimation compared with CLS pooling and patch-mean pooling. We also test several plausible explanations for the remaining error, including camera-height metadata, simple geometry, feature-statistic alignment, and domain-adversarial learning (Sun and Saenko, 2016; Ganin et al., 2016).",
        "The contribution is an original cross-greenhouse evaluation of foundation-model feature aggregation for maize height phenotyping. Unlike UAS/SfM, stereo-vision, or protected-facility geometry pipelines, this work asks whether frozen foundation-model features can transfer across greenhouse domains when only source-domain labels are available. We report both successful and unsuccessful interventions because negative diagnostic results are important for identifying the actual bottleneck. The CEA-oriented interpretation is not that the method fully solves external maize height estimation, but that attention-weighted pooling is a strong representation improvement and that early-stage plants remain the key domain-transfer target.",
    ]:
        add_p(doc, para)

    doc.add_heading("2. Materials and methods", level=1)
    doc.add_heading("2.1. Study design and datasets", level=2)
    for para in [
        "Model development used a source-domain hand-bounding-box dataset and external testing used the independent DATA325 target greenhouse. Target-domain annotations were restricted to final evaluation and diagnostic visualization. They were not used for training, hyperparameter tuning, model selection, or early stopping. The reported DATA325 results therefore represent zero-shot cross-greenhouse transfer.",
        "All compared DATA325 evaluations used the same 82 manually annotated plant boxes, the same plant-height labels, and the same image files. Manual ROIs were used because object detection, tracking, and segmentation can introduce their own errors under complex agricultural backgrounds (Ariza-Sentis et al., 2024; Xing et al., 2023; Kirillov et al., 2023). This design isolates the height-regression and representation-transfer problem from automatic detector quality. Per-image camera height was resolved from the capture-height mapping file rather than from a single global default. The correction changed DATA325 MAE by only about 0.03 cm, indicating that camera-height metadata mismatch was not the main source of error.",
        "The source-domain hand-box dataset contains 156 samples. Plant heights range from 64.00 to 270.00 cm, with mean 140.50 cm and standard deviation 66.30 cm. Source camera heights include 106 and 140 cm. DATA325 contains 75 original images; 25 completed images currently provide 82 evaluated boxes. DATA325 height labels range from 30.00 to 178.00 cm, with median 97.00 cm and mean 102.79 cm. DATA325 contains a larger fraction of plants below 80 cm than the source set.",
    ]:
        add_p(doc, para)
    add_p(doc, "Figure 1 makes the evaluation protocol auditable by showing source-domain hand-box images, independent DATA325 target images with manual boxes, the extracted ROI crops, and the exact height-prediction records used in the external evaluation.")
    add_named_figure(doc, "fig_real_protocol")
    add_named_figure(doc, "fig3")

    doc.add_heading("2.2. DINOv3-DCF architecture", level=2)
    for para in [
        "The pipeline receives a plant image and a region of interest (ROI). DATA325 ROIs are manually drawn in the current benchmark to isolate the height-regression problem from detector errors. In a deployed greenhouse system, the same ROI-level pipeline can be connected to an automatic detector or segmentation model, but detector quality is outside the scope of this controlled zero-shot evaluation.",
        "Each ROI is resized to 224 x 224 pixels and passed through a frozen DINOv3 ViT-L backbone. The visual output has a 1024-dimensional hidden representation. The DCF head concatenates this visual descriptor with one scalar camera-height input, forming a 1025-dimensional input vector. The head outputs a 64-dimensional phytomer representation, interpreted as 16 phytomers with four parameters each. Plant height is computed from internode-length-related components of this representation.",
        "The DINOv3 backbone remains frozen in all main experiments. This isolates the contribution of feature aggregation and DCF training, reduces computational cost, and improves reproducibility. The DCF head is the trainable module for height estimation, except in the DANN experiment where an additional domain classifier is trained for adversarial domain discrimination.",
    ]:
        add_p(doc, para)
    add_named_figure(doc, "fig1")

    doc.add_heading("2.3. Feature aggregation modes", level=2)
    for para in [
        "Three feature aggregation modes were evaluated. The CLS baseline uses the first token of the final hidden state. Patch-mean aggregation averages all patch tokens and removes the explicit CLS bottleneck. Attention-weighted aggregation uses the final-layer CLS-to-patch attention map: attention weights from the CLS token to patch tokens are averaged across heads, normalized over patches, and used to compute a weighted sum of final-layer patch tokens.",
        "Attention-weighted aggregation keeps the feature dimension unchanged at 1024, so the DCF input interface remains compatible with the original 1025-dimensional design. The change is localized to feature extraction. This experimental control means that differences between CLS, patch-mean, and attention-weighted models can be attributed primarily to feature aggregation rather than to a redesigned regressor.",
        "The biological motivation is that a maize ROI contains plant tissue as well as pot edges, bench structures, substrate, labels, and background objects. The CLS token can summarize all of these signals, while attention-weighted patch aggregation reduces the contribution of less salient patches. It is not a segmentation mask, but it is a lightweight mechanism for suppressing background-dominated regions without training a new visual encoder.",
    ]:
        add_p(doc, para)
    add_named_figure(doc, "fig2")

    doc.add_heading("2.4. Training variants and inference", level=2)
    for para in [
        "The old baseline used the original CLS-only feature bundle and the original DCF checkpoint. To separate representation effects from training artifacts, CLS retrain, patch-mean retrain, and attention-weighted retrain models were built from newly extracted bundles using the same sample identities, bounding boxes, camera heights, and height labels. Only the visual feature values changed across these three bundles.",
        "The attention-augmentation model used attention-weighted features and training-time visual perturbations. The augmentation strategy combined background and color-domain perturbations to reduce sensitivity to greenhouse-specific appearance. Background augmentation perturbed peripheral ROI regions, while color augmentation sampled brightness, contrast, saturation, and hue changes. The DINOv3 backbone remained frozen.",
        "Test-time augmentation (TTA8) was applied only at inference. For each ROI, one prediction used the original crop and seven predictions used independent color perturbations. The reported prediction is the mean of the eight outputs. Camera height was held constant across the TTA samples.",
    ]:
        add_p(doc, para)

    doc.add_heading("2.5. Domain-shift diagnostics", level=2)
    for para in [
        "We tested several alternative explanations for the external error. First, DATA325 camera-height inputs were corrected using per-image capture-height mapping. Second, two geometry routes were evaluated: attention-derived top-bottom localization and the ratio of bounding-box pixel height to image height multiplied by camera height. Third, source and target feature statistics were computed for attention-weighted features, and target features were aligned to source mean and standard deviation by per-dimension standardization, following the general motivation of correlation-alignment methods (Sun and Saenko, 2016). Fourth, a DANN-style domain-adversarial model was trained with source labels and unlabeled target images (Ganin et al., 2016).",
        "These experiments were designed as diagnostic controls rather than as exhaustive optimization of each adaptation family. They identify whether the observed error is primarily metadata-driven, geometry-driven, marginal feature-distribution-driven, or reducible through the tested adversarial setup. This framing follows the broader distribution-shift literature, where source-domain accuracy alone is not sufficient evidence of deployment robustness (Koh et al., 2021; Gulrajani and Lopez-Paz, 2021).",
    ]:
        add_p(doc, para)

    doc.add_heading("2.6. Evaluation metrics", level=2)
    add_p(doc, "Evaluation uses mean absolute error (MAE), root mean squared error (RMSE), and mean absolute percentage error (MAPE). MAPE is computed at the sample level as |prediction - ground truth| / ground truth x 100% and then averaged. DATA325 results are grouped into height bins of 0-80, 80-100, 100-120, 120-140, 140-160, and 160-180 cm to locate stage-dependent error.")

    doc.add_heading("3. Results", level=1)
    doc.add_heading("3.1. DATA325 differs from the source domain in height distribution and feature space", level=2)
    add_p(doc, "DATA325 differs from the source domain in both greenhouse appearance and plant-stage composition. The target set has a larger fraction of plants below 80 cm, whereas the source-domain training set extends to taller plants. This creates a difficult evaluation condition because the model must transfer across both visual background and stage distribution.")
    add_p(doc, f"Feature analysis confirmed a measurable domain gap. A t-SNE visualization using 20 source and 20 DATA325 ROI features showed a centroid distance of {fmt(m['centroid_distance'])}. In high-dimensional feature statistics, the relative mean shift was {fmt(m['mean_shift'])}% and the relative standard-deviation shift was {fmt(m['std_shift'])}%. These results support the interpretation that source and target domains are separated in DINOv3 feature space.")
    add_named_figure(doc, "fig4")

    doc.add_heading("3.2. Attention-weighted aggregation is the most effective single change", level=2)
    add_p(doc, "The original CLS baseline produced 41.76 cm MAE on DATA325. Retraining a CLS model did not improve external performance; it increased MAE to 45.86 cm. Patch-mean aggregation reduced MAE to 33.53 cm, indicating that moving away from the CLS token is beneficial. Attention-weighted patch aggregation produced the largest single improvement, reducing MAE to 30.41 cm while maintaining low source-domain MAE.")
    add_p(doc, "Visual augmentation and TTA8 provided additional but smaller gains. Attention plus augmentation reached 29.89 cm MAE, and attention plus augmentation plus TTA8 reached 29.57 cm MAE, 38.98 cm RMSE, and 36.14% MAPE. The improvement pattern shows a strong initial gain from representation change followed by diminishing returns from augmentation and test-time averaging.")
    add_named_figure(doc, "fig5")
    add_ablation_table(doc, m)

    doc.add_heading("3.3. Remaining error concentrates in early-stage plants", level=2)
    add_p(doc, "Errors were not uniformly distributed across height ranges. In the 0-80 cm bin, MAPE was 55.11%, much higher than in taller bins. For the 120-140, 140-160, and 160-180 cm bins, relative errors were lower, suggesting that larger plants provide more stable ROI structure and more reliable attention patterns.")
    add_p(doc, "This pattern is consistent with visual inspection. Low-stage ROIs often contain a high fraction of pot, substrate, bench, labels, and background relative to plant tissue. Even when the attention map highlights plant-like areas, the representation contains less structural evidence for height. Taller plants fill more of the ROI, making the visual descriptor more stable under cross-domain changes.")
    add_named_figure(doc, "fig6")
    add_height_bin_table(doc, m)
    add_p(doc, "To connect the quantitative height-bin result to the actual imaging conditions, Figure 8 groups real DATA325 photographs by growth stage and absolute-error rank. The high-error examples show that early plants and cluttered backgrounds are not abstract statistical labels; they are visible target-domain image conditions.")
    add_named_figure(doc, "fig_real_stage_error_gallery")

    doc.add_heading("3.4. Diagnostic experiments reject several simple explanations", level=2)
    add_p(doc, f"The geometry experiments did not meet the threshold for a usable prior. The bounding-box-ratio prior reached Pearson r={fmt(m['bbox_r'], 3)} and MAE {fmt(m['bbox_mae'])} cm. Attention-map thresholding and vertical projection were weaker, with r={fmt(m['attn_threshold_r'], 3)} and r={fmt(m['attn_projection_r'], 3)}, respectively. These results show that the available camera-height definition and ROI geometry do not form a closed physical height model.")
    add_p(doc, "Feature-statistic alignment also failed. Although the source and target domains had measurable mean and variance shifts, mapping target features to source mean and standard deviation worsened DATA325 MAE to 45.98 cm. This indicates that the domain difference is structural rather than a purely marginal channel-wise shift.")
    add_p(doc, f"The tested DANN configuration did not improve target-domain performance. Its source-domain MAE remained acceptable at 2.45 cm, but DATA325 MAE was 34.20 cm. The domain classifier converged to approximately {fmt(m['dann_final_acc'] * 100, 0)}% accuracy, meaning source and target representations remained easily separable. In this setting, adversarial training did not produce the intended domain confusion.")
    add_named_figure(doc, "fig7")

    doc.add_heading("3.5. Qualitative DATA325 evidence supports the stage-dependent error pattern", level=2)
    add_p(doc, "Real DATA325 examples show the same pattern as the height-bin analysis. Low-stage plants often have sparse plant pixels and dispersed attention, and their prediction errors can exceed 40 cm. High-stage examples fill more of the ROI and show more concentrated attention, with lower relative errors. The qualitative evidence does not replace quantitative evaluation, but it explains why the 0-80 cm group remains the main bottleneck.")
    add_named_figure(doc, "fig8")

    doc.add_heading("4. Discussion", level=1)
    for para in [
        "The main result is that feature aggregation in a frozen foundation-model backbone strongly affects cross-greenhouse phenotyping transfer. CLS-only features are convenient, but they are a global image summary and can encode greenhouse context. Attention-weighted patch aggregation is a small post-processing change, yet it produced the largest reduction in DATA325 MAE. This supports the view that plant-focused pooling is more important than simply retraining the same CLS-based regressor.",
        "From a Computers and Electronics in Agriculture perspective, the result is useful because it identifies a practical representation-level intervention for agricultural computer vision systems that must operate across facilities. Earlier CEA crop-height studies often recover height through UAS/SfM crop-height models, stereo geometry, or facility-specific machine vision (Chang et al., 2017; Xie et al., 2021; Kim et al., 2021; Jayasuriya et al., 2024). The present work does not try to replace those geometry-rich systems. Instead, it tests whether frozen foundation-model features can support zero-shot transfer from one greenhouse to another when only ROI images and camera-height metadata are available.",
        "This distinction is important for deployment. Geometry-based systems can be accurate when the imaging platform, terrain or bench reference, and calibration assumptions are well controlled. ROI-level greenhouse images are cheaper to process, but they expose the model to background, pot, and growth-stage shifts. The current results show that attention-weighted patch pooling reduces this representation problem, but they also show that frozen features alone do not fully recover a physically stable height signal in early-stage plants.",
        "The DATA325 benchmark is small, but it has a specific diagnostic role. Multi-environment phenotyping work has shown that environmental and spatial variation can distort growth-stage estimates and cultivar comparisons (Che et al., 2024). Similarly, the present target set reveals a cross-greenhouse failure mode that source-domain validation would miss. The value of the dataset is therefore not only its sample size, but its ability to expose a realistic deployment shift with fixed images, fixed manual ROIs, and reproducible evaluation JSON files.",
        "The current architecture nevertheless has a clear performance ceiling on DATA325. Even the best variant remains near 30 cm MAE, which is insufficient for precise agronomic measurement in early growth stages. The main contribution is therefore a controlled cross-domain benchmark and an empirically validated representation improvement, not a claim of solved greenhouse height estimation.",
        "The diagnostic experiments clarify the method boundary. Camera-height correction ruled out a metadata-driven explanation. Attention and bounding-box geometry ruled out a simple height-ratio solution. Feature-statistic alignment ruled out a purely marginal distribution shift. The tested DANN result showed that this adversarial setup was insufficient to align domains. Together, these findings support a structural, stage-dependent interpretation of the remaining shift.",
        "Recent CEA work on multi-environment crop monitoring provides useful context. Che et al. (2024) showed that UAV-based soybean growth-stage estimation requires explicit spatial-heterogeneity and climate corrections across environments. Jayasuriya et al. (2024) reported that protected-facility height estimation is sensitive to camera angle, plant architecture, and background complexity. The present results are consistent with these findings: cross-greenhouse transfer is not only a sensor-calibration problem but also a representation problem that depends on how visual features aggregate plant and background information.",
        "Future work should focus on the low-stage regime, where the current best model still reaches 55.11% MAPE. A concrete target is reducing 0-80 cm MAPE below 30% while maintaining performance in taller bins. Promising directions include stage-aware models, segmentation-guided ROI normalization, explicit plant-mask pooling, target-domain self-supervised adaptation, and multi-greenhouse pretraining. Segment Anything-style masks and precision-farming detection pipelines could help reduce pot and background contamination before feature extraction (Kirillov et al., 2023; Ariza-Sentis et al., 2024). Plant-mask pooling and 3D structural priors should also be evaluated, because plant phenotyping studies increasingly emphasize part-level annotation and structural data for reliable trait extraction (Liu et al., 2020; Reena et al., 2025). A better DANN design may also be possible, but the present result suggests that source ROI features and target images must be matched more carefully before adversarial alignment can be useful.",
    ]:
        add_p(doc, para)

    doc.add_heading("4.1. Limitations", level=2)
    for para in [
        "The DATA325 evaluation set contains 82 annotated boxes from 25 completed images. This is sufficient to reveal a strong cross-domain failure mode, but larger target-domain collections are needed to estimate performance across cultivars, developmental stages, camera configurations, and greenhouse layouts.",
        "The current benchmark evaluates manually drawn DATA325 boxes to isolate DCF regression from detector errors. A complete greenhouse deployment will require robust automatic ROI detection and quality control.",
        "DINOv3 remains frozen. This improves reproducibility and controls the experiment, but it limits the ability to learn plant-specific invariances. Future studies should compare frozen features with carefully regularized fine-tuning or adapter-based tuning across multiple greenhouses.",
    ]:
        add_p(doc, para)

    doc.add_heading("5. Conclusions", level=1)
    add_p(doc, "We evaluated DINOv3-DCF for maize plant-height phenotyping under a strict zero-shot cross-greenhouse protocol. Attention-weighted patch aggregation reduced DATA325 MAE from 41.76 to 30.41 cm, and the best Attn+aug+TTA8 variant reached 29.57 cm. The remaining error is dominated by plants below 80 cm. Camera-height correction, simple geometry, feature-statistic alignment, and the tested DANN configuration did not resolve the gap. These results define a practical representation improvement and a clear next target: stage-aware, plant-focused adaptation for early maize growth.")

    doc.add_heading("CRediT authorship contribution statement", level=1)
    add_p(doc, f"{AUTHOR_NAME}: Conceptualization, Methodology, Software, Validation, Formal analysis, Investigation, Data curation, Writing - original draft, Writing - review & editing, and Visualization.")
    doc.add_heading("Declaration of competing interest", level=1)
    add_p(doc, "The authors declare no competing interests.")
    doc.add_heading("Funding", level=1)
    add_p(doc, "No external funding was reported for this work.")
    doc.add_heading("Acknowledgements", level=1)
    add_p(doc, "None.")
    doc.add_heading("Data availability", level=1)
    add_p(doc, f"DATA325 raw images, manual bounding boxes, plant-height labels, camera-height mapping, evaluation outputs, selected DCF checkpoints, figure-generation scripts, and reproducibility notes have been organized for public release in the accompanying open-source repository: {REPOSITORY_URL}. The archive file prepared for release is {OPEN_RELEASE_ARCHIVE}.")
    doc.add_heading("Declaration of generative AI and AI-assisted technologies in the writing process", level=1)
    add_p(doc, "During preparation of this work the authors used AI-assisted writing and formatting tools to help draft and organize text. After using these tools, the authors reviewed and edited the content as needed and take full responsibility for the content of the publication. No generative AI was used to create or modify the submitted scientific figures, graphical abstract, or result images.")

    doc.add_heading("References", level=1)
    refs = [
        "Araus, J.L., Cairns, J.E., 2014. Field high-throughput phenotyping: the new crop breeding frontier. Trends in Plant Science 19, 52-61.",
        "Ariza-Sentis, M., Velez, S., Martinez-Pena, R., Baja, H., Valente, J., 2024. Object detection and tracking in Precision Farming: a systematic review. Computers and Electronics in Agriculture 219, 108757. doi:10.1016/j.compag.2024.108757.",
        "Caron, M., Touvron, H., Misra, I., Jegou, H., Mairal, J., Bojanowski, P., Joulin, A., 2021. Emerging properties in self-supervised vision transformers. Proceedings of the IEEE/CVF International Conference on Computer Vision, 9650-9660.",
        "Chang, A., Jung, J., Maeda, M.M., Landivar, J., 2017. Crop height monitoring with digital imagery from Unmanned Aerial System (UAS). Computers and Electronics in Agriculture 141, 232-237. doi:10.1016/j.compag.2017.07.008.",
        "Che, Y., Gu, Y., Bai, D., Li, D., Li, J., Zhao, C., Wang, Q., Qiu, H., et al., 2024. Accurately estimate soybean growth stages from UAV imagery by accounting for spatial heterogeneity and climate factors across multiple environments. Computers and Electronics in Agriculture 225, 109313. doi:10.1016/j.compag.2024.109313.",
        "Dosovitskiy, A., Beyer, L., Kolesnikov, A., et al., 2021. An image is worth 16x16 words: transformers for image recognition at scale. International Conference on Learning Representations.",
        "Ganin, Y., Ustinova, E., Ajakan, H., et al., 2016. Domain-adversarial training of neural networks. Journal of Machine Learning Research 17, 1-35.",
        "Gulrajani, I., Lopez-Paz, D., 2021. In search of lost domain generalization. International Conference on Learning Representations.",
        "Jayasuriya, N., Guo, Y., Hu, W., Ghannoum, O., 2024. Machine vision based plant height estimation for protected crop facilities. Computers and Electronics in Agriculture 218, 108669. doi:10.1016/j.compag.2024.108669.",
        "Kamilaris, A., Prenafeta-Boldu, F.X., 2018. Deep learning in agriculture: a survey. Computers and Electronics in Agriculture 147, 70-90. doi:10.1016/j.compag.2018.02.016.",
        "Kim, W., Lee, D., Kim, Y., Kim, T., Lee, W., Choi, C., 2021. Stereo-vision-based crop height estimation for agricultural robots. Computers and Electronics in Agriculture 181, 105937. doi:10.1016/j.compag.2020.105937.",
        "Kirillov, A., Mintun, E., Ravi, N., Mao, H., Rolland, C., Gustafson, L., Xiao, T., Whitehead, S., et al., 2023. Segment anything. arXiv:2304.02643.",
        "Koh, P.W., Sagawa, S., Marklund, H., Xie, S.M., Zhang, M., Balsubramani, A., Hu, W., Yasunaga, M., et al., 2021. WILDS: a benchmark of in-the-wild distribution shifts. Proceedings of the 38th International Conference on Machine Learning 139, 5637-5664.",
        "Li, L., Zhang, Q., Huang, D., 2014. A review of imaging techniques for plant phenotyping. Sensors 14, 20078-20111.",
        "Li, M., Sui, R., Meng, Y., Yan, H., 2019. A real-time fuzzy decision support system for alfalfa irrigation. Computers and Electronics in Agriculture 163, 104870. doi:10.1016/j.compag.2019.104870.",
        "Li, Z., Guo, R., Li, M., Chen, Y., Li, G., 2020. A review of computer vision technologies for plant phenotyping. Computers and Electronics in Agriculture 176, 105672. doi:10.1016/j.compag.2020.105672.",
        "Liu, H., Bruning, B., Garnett, T., Berger, B., 2020. Hyperspectral imaging and 3D technologies for plant phenotyping: from satellite to close-range sensing. Computers and Electronics in Agriculture 175, 105621. doi:10.1016/j.compag.2020.105621.",
        "Oquab, M., Darcet, T., Moutakanni, T., et al., 2024. DINOv2: learning robust visual features without supervision. Transactions on Machine Learning Research.",
        "Patricio, D.I., Rieder, R., 2018. Computer vision and artificial intelligence in precision agriculture for grain crops: a systematic review. Computers and Electronics in Agriculture 153, 69-81. doi:10.1016/j.compag.2018.08.001.",
        "Reena, Doonan, J.H., Williams, K., Corke, F.M.K., Zhang, H., Batke, S., Liu, Y., 2025. Wheat3D PartNet: annotated dataset for 3D wheat part segmentation. Computers and Electronics in Agriculture 238, 110697. doi:10.1016/j.compag.2025.110697.",
        "Shorten, C., Khoshgoftaar, T.M., 2019. A survey on image data augmentation for deep learning. Journal of Big Data 6, 60.",
        "Simeoni, O., Vo, H.V., Seitzer, M., et al., 2025. DINOv3. arXiv:2508.10104. doi:10.48550/arXiv.2508.10104. Official source repository: https://github.com/facebookresearch/dinov3. Accessed 28 Apr 2026.",
        "Sun, B., Saenko, K., 2016. Deep CORAL: correlation alignment for deep domain adaptation. European Conference on Computer Vision Workshops, 443-450. doi:10.1007/978-3-319-49409-8_35.",
        "Ubbens, J.R., Stavness, I., 2017. Deep Plant Phenomics: a deep learning platform for complex plant phenotyping tasks. Frontiers in Plant Science 8, 1190.",
        "Veramendi, W., Cruvinel, P., 2024. Method for maize plants counting and crop evaluation based on multispectral images analysis. Computers and Electronics in Agriculture 216, 108470. doi:10.1016/j.compag.2023.108470.",
        "Wang, D., Shelhamer, E., Liu, S., Olshausen, B., Darrell, T., 2021. Tent: fully test-time adaptation by entropy minimization. International Conference on Learning Representations.",
        "Xie, T., Li, J., Yang, C., Jiang, Z., Chen, Y., Guo, L., Zhang, J., 2021. Crop height estimation based on UAV images: methods, errors, and strategies. Computers and Electronics in Agriculture 185, 106155. doi:10.1016/j.compag.2021.106155.",
        "Xing, Z., Zhang, Z., Shi, R., Guo, Q., Zeng, C., 2023. Filament-necking localization method via combining improved PSO with rotated rectangle algorithm for safflower-picking robots. Computers and Electronics in Agriculture 215, 108464. doi:10.1016/j.compag.2023.108464.",
    ]
    for ref in refs:
        add_p(doc, ref)

    out = OUT / "manuscript_cea.docx"
    elsevier = OUT / "manuscript_cea_elsevier_style.docx"
    doc.save(out)
    doc.save(elsevier)
    build_up_word_equations(out)
    build_up_word_equations(elsevier)
    return out


def build_highlights() -> Path:
    highlights = [
        "Real greenhouse ROIs test zero-shot maize height transfer.",
        "Attention pooling reduced DATA325 MAE from 41.76 to 30.41 cm.",
        "The best Attn+aug+TTA8 variant reached 29.57 cm MAE on DATA325.",
        "Errors concentrated below 80 cm, where MAPE reached 55.11%.",
        "Geometry, feature-statistic alignment, and DANN did not close the gap.",
    ]
    for item in highlights:
        if len(item) > 85:
            raise ValueError(f"Highlight exceeds Elsevier 85-character guidance: {item}")
    doc = setup_doc()
    doc.add_heading("Highlights", level=1)
    for item in highlights:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.add_run(f"- {item}")
    out = OUT / "highlights.docx"
    doc.save(out)
    (OUT / "highlights.txt").write_text("\n".join(f"- {item}" for item in highlights) + "\n", encoding="utf-8")
    return out


def build_cover_letter() -> Path:
    doc = setup_doc()
    doc.add_heading("Cover letter", level=1)
    paragraphs = [
        "Dear Editors,",
        f"I submit the manuscript entitled \"{TITLE}\" for consideration as an Original research paper in Computers and Electronics in Agriculture.",
        "The manuscript addresses a practical agricultural computer vision problem: maize height estimation that transfers from one greenhouse to another without target-domain training labels. It is positioned against recent Computers and Electronics in Agriculture work on crop-height estimation, protected-crop machine vision, plant phenotyping, object detection, and multi-environment visual phenotyping. The study combines frozen DINOv3 visual features with a DiffCorn-Fusion regression head and uses real greenhouse images, manual ROI crops, attention visualization, and external error galleries to show that attention-weighted patch aggregation substantially improves zero-shot external performance.",
        "The main result is that attention-weighted pooling reduced DATA325 MAE from 41.76 to 30.41 cm, and the best attention plus augmentation plus TTA8 variant reached 29.57 cm MAE. The manuscript also reports diagnostic negative results: camera-height correction, simple geometry, feature-statistic alignment, and the tested DANN setup did not close the remaining domain gap. These controls identify early-stage plants below 80 cm as the primary bottleneck.",
        "The work fits the scope of Computers and Electronics in Agriculture because it evaluates a computational imaging method for crop phenotyping, quantifies cross-greenhouse domain shift, and provides reproducible evaluation assets for future agricultural computer vision research.",
        "The revised figure set contains ten main figures, including newly generated real-image protocol and DATA325 stage-error galleries. All figures and the optional graphical abstract in this CEA package are generated from real images, measured results, or deterministic Python/vector plotting. Generative AI was not used to create or modify submitted scientific figures, graphical abstract, or result images.",
        f"DATA325 raw images, bounding boxes, height labels, evaluation outputs, selected DCF checkpoints, and figure-generation scripts have been organized for public release in the accompanying open-source repository: {REPOSITORY_URL}. The archive file prepared for release is {OPEN_RELEASE_ARCHIVE}. All quantitative results reported in the manuscript can be independently reproduced from the released evaluation JSON files and scripts.",
        "This manuscript is original and is not under consideration elsewhere.",
        "Sincerely,",
        f"{AUTHOR_NAME}",
        f"{AFFILIATION}",
        CORRESPONDING_EMAIL,
    ]
    for para in paragraphs:
        add_cover_p(doc, para)
    out = OUT / "cover_letter_cea.docx"
    doc.save(out)
    return out


def build_supplement(m: dict) -> Path:
    doc = setup_doc()
    doc.add_heading("Supplementary material", level=1)
    add_p(doc, f"Supplementary material for: {TITLE}")
    doc.add_heading("S1. Released and reproducible assets", level=2)
    for item in [
        "DATA325 raw greenhouse images and manually annotated boxes.",
        "Plant-height labels and per-image camera-height mapping used for corrected evaluation.",
        "Evaluation JSON files for CLS, patch-mean, attention-weighted, augmented, TTA8, feature-alignment, and DANN variants.",
        "Figure-generation scripts and non-AI figure assets.",
        "Main tables as editable manuscript tables and CSV sidecars.",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.add_run(f"- {item}")

    doc.add_heading("S2. Feature extraction modes", level=2)
    add_p(doc, "CLS mode uses the first token from the final hidden state. Patch-mean mode averages all patch tokens. Attention-weighted mode averages final-layer CLS-to-patch attention across heads, normalizes over patch positions, and uses the resulting weights to aggregate patch tokens.")

    doc.add_heading("S3. Diagnostic negative results", level=2)
    for item in [
        f"Bounding-box-ratio geometry reached Pearson r={fmt(m['bbox_r'], 3)} and MAE={fmt(m['bbox_mae'])} cm.",
        f"Attention-map thresholding and vertical projection reached r={fmt(m['attn_threshold_r'], 3)} and r={fmt(m['attn_projection_r'], 3)}, respectively.",
        "Feature-statistic alignment worsened DATA325 performance to 45.98 cm MAE.",
        f"The DANN domain classifier finished near {fmt(m['dann_final_acc'] * 100, 0)}% accuracy, indicating unsuccessful domain confusion.",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.add_run(f"- {item}")

    doc.add_heading("S4. Literature-positioning notes", level=2)
    for item in [
        "The revised manuscript cites CEA crop-height studies based on UAS/SfM, stereo vision, and protected-facility machine vision to clarify that DINOv3-DCF targets ROI-level zero-shot representation transfer rather than explicit 3D reconstruction.",
        "The revised manuscript cites CEA review and detection papers to motivate agricultural computer vision, ROI quality, and background clutter as deployment constraints.",
        "The revised manuscript cites distribution-shift and domain-generalization work to frame DATA325 as an external diagnostic benchmark rather than a source-domain validation split.",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.add_run(f"- {item}")

    doc.add_heading("S5. Additional DATA325 visual evidence", level=2)
    add_p(doc, "Two additional real-data visual figures from the Plant Methods preparation are retained outside the main CEA manuscript as optional supplementary evidence. They show the DATA325 annotation protocol and stage-wise error gallery. They are not generated or modified by generative AI.")
    extra = ROOT / "plant_methods_submission_final_public_tone" / "figures_extra"
    for src_name, dst_name, caption in [
        ("fig9_data325_protocol.png", "figS1_data325_protocol.png", "Supplementary Fig. S1. DATA325 annotation and external evaluation protocol."),
        ("fig10_error_pattern_gallery.png", "figS2_error_pattern_gallery.png", "Supplementary Fig. S2. Error-pattern gallery across growth stages."),
    ]:
        src = extra / src_name
        if src.exists():
            dst = OUT / "supplementary_figures"
            dst.mkdir(exist_ok=True)
            shutil.copy2(src, dst / dst_name)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(dst / dst_name), width=Inches(5.9))
            add_caption(doc, caption)

    out = OUT / "supplementary_material.docx"
    doc.save(out)
    return out


def write_sidecars(m: dict) -> None:
    (OUT / "data_availability_statement.md").write_text(
        textwrap.dedent(
            """\
            # Data availability statement

            DATA325 raw images, manual bounding boxes, plant-height labels, camera-height mapping,
            evaluation JSON files, selected DCF checkpoints, figure-generation scripts, and
            reproducibility notes have been organized for public release in the accompanying
            open-source repository: https://github.com/sanyueuy/dinov3-dcf-maize-height.
            The archive file prepared for release is dinov3_dcf_maize_height_open_release_v0_1_0.zip.

            If a Zenodo DOI is minted later, the DOI can be added before journal submission.
            The DINOv3 foundation-model weights are not redistributed and must be obtained under
            their upstream license.
            """
        ),
        encoding="utf-8",
    )
    (OUT / "submission_checklist_cea.md").write_text(
        textwrap.dedent(
            """\
            # Computers and Electronics in Agriculture submission checklist

            ## Included in this package
            - manuscript_cea.docx
            - highlights.docx and highlights.txt
            - cover_letter_cea.docx
            - supplementary_material.docx
            - graphical_abstract_non_ai.png/pdf (optional upload)
            - figures/Figure_1...Figure_10 as PNG/PDF
            - tables/table1_ablation.csv and table2_height_bins.csv
            - reproducibility_json/*.json
            - data_availability_statement.md
            - submission_asset_manifest.md

            ## Still required before online submission
            - Optional ORCID ID for Hong Wu
            - Optional archive DOI if a Zenodo release is minted
            - Author approval of AI-assisted writing disclosure
            - Final check in Elsevier submission system for article-type-specific fields
            - Confirm whether "No external funding was reported" is the final funding statement

            ## CEA-specific handling
            - Article type is prepared as Original research paper.
            - Abstract is non-structured and below 250 words.
            - Highlights are provided as a separate file.
            - Figures are non-AI assets generated from real images, measured results, or deterministic plotting.
            - Main figures include real-image DATA325 annotation, ROI, and error-gallery panels.
            - Tables are editable in the manuscript and also provided as CSV sidecars.
            """
        ),
        encoding="utf-8",
    )

    manifest_lines = [
        "# CEA submission asset manifest",
        "",
        "## Main files",
        "- manuscript_cea.docx",
        "- highlights.docx",
        "- highlights.txt",
        "- cover_letter_cea.docx",
        "- supplementary_material.docx",
        "- graphical_abstract_non_ai.png",
        "- graphical_abstract_non_ai.pdf",
        "",
        "## Figures",
    ]
    for _, upload_stem, caption, _ in FIGURES:
        manifest_lines.append(f"- figures/{upload_stem}.png and .pdf: {caption}")
    manifest_lines.extend(
        [
            "",
            "## Real-image figure generation",
            "- figures/Figure_1_Real_image_protocol.* is generated from DATA325 captured greenhouse photographs, manual bounding boxes, ROI crops, and Attn+aug+TTA8 prediction records.",
            "- figures/Figure_8_DATA325_stage_error_gallery.* is generated from DATA325 photographs grouped by ground-truth height bin and absolute-error rank.",
            "- figures/Figure_9_Diagnostic_experiments.* is redrawn deterministically from bbox geometry, feature-alignment, and DANN-history JSON files.",
            "",
            "## Open-source release package",
            f"- {OPEN_RELEASE_ARCHIVE}: local public-release archive containing DATA325 images, cleaned annotations, sanitized evaluation JSON, selected DCF checkpoints, reproducibility scripts, licenses, citation metadata, and checksums. Public repository: {REPOSITORY_URL}.",
            "",
            "## Tables",
            "- tables/table1_ablation.csv",
            "- tables/table2_height_bins.csv",
            "",
            "## Newly added literature sources",
            "- CEA crop-height and protected-facility machine vision: Chang et al. 2017; Xie et al. 2021; Kim et al. 2021; Jayasuriya et al. 2024.",
            "- CEA agricultural computer vision and phenotyping reviews: Patricio and Rieder 2018; Kamilaris and Prenafeta-Boldu 2018; Li et al. 2020; Liu et al. 2020; Ariza-Sentis et al. 2024.",
            "- CEA crop-specific and multi-environment examples: Li et al. 2019; Xing et al. 2023; Che et al. 2024; Veramendi and Cruvinel 2024; Reena et al. 2025.",
            "- Domain-shift and foundation-model background: Sun and Saenko 2016; Koh et al. 2021; Gulrajani and Lopez-Paz 2021; Caron et al. 2021; Kirillov et al. 2023.",
            "",
            "## Reproducibility JSON",
        ]
    )
    for p in sorted(REPRO_OUT.glob("*.json")):
        manifest_lines.append(f"- reproducibility_json/{p.name}")
    (OUT / "submission_asset_manifest.md").write_text("\n".join(manifest_lines) + "\n", encoding="utf-8")


def zip_package() -> Path:
    zip_path = OUT / "cea_submission_package.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for p in OUT.rglob("*"):
            if not p.is_file() or p == zip_path:
                continue
            if "qa_render" in p.parts:
                continue
            z.write(p, p.relative_to(OUT))
    return zip_path


# ---------------------------------------------------------------------------
# CEA revision override block.
# This block updates the package to the Wheat3D-style DATA325 benchmark framing
# without removing the earlier helper functions, tables, and figure builders.
# ---------------------------------------------------------------------------

CEA_EXP_DIR = ROOT / "experiments" / "cea_revision"
MASK_EXAMPLE_DIR = CEA_EXP_DIR / "mask_examples"
SUPP_FIG_OUT = OUT / "supplementary_figures"

TITLE = (
    "A reproducible external-greenhouse benchmark for diagnosing DINOv3 "
    "feature transfer in maize height estimation"
)

FIGURES = [
    (
        "fig_real_protocol",
        "Figure_1_DATA325_acquisition_annotation_ROI_extraction",
        "Fig. 1. DATA325 acquisition, annotation, and ROI extraction. Source hand-box examples, independent target-greenhouse images, manual DATA325 boxes, ROI crops, height/camera metadata, and per-box prediction records are shown to make the benchmark auditable.",
        6.15,
    ),
    (
        "fig2_benchmark_comparison",
        "Figure_2_Source_DATA325_benchmark_comparison",
        "Fig. 2. Source and DATA325 benchmark comparison. Real source hand-box examples and real DATA325 target images illustrate the external-domain shift in greenhouse appearance, height range, camera height, and evaluation role.",
        6.15,
    ),
    (
        "fig3_distribution",
        "Figure_3_DATA325_distribution_imbalance",
        "Fig. 3. DATA325 distribution and imbalance. The target benchmark contains many early-stage plants, varied camera heights, heterogeneous bbox areas, and different foreground fractions across manually cropped ROIs.",
        6.15,
    ),
    (
        "fig1",
        "Figure_4_DINOv3_DCF_zero_shot_workflow",
        "Fig. 4. Mind-map workflow for DATA325 and DINOv3-DiffCorn-Fusion (DINOv3-DCF). The benchmark, frozen tokens, pooling equation, camera-height conditioning, structured latent head, zero-shot metrics, and failure diagnostics are shown as connected branches.",
        6.15,
    ),
    (
        "fig6_attention_roi",
        "Figure_5_Attention_pooling_real_ROIs",
        "Fig. 5. Attention pooling behavior on real maize ROIs. Real DATA325 crops are paired with deterministic foreground diagnostics and the attention-weighted token-pooling mechanism; the mask overlay is a QA diagnostic, not generated experimental evidence.",
        6.15,
    ),
    (
        "fig7_domain_shift_thumb",
        "Figure_6_Feature_domain_shift_with_ROI_thumbnails",
        "Fig. 6. Feature-space domain shift with ROI thumbnails. DINOv3 feature embeddings separate source and DATA325 ROIs, and thumbnail panels show that the clusters correspond to visibly different greenhouse and growth-stage conditions.",
        6.15,
    ),
    (
        "fig8_ablation_ci",
        "Figure_7_Main_ablation_results_with_CI",
        "Fig. 7. Main model and baseline ablation with bootstrap confidence intervals. Attention-weighted pooling provides the main DINOv3-DCF gain, while source-trained morphometric baselines show how much manual bbox geometry contributes under the same external target set.",
        6.15,
    ),
    (
        "fig9_resampling_stats",
        "Figure_8_Seed_robustness_paired_comparison",
        "Fig. 8. Seed robustness and paired comparison. Independent DCF-head retraining confirms that attention pooling remains the most stable external feature mode among CLS, patch-mean, attention, and attention+augmentation under TTA1 evaluation.",
        6.15,
    ),
    (
        "fig10_height_bin_ci",
        "Figure_9_Height_bin_error_early_stage_failure",
        "Fig. 9. Height-bin error and early-stage failure concentration. Bootstrap intervals confirm that plants below 80 cm dominate relative error and remain the key target for future stage-aware adaptation.",
        6.15,
    ),
    (
        "fig12_real_photo_matrix",
        "Figure_10_Real_photo_DATA325_prediction_matrix",
        "Fig. 10. Real-photo DATA325 prediction matrix. Each case shows a manual ground-truth box on the original image, the ROI used for model prediction, and a deterministic height-error diagnostic. \"Difference\" denotes plant-height error, foreground mask, and failure category; it is not a pixel-level segmentation or 3D point-cloud difference.",
        6.15,
    ),
]

SUPPLEMENTARY_FIGURES = [
    (
        "fig4_preprocessing",
        "Supplementary_Figure_1_Preprocessing_ROI_quality_examples",
        "Supplementary Fig. 1. Preprocessing and ROI-quality examples. Raw DATA325 images are converted to manual boxes, ROI crops, resized model inputs, and deterministic plant-mask diagnostics used only to quantify foreground/background contamination.",
        5.7,
    ),
    (
        "fig_real_stage_error_gallery",
        "Supplementary_Figure_2_STAGE_wise_DATA325_qualitative_gallery",
        "Supplementary Fig. 2. Stage-wise DATA325 qualitative gallery. Low-, mid-, and tall-stage real images are shown with manual boxes, ROI crops, ground truth, prediction, and absolute error.",
        5.7,
    ),
    (
        "fig8",
        "Supplementary_Figure_3_Attention_error_overlay_gallery",
        "Supplementary Fig. 3. Error-overlay gallery for real DATA325 ROIs. Successful, over-estimated, under-estimated, sparse, and uncertain cases are shown with ROI crops, deterministic plant-focus overlays, ground truth, prediction, and error.",
        5.7,
    ),
    (
        "fig15_release_map",
        "Supplementary_Figure_4_Open_release_future_deployment_map",
        "Supplementary Fig. 4. Open-release and future-deployment map. The release packages DATA325 images, annotations, predictions, diagnostics, scripts, and checkpoints, while future work connects automatic detection, segmentation-guided ROI normalization, and multi-greenhouse adaptation.",
        5.7,
    ),
    (
        "fig11_roi_contamination",
        "Supplementary_Figure_5_ROI_contamination_morphometric_diagnostics",
        "Supplementary Fig. 5. ROI contamination and morphometric diagnostics. Foreground/background fractions from deterministic color-index masks are weakly correlated with absolute error, while source-trained bbox/mask baselines quantify the strength and limits of simple agricultural geometry.",
        5.7,
    ),
    (
        "fig11_qualitative_combined",
        "Supplementary_Figure_6_Qualitative_stage_attention_error_gallery",
        "Supplementary Fig. 6. Qualitative stage-wise and attention/error gallery. Real DATA325 examples combine low-, mid-, and tall-stage cases with success, over-estimation, under-estimation, sparse-plant, and high-uncertainty ROIs.",
        5.7,
    ),
    (
        "fig7",
        "Supplementary_Figure_7_Diagnostic_negative_controls",
        "Supplementary Fig. 7. Diagnostic negative controls. Camera-height correction, bbox geometry, feature-statistic alignment, and DANN do not remove the external-greenhouse gap, so the remaining error is not a single metadata or marginal-alignment artifact.",
        5.7,
    ),
    (
        "fig12_extended_real_photo_matrix",
        "Supplementary_Figure_8_Extended_DATA325_real_photo_matrix",
        "Supplementary Fig. 8. Extended real-photo DATA325 matrix covering successful predictions, over-estimation, under-estimation, early sparse plants, high-background ROIs, and high TTA uncertainty.",
        5.7,
    ),
]
FIG_NAME_BY_STEM = {source_stem: upload_stem for source_stem, upload_stem, _, _ in FIGURES + SUPPLEMENTARY_FIGURES}

REPRO_FILES = [
    ROOT / "data325_zero_shot_attn_aug_tta8" / "data325_zero_shot_comparison_attn_aug_tta8.json",
    ROOT / "data325_zero_shot_corrected_camheight" / "data325_zero_shot_comparison_corrected_camheight.json",
    ROOT / "data325_zero_shot_attn_aug" / "data325_zero_shot_comparison_attn_aug.json",
    ROOT / "data325_zero_shot_attn_aug_featurealign" / "data325_zero_shot_comparison_attn_aug_featurealign.json",
    ROOT / "data325_zero_shot_dann" / "data325_zero_shot_comparison_dann.json",
    ROOT / "tsne_source_vs_data325.json",
    ROOT / "data325_zero_shot_attn_aug_featurealign" / "data325_feature_stats_attn_aug_featurealign.json",
    ROOT / "bbox_geometry_prior_phaseA.json",
    ROOT / "attention_geometry_prior_phaseA.json",
    ROOT / "checkpoints" / "diffcorn_fusion_hand_dann_history.json",
    CEA_EXP_DIR / "bootstrap_ci.json",
    CEA_EXP_DIR / "paired_tests.json",
    CEA_EXP_DIR / "roi_quality_summary.json",
    CEA_EXP_DIR / "morphometric_baseline.json",
    CEA_EXP_DIR / "source_morphometric_baseline.json",
    CEA_EXP_DIR / "uncertainty_diagnostic.json",
    CEA_EXP_DIR / "error_taxonomy_summary.json",
    CEA_EXP_DIR / "height_bin_bootstrap.json",
    CEA_EXP_DIR / "resampling_robustness.json",
    CEA_EXP_DIR / "seed_retraining_summary.json",
    CEA_EXP_DIR / "seed_retraining" / "data325_eval" / "data325_zero_shot_comparison_seed_retraining.json",
    CEA_EXP_DIR / "cea_revision_summary.json",
]


def add_figure(doc: Document, stem: str, caption: str, width_in: float) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(FIG_OUT / f"{FIG_NAME_BY_STEM.get(stem, stem)}.png"), width=Inches(min(width_in, 6.05)))
    add_caption(doc, caption)


def add_display_equation(doc: Document, equation_text: str, number: int) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        borders.append(elem)
    tbl_pr.append(borders)
    eq_cell, num_cell = table.rows[0].cells
    eq_cell.width = Inches(5.35)
    num_cell.width = Inches(0.7)
    for cell in (eq_cell, num_cell):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0

    p = eq_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_run = p.add_run(equation_text)
    eq_run.font.name = "Cambria Math"
    eq_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    eq_run.font.size = Pt(10)

    p_num = num_cell.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    num_run = p_num.add_run(f"Eq. ({number})")
    num_run.font.name = "Times New Roman"
    num_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    num_run.font.size = Pt(9.2)


def build_up_word_equations(path: Path) -> None:
    """Convert linear equation paragraphs into native Word Office Math objects."""
    script = f"""
$ErrorActionPreference = 'Stop'
$docPath = '{str(path).replace("'", "''")}'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
try {{
    $doc = $word.Documents.Open($docPath, $false, $false)
    foreach ($para in @($doc.Paragraphs)) {{
        $text = $para.Range.Text.Trim()
        $fontName = $para.Range.Font.Name
        if ($text.Length -gt 0 -and $text -notmatch 'Eq\\. \\([0-9]+\\)' -and $fontName -eq 'Cambria Math' -and $para.Range.OMaths.Count -eq 0) {{
            $eqRange = $para.Range.Duplicate
            while ($eqRange.End -gt $eqRange.Start) {{
                $last = $doc.Range($eqRange.End - 1, $eqRange.End).Text
                if ($last -eq "`r" -or ([int][char]$last[0]) -eq 7) {{
                    $eqRange.End = $eqRange.End - 1
                }} else {{
                    break
                }}
            }}
            if ($eqRange.Text.Trim().Length -gt 0) {{
                [void]$doc.OMaths.Add($eqRange)
                $doc.OMaths.Item($doc.OMaths.Count).BuildUp()
            }}
        }}
    }}
    $doc.Save()
    $doc.Close()
}} finally {{
    $word.Quit()
}}
"""
    result = subprocess.run(
        ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script],
        cwd=str(ROOT),
        capture_output=True,
        text=True,
        timeout=180,
    )
    if result.returncode != 0:
        raise RuntimeError(
            f"Word equation build-up failed for {path.name}:\n{result.stdout}\n{result.stderr}"
        )


def add_cover_p(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Els-body-text")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(9.5)


def bbox_xyxy(bbox: list[float], image_size: tuple[int, int] | None = None) -> tuple[int, int, int, int]:
    x1, y1, x2, y2 = [float(v) for v in bbox]
    if x2 <= x1 or y2 <= y1:
        x2 = x1 + max(1.0, x2)
        y2 = y1 + max(1.0, y2)
    if image_size is not None:
        w, h = image_size
        x1 = min(max(x1, 0.0), float(w - 1))
        y1 = min(max(y1, 0.0), float(h - 1))
        x2 = min(max(x2, x1 + 1.0), float(w))
        y2 = min(max(y2, y1 + 1.0), float(h))
    return int(round(x1)), int(round(y1)), int(round(x2)), int(round(y2))


def draw_bbox_on_image(path: Path, bbox: list[float], color: tuple[int, int, int] = (0, 158, 115)) -> Image.Image:
    image = Image.open(path).convert("RGB")
    draw = ImageDraw.Draw(image)
    x1, y1, x2, y2 = bbox_xyxy(bbox, image.size)
    width = max(5, image.width // 260)
    draw.rectangle((x1, y1, x2, y2), outline=color, width=width)
    return image


def crop_roi(path: Path, bbox: list[float], pad: int = 8) -> Image.Image:
    image = Image.open(path).convert("RGB")
    x1, y1, x2, y2 = bbox_xyxy(bbox, image.size)
    return image.crop((max(0, x1 - pad), max(0, y1 - pad), min(image.width, x2 + pad), min(image.height, y2 + pad)))


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8")) if path.exists() else {}


def load_csv_rows(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def load_revision_results() -> dict:
    return {
        "summary": load_json(CEA_EXP_DIR / "cea_revision_summary.json"),
        "bootstrap": load_json(CEA_EXP_DIR / "bootstrap_ci.json"),
        "paired": load_json(CEA_EXP_DIR / "paired_tests.json"),
        "roi_summary": load_json(CEA_EXP_DIR / "roi_quality_summary.json"),
        "roi_metrics": load_csv_rows(CEA_EXP_DIR / "roi_quality_metrics.csv"),
        "morphometric": load_json(CEA_EXP_DIR / "morphometric_baseline.json"),
        "source_morphometric": load_json(CEA_EXP_DIR / "source_morphometric_baseline.json"),
        "uncertainty": load_json(CEA_EXP_DIR / "uncertainty_diagnostic.json"),
        "taxonomy": load_csv_rows(CEA_EXP_DIR / "error_taxonomy.csv"),
        "taxonomy_summary": load_json(CEA_EXP_DIR / "error_taxonomy_summary.json"),
        "height_bins": load_json(CEA_EXP_DIR / "height_bin_bootstrap.json"),
        "resampling": load_json(CEA_EXP_DIR / "resampling_robustness.json"),
        "seed_retraining": load_json(CEA_EXP_DIR / "seed_retraining_summary.json"),
    }


def sanitize_json_value(value):
    if isinstance(value, dict):
        out = {}
        for key, item in value.items():
            if key == "image_path" and isinstance(item, str):
                out[key] = "data/DATA325/images/" + Path(item).name
            elif key == "checkpoint_path" and isinstance(item, str):
                out[key] = "checkpoints/" + Path(item).name
            elif key == "annotations_file" and isinstance(item, str):
                out[key] = "data/DATA325/annotations/data325_annotations.csv"
            elif key == "capture_height_mapping_csv" and isinstance(item, str):
                out[key] = "data/DATA325/annotations/camera_height_mapping.csv"
            elif key == "source_bundle" and isinstance(item, str):
                out[key] = "data/source_feature_bundles/" + Path(item).name
            elif key == "target_image_root" and isinstance(item, str):
                out[key] = "data/DATA325/images"
            elif key in {"path", "output_dir"} and isinstance(item, str):
                out[key] = "results/reproducibility_json/" + Path(item).name
            else:
                out[key] = sanitize_json_value(item)
        return out
    if isinstance(value, list):
        return [sanitize_json_value(item) for item in value]
    if isinstance(value, str):
        replacements = [
            (r"C:\Users\Wuhon\OneDrive\桌面\paper\data325", "data/DATA325/images"),
            (r"D:\OneDrive\桌面\paper\data325", "data/DATA325/images"),
            (r"D:\wechatfile", "external_metadata"),
            (r"D:\cornTrain\DINOV3\checkpoints", "checkpoints"),
            (r"D:\cornTrain\DINOV3", "."),
            (r"D:\cornTrain\dcf-bbox-eval-tool\data", "data/DATA325/annotations"),
            (r"C:\Users\Wuhon", "."),
        ]
        out = value
        for old, new in replacements:
            out = out.replace(old, new)
        return out.replace("\\", "/")
    return value


def canonical_ablation_rows(m: dict) -> list[dict]:
    rev = load_revision_results()
    summary = rev["summary"].get("model_summary", {})
    by_method = {row["method"]: dict(row) for row in m["ablation"]}

    def row(method: str, feature: str, key: str, source_method: str, note: str = "") -> dict:
        base = by_method.get(source_method, {})
        s = summary[key]
        return {
            "method": method,
            "feature": feature,
            "source_mae": base.get("source_mae", float("nan")),
            "data325_mae": s["mae_cm"],
            "data325_rmse": s["rmse_cm"],
            "data325_mape": s["mape_percent"],
            "note": note,
        }

    rows = [
        row("Old baseline", "CLS", "old", "Old baseline"),
        row("CLS retrain", "CLS", "cls", "CLS retrain"),
        row("Patch mean", "Patch mean", "patch_mean", "Patch mean"),
        row("Attn-weighted", "Attn-weighted", "attn", "Attn-weighted"),
        row("Corrected cam height", "Attn-weighted + camera-height correction", "corrected_camheight", "Attn-weighted", "diagnostic"),
        row("Attn+aug", "Attn-weighted + aug", "attn_aug", "Attn+aug"),
        row("Attn+aug+TTA8", "Attn-weighted + aug + TTA8", "attn_aug_tta8", "Attn+aug+TTA8", "best"),
        row("Feat. align", "Attn-weighted + aug + feature alignment", "attn_aug_featurealign", "Feat. align", "failed"),
    ]
    if "DANN" in by_method:
        rows.append(by_method["DANN"])
    source_morph = rev.get("source_morphometric", {})
    for model_key, label in [
        ("ridge_cv", "Source morph. RidgeCV"),
        ("random_forest", "Source morph. RF"),
    ]:
        model = source_morph.get("models", {}).get(model_key, {})
        summary_row = model.get("summary", {})
        if summary_row:
            rows.append(
                {
                    "method": label,
                    "feature": "source-trained bbox + camera + ExG mask geometry",
                    "source_mae": float("nan"),
                    "data325_mae": summary_row.get("mae_cm", float("nan")),
                    "data325_rmse": summary_row.get("rmse_cm", float("nan")),
                    "data325_mape": summary_row.get("mape_percent", float("nan")),
                    "note": "source-trained geometry baseline",
                }
            )
    return rows


def to_float(row: dict, key: str, default: float = float("nan")) -> float:
    try:
        return float(row.get(key, default))
    except (TypeError, ValueError):
        return default


def choose_by_height(records: list[dict], lo: float, hi: float, rank: str = "low") -> dict:
    subset = [r for r in records if lo <= float(r["true_height_cm"]) < hi]
    if not subset:
        subset = records
    subset = sorted(subset, key=lambda r: float(r["abs_error_cm"]))
    if rank == "high":
        return subset[-1]
    if rank == "mid":
        return subset[len(subset) // 2]
    return subset[0]


def save_matplotlib(fig, upload_stem: str) -> None:
    fig.savefig(FIG_OUT / f"{upload_stem}.png", bbox_inches="tight", pad_inches=0.08, dpi=300)
    fig.savefig(FIG_OUT / f"{upload_stem}.pdf", bbox_inches="tight", pad_inches=0.08)
    plt.close(fig)


def green_mask_array(roi: Image.Image) -> np.ndarray:
    arr = np.asarray(roi.convert("RGB"), dtype=np.float32)
    r, g, b = arr[..., 0], arr[..., 1], arr[..., 2]
    exg = 2.0 * g - r - b
    chroma = (g > r * 1.04) & (g > b * 1.04)
    dynamic = exg > np.quantile(exg, 0.60)
    bright = arr.mean(axis=2) > 20
    return (dynamic | chroma) & bright


def mask_overlay(roi: Image.Image, alpha: int = 105) -> Image.Image:
    small = roi.convert("RGB")
    mask = Image.fromarray((green_mask_array(small).astype(np.uint8) * 255), mode="L")
    overlay = small.convert("RGBA")
    green = Image.new("RGBA", overlay.size, (35, 170, 105, alpha))
    return Image.composite(green, overlay, mask).convert("RGB")


def build_fig2_benchmark_comparison(records: list[dict], rev: dict) -> None:
    upload = FIG_NAME_BY_STEM["fig2_benchmark_comparison"]
    canvas = Image.new("RGB", (2400, 1600), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(52, True)
    head_font = pil_font(34, True)
    body_font = pil_font(25)
    small_font = pil_font(21)
    blue = (68, 114, 196)
    green = (29, 158, 117)
    orange = (216, 90, 48)
    gray = (110, 114, 120)
    black = (32, 33, 36)
    draw.text((70, 45), "Source vs DATA325 benchmark comparison", font=title_font, fill=black)
    draw.text((70, 112), "Real photographs show why DATA325 is an external domain-shift benchmark, not a random validation split.", font=body_font, fill=gray)

    draw.rectangle((70, 190, 1130, 1325), fill=(247, 249, 251), outline=blue, width=4)
    draw.rectangle((1270, 190, 2330, 1325), fill=(247, 249, 251), outline=green, width=4)
    draw.text((105, 220), "Source hand-bbox set", font=head_font, fill=blue)
    draw.text((1305, 220), "DATA325 target greenhouse", font=head_font, fill=green)

    source_files = [
        source_sample("106-64-*_bbox.jpg"),
        source_sample("140-121-*_bbox.jpg"),
        source_sample("140-170-*_bbox.jpg"),
        source_sample("140-270-*_bbox.jpg"),
    ]
    target_records = [
        choose_by_height(records, 0, 80, "low"),
        choose_by_height(records, 80, 120, "mid"),
        choose_by_height(records, 120, 160, "low"),
        choose_by_height(records, 160, 260, "low"),
    ]
    grid_boxes = [(110, 285, 565, 590), (620, 285, 1075, 590), (110, 655, 565, 960), (620, 655, 1075, 960)]
    for path, box in zip(source_files, grid_boxes):
        img = Image.open(path).convert("RGB")
        paste_fit(canvas, img, box, (250, 250, 250))
        draw.text((box[0], box[3] + 12), path.stem.replace("_bbox", ""), font=small_font, fill=gray)
    grid_boxes_r = [(1310, 285, 1765, 590), (1820, 285, 2275, 590), (1310, 655, 1765, 960), (1820, 655, 2275, 960)]
    for rec, box in zip(target_records, grid_boxes_r):
        img = draw_bbox_on_image(DATA325_IMAGE_DIR / rec["file_name"], rec["bbox"], green)
        paste_fit(canvas, img, box, (250, 250, 250))
        draw.text((box[0], box[3] + 12), f"{rec['true_height_cm']:.0f} cm, {rec['box_id']}", font=small_font, fill=gray)

    table_y = 1040
    rows = [
        ("Role", "Training/calibration", "External zero-shot evaluation"),
        ("Samples", "156 hand-box ROIs", "82 evaluated boxes from 25 images"),
        ("Height range", "64-270 cm", "30-178 cm"),
        ("Camera heights", "106 and 140 cm", "mapped per image, mainly 140 cm"),
        ("Leakage control", "Source-domain only", "No target labels used for training"),
    ]
    col_x = [115, 475, 1360, 1910]
    for i, (k, a, b) in enumerate(rows):
        y = table_y + i * 52
        draw.text((115, y), k, font=body_font, fill=black)
        draw.text((470, y), a, font=body_font, fill=blue)
        draw.text((1310, y), k, font=body_font, fill=black)
        draw_wrapped(draw, (1665, y), b, body_font, green, 565, 2)

    draw.rectangle((70, 1370, 2330, 1485), fill=(248, 238, 233), outline=orange, width=3)
    draw.text((105, 1395), "Benchmark framing", font=head_font, fill=orange)
    draw_wrapped(
        draw,
        (470, 1390),
        "DATA325 is used as a diagnostic external-greenhouse benchmark: the goal is to measure whether frozen foundation features and ROI-level aggregation transfer across real greenhouse conditions.",
        body_font,
        black,
        1750,
        2,
    )
    save_png_pdf(canvas, upload)


def build_fig3_distribution(records: list[dict], rev: dict) -> None:
    upload = FIG_NAME_BY_STEM["fig3_distribution"]
    roi = rev["roi_metrics"]
    heights = np.array([float(r["true_height_cm"]) for r in records])
    cams = np.array([float(r["camera_height_cm"]) for r in records])
    bbox_area = np.array([to_float(r, "bbox_area_fraction") for r in roi])
    fg = np.array([to_float(r, "foreground_fraction") for r in roi])
    bins = [0, 80, 120, 160, 220]
    plt.rcParams.update({"font.family": "DejaVu Sans", "font.size": 10})
    fig, axs = plt.subplots(2, 2, figsize=(12.7, 7.6), dpi=220)
    axs = axs.ravel()
    axs[0].hist(heights, bins=bins, color=f"#{GREEN}", edgecolor="white")
    axs[0].axvline(80, color=f"#{ORANGE}", linestyle="--", lw=2)
    axs[0].set_title("(a) DATA325 height distribution")
    axs[0].set_xlabel("Ground-truth height (cm)")
    axs[0].set_ylabel("Boxes")
    axs[1].hist(cams, bins=8, color=f"#{BLUE}", edgecolor="white")
    axs[1].set_title("(b) Camera-height metadata")
    axs[1].set_xlabel("Camera height (cm)")
    axs[1].set_ylabel("Boxes")
    axs[2].hist(bbox_area[np.isfinite(bbox_area)] * 100, bins=12, color="#7E57C2", edgecolor="white")
    axs[2].set_title("(c) Manual bbox area")
    axs[2].set_xlabel("BBox area (% of image)")
    axs[2].set_ylabel("Boxes")
    axs[3].hist(fg[np.isfinite(fg)] * 100, bins=12, color=f"#{ORANGE}", edgecolor="white")
    axs[3].set_title("(d) Foreground fraction from color-index mask")
    axs[3].set_xlabel("Foreground fraction (%)")
    axs[3].set_ylabel("Boxes")
    for ax in axs:
        ax.grid(axis="y", color="#D8DEE6", lw=0.6)
        ax.spines[["top", "right"]].set_visible(False)
    fig.suptitle("DATA325 distribution and imbalance", fontsize=16, weight="bold")
    fig.tight_layout(rect=[0.02, 0.02, 0.98, 0.93])
    fig.subplots_adjust(hspace=0.48, wspace=0.28)
    save_matplotlib(fig, upload)


def record_key(record: dict) -> tuple[str, str]:
    return (str(record.get("file_name", "")), str(record.get("box_id", "")))


def row_lookup(rows: list[dict[str, str]]) -> dict[tuple[str, str], dict[str, str]]:
    return {(str(row.get("file_name", "")), str(row.get("box_id", ""))): row for row in rows}


def record_stage(record: dict) -> str:
    h = float(record["true_height_cm"])
    if h < 80:
        return "early <80 cm"
    if h < 120:
        return "mid 80-120 cm"
    return "tall >=120 cm"


def category_text(record: dict, taxonomy: dict[tuple[str, str], dict[str, str]]) -> str:
    row = taxonomy.get(record_key(record), {})
    cat = row.get("primary_error_category") or row.get("category") or ""
    return cat.replace("_", " ").strip() or "height-error diagnostic"


def choose_error_rank(records: list[dict], lo: float, hi: float | None, rank: str) -> dict:
    subset = [r for r in records if float(r["true_height_cm"]) >= lo and (hi is None or float(r["true_height_cm"]) < hi)]
    if not subset:
        subset = records
    subset = sorted(subset, key=lambda r: float(r["abs_error_cm"]))
    if rank == "high":
        return subset[-1]
    if rank == "mid":
        return subset[len(subset) // 2]
    return subset[0]


def build_fig4_mindmap_workflow() -> None:
    upload = FIG_NAME_BY_STEM["fig1"]
    canvas = Image.new("RGB", (2600, 1700), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(54, True)
    center_font = pil_font(36, True)
    head_font = pil_font(30, True)
    body_font = pil_font(23)
    formula_font = pil_font(23, True)
    small_font = pil_font(20)
    black = (32, 33, 36)
    gray = (98, 105, 115)
    green = (29, 158, 117)
    blue = (68, 114, 196)
    orange = (216, 90, 48)
    purple = (116, 82, 180)
    teal = (0, 135, 150)
    slate = (82, 92, 105)
    colors = [green, blue, orange, purple, teal, slate]

    draw.text((90, 55), "DATA325 external-greenhouse diagnostic benchmark", font=title_font, fill=black)
    draw.text(
        (90, 125),
        "A mind-map view of how real images, frozen DINOv3 tokens, DCF prediction, and diagnostic release assets connect.",
        font=body_font,
        fill=gray,
    )

    center = (910, 600, 1690, 1040)
    draw.rounded_rectangle(center, radius=42, fill=(245, 249, 247), outline=green, width=5)
    draw_wrapped(
        draw,
        (center[0] + 62, center[1] + 95),
        "DATA325 external-greenhouse diagnostic benchmark",
        center_font,
        black,
        center[2] - center[0] - 124,
        8,
    )
    draw_wrapped(
        draw,
        (center[0] + 65, center[1] + 255),
        "82 evaluated boxes from 25 real target-greenhouse images, held out from source-domain DCF training.",
        body_font,
        gray,
        center[2] - center[0] - 130,
        5,
    )

    branches = [
        ((105, 285, 735, 520), "Real images / ROI labels", "I, b, h, h_cam", "Manual boxes isolate model transfer from detector errors."),
        ((105, 690, 735, 925), "DINOv3 frozen tokens", "{z_cls, z_i}=F_DINO(r)", "The backbone is frozen; only the DCF head is trained on source ROIs."),
        ((105, 1095, 735, 1330), "Feature pooling", "f_attn=sum_i alpha_i z_i", "CLS, patch-mean, and attention pooling test ROI representation transfer."),
        ((1865, 285, 2495, 520), "Camera-height context", "x=[f; h_cam/200]", "A scalar acquisition context is concatenated with visual features."),
        ((1865, 690, 2495, 925), "DCF latent head", "h_hat=sum_k p[k,0]", "The 64D output is a phytomer-inspired latent vector, supervised by derived height."),
        ((1865, 1095, 2495, 1330), "Metrics, failures, release", "MAE, RMSE, MAPE, CI", "Bootstrap, paired tests, error taxonomy, and open files make the benchmark auditable."),
    ]
    anchor_points = [
        (center[0], center[1] + 70),
        (center[0], center[1] + 220),
        (center[0], center[3] - 70),
        (center[2], center[1] + 70),
        (center[2], center[1] + 220),
        (center[2], center[3] - 70),
    ]
    for idx, (box, title, formula, note) in enumerate(branches):
        color = colors[idx]
        bx0, by0, bx1, by1 = box
        if bx1 < center[0]:
            end = (bx1, (by0 + by1) // 2)
        else:
            end = (bx0, (by0 + by1) // 2)
        draw.line((anchor_points[idx][0], anchor_points[idx][1], end[0], end[1]), fill=color, width=6)
        draw.ellipse((end[0] - 11, end[1] - 11, end[0] + 11, end[1] + 11), fill=color)
        draw.rounded_rectangle(box, radius=28, fill=(250, 252, 252), outline=color, width=4)
        draw.text((bx0 + 30, by0 + 25), title, font=head_font, fill=color)
        draw.rounded_rectangle((bx0 + 30, by0 + 75, bx1 - 30, by0 + 126), radius=14, fill=(246, 247, 249), outline=(224, 229, 235), width=2)
        draw.text((bx0 + 48, by0 + 88), formula, font=formula_font, fill=black)
        draw_wrapped(draw, (bx0 + 32, by0 + 145), note, body_font, gray, bx1 - bx0 - 64, 4)

    draw.rounded_rectangle((210, 1430, 2390, 1575), radius=26, fill=(251, 247, 240), outline=orange, width=3)
    draw.text((250, 1460), "Interpretation boundary", font=head_font, fill=orange)
    draw_wrapped(
        draw,
        (650, 1455),
        "The figure summarizes the evaluation pipeline. Evidence figures elsewhere use only real DATA325/source images, recorded boxes, model outputs, and deterministic diagnostics.",
        body_font,
        black,
        1670,
        5,
    )
    draw.text((90, 1620), "Deterministic drawing script; no generative image editing was used.", font=small_font, fill=gray)
    save_png_pdf(canvas, upload)


def draw_height_difference_panel(
    canvas: Image.Image,
    record: dict,
    box: tuple[int, int, int, int],
    category: str,
    label: str,
) -> None:
    draw = ImageDraw.Draw(canvas)
    x0, y0, x1, y1 = box
    head_font = pil_font(20, True)
    body_font = pil_font(17)
    small_font = pil_font(14)
    black = (32, 33, 36)
    gray = (99, 107, 116)
    green = (29, 158, 117)
    orange = (216, 90, 48)
    blue = (68, 114, 196)
    draw.rectangle(box, fill=(255, 255, 255), outline=(218, 224, 230), width=2)
    draw.text((x0 + 12, y0 + 10), "Difference", font=head_font, fill=black)
    gt = float(record["true_height_cm"])
    pred = float(record["pred_height_cm"])
    ae = abs(pred - gt)
    signed = pred - gt
    scale_max = max(200.0, gt, pred)
    base_y = y0 + 220
    bar_w = 36
    bar_h_gt = int(150 * gt / scale_max)
    bar_h_pred = int(150 * pred / scale_max)
    gx = x0 + 35
    px = x0 + 90
    draw.rectangle((gx, base_y - bar_h_gt, gx + bar_w, base_y), fill=green)
    draw.rectangle((px, base_y - bar_h_pred, px + bar_w, base_y), fill=orange if signed > 0 else blue)
    draw.line((x0 + 25, base_y, x0 + 145, base_y), fill=(170, 178, 188), width=2)
    draw.text((gx - 2, base_y + 8), "GT", font=small_font, fill=green)
    draw.text((px - 8, base_y + 8), "Pred", font=small_font, fill=orange if signed > 0 else blue)
    draw.text((x0 + 150, y0 + 58), f"AE {ae:.1f} cm", font=head_font, fill=orange if ae >= 30 else black)
    draw.text((x0 + 150, y0 + 90), f"Delta {signed:+.1f} cm", font=body_font, fill=orange if signed > 0 else blue)
    draw_wrapped(draw, (x0 + 150, y0 + 123), category, body_font, gray, x1 - x0 - 170, 2)
    draw.text((x0 + 150, y0 + 226), label, font=small_font, fill=gray)

    roi = crop_roi(DATA325_IMAGE_DIR / record["file_name"], record["bbox"])
    overlay = mask_overlay(roi)
    paste_fit(canvas, overlay, (x0 + 255, y0 + 55, x1 - 16, y1 - 18), (249, 250, 251))


def build_real_photo_prediction_matrix(records: list[dict], rev: dict) -> None:
    upload = FIG_NAME_BY_STEM["fig12_real_photo_matrix"]
    taxonomy = row_lookup(rev["taxonomy"])
    selections = [
        ("Early <80 cm", [
            choose_error_rank(records, 0, 80, "low"),
            choose_error_rank(records, 0, 80, "mid"),
            choose_error_rank(records, 0, 80, "high"),
        ]),
        ("Mid 80-120 cm", [
            choose_error_rank(records, 80, 120, "low"),
            choose_error_rank(records, 80, 120, "mid"),
            choose_error_rank(records, 80, 120, "high"),
        ]),
        ("Tall >=120 cm", [
            choose_error_rank(records, 120, None, "low"),
            choose_error_rank(records, 120, None, "mid"),
            choose_error_rank(records, 120, None, "high"),
        ]),
    ]
    canvas = Image.new("RGB", (3600, 2550), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(58, True)
    stage_font = pil_font(34, True)
    head_font = pil_font(25, True)
    body_font = pil_font(20)
    small_font = pil_font(17)
    black = (32, 33, 36)
    gray = (100, 107, 116)
    green = (29, 158, 117)
    orange = (216, 90, 48)
    blue = (68, 114, 196)
    draw.text((80, 55), "Real-photo DATA325 prediction matrix", font=title_font, fill=black)
    draw.text(
        (80, 130),
        "Rows follow growth stage; columns select low-, median-, and high-error cases within each stage.",
        font=body_font,
        fill=gray,
    )

    group_w = 1120
    row_h = 735
    start_y = 250
    start_x = 125
    panel_gap = 15
    panel_w = 345
    panel_h = 465
    sample_labels = ["low error", "median error", "high error"]
    for row_idx, (stage, row_records) in enumerate(selections):
        y = start_y + row_idx * row_h
        draw.rounded_rectangle((45, y - 68, 3555, y + row_h - 78), radius=22, fill=(249, 251, 252), outline=(224, 229, 235), width=2)
        draw.text((80, y - 50), stage, font=stage_font, fill=green if row_idx == 0 else blue if row_idx == 1 else orange)
        for col_idx, record in enumerate(row_records):
            x = start_x + col_idx * group_w
            draw.text((x + 445, y - 48), sample_labels[col_idx], font=head_font, fill=gray)
            raw_box = (x, y, x + panel_w, y + panel_h)
            pred_box = (x + panel_w + panel_gap, y, x + panel_w * 2 + panel_gap, y + panel_h)
            diff_box = (x + panel_w * 2 + panel_gap * 2, y, x + panel_w * 3 + panel_gap * 2, y + panel_h)

            path = DATA325_IMAGE_DIR / record["file_name"]
            raw = draw_bbox_on_image(path, record["bbox"], green)
            paste_fit(canvas, raw, raw_box, (250, 250, 250))
            draw.rectangle(raw_box, outline=(218, 224, 230), width=2)
            draw.text((raw_box[0] + 10, raw_box[1] + 10), "Ground truth", font=head_font, fill=black)
            draw.rectangle((raw_box[0] + 8, raw_box[3] - 72, raw_box[2] - 8, raw_box[3] - 8), fill=(255, 255, 255), outline=(224, 229, 235), width=1)
            draw.text((raw_box[0] + 18, raw_box[3] - 62), f"GT {float(record['true_height_cm']):.0f} cm", font=body_font, fill=green)
            draw.text((raw_box[0] + 18, raw_box[3] - 33), f"Camera {float(record['camera_height_cm']):.0f} cm", font=small_font, fill=gray)

            roi = crop_roi(path, record["bbox"])
            paste_fit(canvas, roi, pred_box, (250, 250, 250))
            draw.rectangle(pred_box, outline=(218, 224, 230), width=2)
            draw.text((pred_box[0] + 10, pred_box[1] + 10), "Prediction", font=head_font, fill=black)
            draw.rectangle((pred_box[0] + 8, pred_box[3] - 78, pred_box[2] - 8, pred_box[3] - 8), fill=(255, 255, 255), outline=(224, 229, 235), width=1)
            draw.text((pred_box[0] + 18, pred_box[3] - 67), f"Pred {float(record['pred_height_cm']):.1f} cm", font=body_font, fill=orange)
            draw.text((pred_box[0] + 18, pred_box[3] - 38), f"TTA std {float(record.get('pred_std_cm', 0.0)):.1f} cm", font=small_font, fill=gray)

            draw_height_difference_panel(canvas, record, diff_box, category_text(record, taxonomy), record["box_id"])

    draw.text(
        (80, 2440),
        "Difference panels report height-error diagnostics and deterministic foreground masks, not pixel-level or 3D difference ground truth.",
        font=body_font,
        fill=gray,
    )
    save_png_pdf(canvas, upload)


def unique_record_add(out: list[tuple[str, dict]], used: set[tuple[str, str]], label: str, candidates: list[dict], count: int) -> None:
    for rec in candidates:
        key = record_key(rec)
        if key in used:
            continue
        out.append((label, rec))
        used.add(key)
        if sum(1 for item_label, _ in out if item_label == label) >= count:
            return


def build_extended_real_photo_matrix(records: list[dict], rev: dict) -> None:
    upload = FIG_NAME_BY_STEM["fig12_extended_real_photo_matrix"]
    taxonomy = row_lookup(rev["taxonomy"])
    quality = row_lookup(rev["roi_metrics"])
    by_error = sorted(records, key=lambda r: float(r["abs_error_cm"]))
    over = sorted(records, key=lambda r: float(r["pred_height_cm"]) - float(r["true_height_cm"]), reverse=True)
    under = sorted(records, key=lambda r: float(r["pred_height_cm"]) - float(r["true_height_cm"]))
    early = sorted([r for r in records if float(r["true_height_cm"]) < 80], key=lambda r: float(r["abs_error_cm"]), reverse=True)

    def bg_score(rec: dict) -> float:
        row = quality.get(record_key(rec), {})
        return float(row.get("background_fraction") or 0.0)

    high_bg = sorted(records, key=bg_score, reverse=True)
    high_unc = sorted(records, key=lambda r: float(r.get("pred_std_cm", 0.0)), reverse=True)
    selected: list[tuple[str, dict]] = []
    used: set[tuple[str, str]] = set()
    unique_record_add(selected, used, "success", by_error, 3)
    unique_record_add(selected, used, "over-estimation", over, 3)
    unique_record_add(selected, used, "under-estimation", under, 3)
    unique_record_add(selected, used, "early sparse", early, 3)
    unique_record_add(selected, used, "high background", high_bg, 3)
    unique_record_add(selected, used, "high uncertainty", high_unc, 3)

    canvas = Image.new("RGB", (3600, 3050), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(56, True)
    head_font = pil_font(24, True)
    body_font = pil_font(18)
    small_font = pil_font(15)
    black = (32, 33, 36)
    gray = (98, 105, 115)
    green = (29, 158, 117)
    orange = (216, 90, 48)
    blue = (68, 114, 196)
    draw.text((80, 55), "Extended DATA325 real-photo matrix", font=title_font, fill=black)
    draw.text(
        (80, 128),
        "Each card uses the original image, manual bbox, ROI crop, deterministic foreground mask, and recorded DINOv3-DCF prediction.",
        font=body_font,
        fill=gray,
    )
    card_w = 560
    card_h = 880
    x0 = 70
    y0 = 225
    gap_x = 28
    gap_y = 55
    for idx, (label, rec) in enumerate(selected[:18]):
        row = idx // 6
        col = idx % 6
        x = x0 + col * (card_w + gap_x)
        y = y0 + row * (card_h + gap_y)
        draw.rounded_rectangle((x, y, x + card_w, y + card_h), radius=20, fill=(249, 251, 252), outline=(220, 226, 232), width=2)
        color = green if label == "success" else orange if "over" in label or "early" in label else blue
        draw.text((x + 18, y + 16), label, font=head_font, fill=color)
        draw.text((x + 18, y + 48), record_stage(rec), font=small_font, fill=gray)
        path = DATA325_IMAGE_DIR / rec["file_name"]
        raw = draw_bbox_on_image(path, rec["bbox"], green)
        paste_fit(canvas, raw, (x + 18, y + 82, x + card_w - 18, y + 380), (250, 250, 250))
        roi = crop_roi(path, rec["bbox"])
        overlay = mask_overlay(roi)
        paste_fit(canvas, roi, (x + 18, y + 405, x + card_w // 2 - 8, y + 620), (250, 250, 250))
        paste_fit(canvas, overlay, (x + card_w // 2 + 8, y + 405, x + card_w - 18, y + 620), (250, 250, 250))
        gt = float(rec["true_height_cm"])
        pred = float(rec["pred_height_cm"])
        ae = abs(pred - gt)
        std = float(rec.get("pred_std_cm", 0.0))
        draw.text((x + 18, y + 650), f"GT {gt:.0f} cm | pred {pred:.1f} cm | AE {ae:.1f} cm", font=body_font, fill=black)
        draw.text((x + 18, y + 681), f"TTA std {std:.1f} cm | box {rec['box_id']}", font=small_font, fill=gray)
        draw_wrapped(draw, (x + 18, y + 713), category_text(rec, taxonomy), body_font, gray, card_w - 36, 2)
        q = quality.get(record_key(rec), {})
        fg = q.get("foreground_fraction")
        bg = q.get("background_fraction")
        if fg is not None and bg is not None:
            draw.text((x + 18, y + 820), f"foreground {float(fg) * 100:.0f}% | background {float(bg) * 100:.0f}%", font=small_font, fill=gray)
    draw.text(
        (80, 2970),
        "The matrix is a qualitative audit sheet; mask overlays are deterministic color-index diagnostics and are not segmentation labels.",
        font=body_font,
        fill=gray,
    )
    save_png_pdf(canvas, upload)


def build_fig4_preprocessing(records: list[dict], rev: dict) -> None:
    upload = FIG_NAME_BY_STEM["fig4_preprocessing"]
    recs = [choose_by_height(records, 0, 80, "mid"), choose_by_height(records, 80, 120, "mid"), choose_by_height(records, 120, 220, "mid")]
    canvas = Image.new("RGB", (2400, 1580), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(50, True)
    head_font = pil_font(28, True)
    body_font = pil_font(22)
    green = (29, 158, 117)
    orange = (216, 90, 48)
    blue = (68, 114, 196)
    gray = (112, 117, 125)
    black = (32, 33, 36)
    draw.text((70, 40), "Preprocessing and ROI-quality diagnostics", font=title_font, fill=black)
    headers = ["Raw target image", "Manual bbox", "ROI crop", "224 x 224 input", "Plant-mask diagnostic"]
    x_boxes = [(70, 190, 490, 520), (535, 190, 955, 520), (1000, 190, 1370, 520), (1415, 190, 1695, 520), (1740, 190, 2330, 520)]
    for x0, _, x1, _ in x_boxes:
        draw.rectangle((x0, 145, x1, 180), fill=(231, 243, 238), outline=green, width=1)
    for i, h in enumerate(headers):
        draw.text((x_boxes[i][0] + 8, 151), h, font=body_font, fill=green)
    y_offsets = [190, 640, 1090]
    for row_idx, rec in enumerate(recs):
        y0 = y_offsets[row_idx]
        boxes = [(70, y0, 490, y0 + 330), (535, y0, 955, y0 + 330), (1000, y0, 1370, y0 + 330), (1415, y0, 1695, y0 + 330), (1740, y0, 2330, y0 + 330)]
        path = DATA325_IMAGE_DIR / rec["file_name"]
        raw = Image.open(path).convert("RGB")
        bb = draw_bbox_on_image(path, rec["bbox"], green if row_idx != 0 else orange)
        roi = crop_roi(path, rec["bbox"])
        resized = roi.resize((224, 224), Image.Resampling.LANCZOS)
        overlay = mask_overlay(roi)
        for img, box in zip([raw, bb, roi, resized, overlay], boxes):
            paste_fit(canvas, img, box, (250, 250, 250))
            draw.rectangle(box, outline=(218, 224, 230), width=2)
        note = f"GT {rec['true_height_cm']:.0f} cm | pred {rec['pred_height_cm']:.1f} cm | AE {rec['abs_error_cm']:.1f} cm"
        draw.text((70, y0 + 345), note, font=head_font, fill=blue if row_idx > 0 else orange)
        draw.line((493, y0 + 165, 532, y0 + 165), fill=gray, width=3)
        draw.line((958, y0 + 165, 997, y0 + 165), fill=gray, width=3)
        draw.line((1373, y0 + 165, 1412, y0 + 165), fill=gray, width=3)
        draw.line((1698, y0 + 165, 1737, y0 + 165), fill=gray, width=3)
    draw_wrapped(
        draw,
        (1745, 1445),
        "Mask panels are deterministic color-index diagnostics for foreground/background quantification. They are not generated images and are not used as training labels.",
        body_font,
        gray,
        570,
        2,
    )
    save_png_pdf(canvas, upload)


def build_fig6_attention_roi(records: list[dict], rev: dict) -> None:
    upload = FIG_NAME_BY_STEM["fig6_attention_roi"]
    recs = [choose_by_height(records, 0, 80, "high"), choose_by_height(records, 80, 120, "mid"), choose_by_height(records, 120, 220, "low")]
    canvas = Image.new("RGB", (2400, 1550), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(50, True)
    head_font = pil_font(30, True)
    body_font = pil_font(24)
    small_font = pil_font(20)
    blue = (68, 114, 196)
    green = (29, 158, 117)
    orange = (216, 90, 48)
    gray = (110, 114, 120)
    black = (32, 33, 36)
    draw.text((70, 40), "Attention-guided pooling on real maize ROIs", font=title_font, fill=black)
    draw.text((70, 108), "Real crops supply the evidence; the token-pooling diagram explains the feature aggregation step.", font=body_font, fill=gray)

    y = 210
    for rec in recs:
        path = DATA325_IMAGE_DIR / rec["file_name"]
        roi = crop_roi(path, rec["bbox"])
        overlay = mask_overlay(roi)
        paste_fit(canvas, roi, (90, y, 500, y + 280), (250, 250, 250))
        paste_fit(canvas, overlay, (535, y, 945, y + 280), (250, 250, 250))
        draw.rectangle((90, y, 500, y + 280), outline=green, width=3)
        draw.rectangle((535, y, 945, y + 280), outline=orange, width=3)
        draw.text((90, y - 34), f"GT {rec['true_height_cm']:.0f} cm, AE {rec['abs_error_cm']:.1f} cm", font=head_font, fill=black)
        draw.text((95, y + 292), "ROI crop", font=small_font, fill=gray)
        draw.text((540, y + 292), "Foreground QA overlay", font=small_font, fill=gray)
        y += 400

    x0, y0 = 1130, 245
    draw.rectangle((x0, y0, 2315, 1305), fill=(247, 249, 251), outline=blue, width=4)
    draw.text((x0 + 35, y0 + 35), "Token aggregation control", font=head_font, fill=blue)
    stages = [
        ("224 x 224 ROI", "Patch tokens + CLS"),
        ("Frozen DINOv3 ViT-L", "No target-domain training"),
        ("CLS-to-patch attention", "Average heads, normalize patches"),
        ("Weighted patch descriptor", "1024D plant-focused feature"),
        ("DCF regression head", "feature + camera height -> plant height"),
    ]
    sy = y0 + 120
    for i, (a, b) in enumerate(stages):
        color = [green, blue, orange, green, blue][i]
        draw.rectangle((x0 + 65, sy, x0 + 1035, sy + 115), fill="white", outline=color, width=3)
        draw.text((x0 + 90, sy + 20), a, font=head_font, fill=color)
        draw.text((x0 + 90, sy + 64), b, font=body_font, fill=black)
        if i < len(stages) - 1:
            draw.line((x0 + 550, sy + 118, x0 + 550, sy + 155), fill=gray, width=3)
            draw.polygon([(x0 + 540, sy + 150), (x0 + 560, sy + 150), (x0 + 550, sy + 168)], fill=gray)
        sy += 175
    draw_wrapped(
        draw,
        (x0 + 65, 1350),
        "The overlay at left is not a DINO attention map; it is a deterministic color-index diagnostic used to quantify ROI contamination. Additional attention/error examples are shown in Supplementary Fig. 3.",
        body_font,
        gray,
        1050,
        2,
    )
    save_png_pdf(canvas, upload)


def build_fig7_domain_shift_thumbnails(records: list[dict], rev: dict) -> None:
    upload = FIG_NAME_BY_STEM["fig7_domain_shift_thumb"]
    canvas = Image.new("RGB", (2400, 1520), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(50, True)
    head_font = pil_font(30, True)
    body_font = pil_font(23)
    blue = (68, 114, 196)
    green = (29, 158, 117)
    orange = (216, 90, 48)
    gray = (110, 114, 120)
    black = (32, 33, 36)
    draw.text((70, 40), "Feature-space domain shift with ROI thumbnails", font=title_font, fill=black)
    tsne = ROOT / "tsne_source_vs_data325.png"
    if tsne.exists():
        img = Image.open(tsne).convert("RGB")
        paste_fit(canvas, img, (70, 155, 1290, 1205), (255, 255, 255))
    draw.rectangle((70, 155, 1290, 1205), outline=(218, 224, 230), width=2)
    draw.text((1325, 155), "Source ROI examples", font=head_font, fill=blue)
    source_files = sorted(SOURCE_SAMPLE_DIR.glob("*_bbox.jpg"))[:12]
    for i, path in enumerate(source_files[:8]):
        col, row = i % 4, i // 4
        box = (1325 + col * 245, 210 + row * 310, 1545 + col * 245, 465 + row * 310)
        paste_fit(canvas, Image.open(path).convert("RGB"), box, (250, 250, 250))
        draw.rectangle(box, outline=blue, width=2)
    draw.text((1325, 850), "DATA325 target ROI examples", font=head_font, fill=green)
    target_recs = [choose_by_height(records, 0, 80, "mid"), choose_by_height(records, 80, 120, "mid"), choose_by_height(records, 120, 160, "mid"), choose_by_height(records, 160, 220, "low")]
    target_recs += sorted(records, key=lambda r: r["abs_error_cm"], reverse=True)[:4]
    for i, rec in enumerate(target_recs[:8]):
        col, row = i % 4, i // 4
        box = (1325 + col * 245, 905 + row * 235, 1545 + col * 245, 1110 + row * 235)
        paste_fit(canvas, crop_roi(DATA325_IMAGE_DIR / rec["file_name"], rec["bbox"]), box, (250, 250, 250))
        draw.rectangle(box, outline=green if rec["abs_error_cm"] < 35 else orange, width=2)
    draw.rectangle((70, 1360, 2330, 1480), fill=(247, 249, 251), outline=orange, width=3)
    draw_wrapped(
        draw,
        (105, 1382),
        "The existing feature map contains 20 source and 20 DATA325 ROI features. The thumbnail panels extend the visual context by showing real source and target crops from the benchmark, but no additional feature points are fabricated.",
        body_font,
        black,
        2140,
        3,
    )
    save_png_pdf(canvas, upload)


def build_fig8_ablation_ci(m: dict, rev: dict) -> None:
    upload = FIG_NAME_BY_STEM["fig8_ablation_ci"]
    model_summary = rev["summary"].get("model_summary", {})
    bootstrap = rev["bootstrap"]
    order = ["old", "cls", "patch_mean", "attn", "corrected_camheight", "attn_aug", "attn_aug_tta8", "attn_aug_featurealign", "dann"]
    labels = {
        "old": "Old CLS",
        "cls": "CLS retrain",
        "patch_mean": "Patch mean",
        "attn": "Attention",
        "corrected_camheight": "Camera fixed",
        "attn_aug": "Attn+aug",
        "attn_aug_tta8": "Attn+aug+TTA8",
        "attn_aug_featurealign": "Feature align",
        "dann": "DANN",
    }
    xs, maes, lows, highs, colors = [], [], [], [], []
    for key in order:
        if key not in model_summary or key not in bootstrap:
            continue
        point = float(model_summary[key]["mae_cm"])
        ci = bootstrap[key]["mae_cm"]
        xs.append(labels[key])
        maes.append(point)
        lows.append(point - float(ci["ci95_low"]))
        highs.append(float(ci["ci95_high"]) - point)
        colors.append(f"#{GREEN}" if "attn" in key and "feature" not in key else f"#{BLUE}" if key in ("old", "cls", "patch_mean") else f"#{ORANGE}")
    source_morph = rev.get("source_morphometric", {}).get("models", {})
    for key, label in [("ridge_cv", "Src morph Ridge"), ("random_forest", "Src morph RF")]:
        model = source_morph.get(key, {})
        if not model:
            continue
        point = float(model["summary"]["mae_cm"])
        ci = model["bootstrap_ci"]["mae_cm"]
        xs.append(label)
        maes.append(point)
        lows.append(point - float(ci["ci95_low"]))
        highs.append(float(ci["ci95_high"]) - point)
        colors.append("#7E57C2" if key == "random_forest" else "#8D99AE")
    fig, ax = plt.subplots(figsize=(13, 6.6), dpi=220)
    ax.bar(range(len(xs)), maes, yerr=[lows, highs], capsize=4, color=colors, edgecolor="#2F3A45", linewidth=0.8)
    ax.set_ylabel("DATA325 MAE (cm)")
    ax.set_title("DATA325 model and source-trained baseline ablations with bootstrap 95% CI", fontsize=15, weight="bold")
    ax.set_xticks(range(len(xs)))
    ax.set_xticklabels(xs, rotation=25, ha="right")
    ax.grid(axis="y", color="#D8DEE6", lw=0.6)
    ax.spines[["top", "right"]].set_visible(False)
    best_idx = xs.index("Attn+aug+TTA8") if "Attn+aug+TTA8" in xs else len(xs) - 1
    ax.text(
        0.02,
        0.94,
        f"Best DINOv3-DCF: {fmt(maes[best_idx])} cm",
        transform=ax.transAxes,
        color=f"#{ORANGE}",
        weight="bold",
        bbox=dict(boxstyle="round,pad=0.25", fc="white", ec=f"#{ORANGE}", lw=1.0),
    )
    if "Src morph RF" in xs:
        rf_idx = xs.index("Src morph RF")
        ax.text(
            0.54,
            0.94,
            f"Source RF geometry: {fmt(maes[rf_idx])} cm",
            transform=ax.transAxes,
            color="#7E57C2",
            weight="bold",
            bbox=dict(boxstyle="round,pad=0.25", fc="white", ec="#7E57C2", lw=1.0),
        )
    save_matplotlib(fig, upload)


def build_fig9_resampling_stats(rev: dict) -> None:
    upload = FIG_NAME_BY_STEM["fig9_resampling_stats"]
    seed = rev.get("seed_retraining", {})
    if not seed.get("available"):
        return
    aggregate = seed["aggregate"]
    rows = seed["rows"]
    order = [k for k in ["cls", "patch_mean", "attn", "attn_aug"] if k in aggregate]
    labels = ["CLS", "Patch mean", "Attention", "Attn+aug"]
    colors = [f"#{BLUE}", f"#{BLUE}", f"#{GREEN}", f"#{ORANGE}"]
    fig, axs = plt.subplots(1, 2, figsize=(13.2, 6.2), dpi=220, constrained_layout=True)
    source_mean = [aggregate[k]["source_best_test_mae_mean_cm"] for k in order]
    source_sd = [aggregate[k]["source_best_test_mae_sd_cm"] for k in order]
    data_mean = [aggregate[k]["data325_mae_mean_cm"] for k in order]
    data_sd = [aggregate[k]["data325_mae_sd_cm"] for k in order]
    axs[0].bar(range(len(order)), source_mean, yerr=source_sd, color=colors, edgecolor="#2F3A45", capsize=4)
    axs[0].set_title("(a) Source test MAE across seeds")
    axs[0].set_ylabel("Source test MAE (cm)")
    axs[0].set_xticks(range(len(order)))
    axs[0].set_xticklabels(labels, rotation=15, ha="right")
    axs[1].bar(range(len(order)), data_mean, yerr=data_sd, color=colors, edgecolor="#2F3A45", capsize=4, alpha=0.75)
    for i, method in enumerate(order):
        vals = [row["data325_mae_cm"] for row in rows if row["method"] == method]
        axs[1].scatter(np.full(len(vals), i), vals, color="#202124", s=28, zorder=3)
    axs[1].set_title("(b) DATA325 zero-shot MAE across seeds")
    axs[1].set_ylabel("DATA325 MAE (cm)")
    axs[1].set_xticks(range(len(order)))
    axs[1].set_xticklabels(labels, rotation=15, ha="right")
    axs[1].annotate("Attention mean\n32.43 cm", xy=(2, data_mean[2]), xytext=(2.35, data_mean[2] + 6), arrowprops=dict(arrowstyle="->", color=f"#{ORANGE}", lw=1.3), color=f"#{ORANGE}", weight="bold")
    for ax in axs:
        ax.grid(axis="y", color="#D8DEE6", lw=0.6)
        ax.spines[["top", "right"]].set_visible(False)
    fig.suptitle("DCF-head re-training robustness over three random seeds (TTA1)", fontsize=15, weight="bold")
    save_matplotlib(fig, upload)


def build_fig10_height_bin_ci(rev: dict) -> None:
    upload = FIG_NAME_BY_STEM["fig10_height_bin_ci"]
    bins = rev["height_bins"]
    order = [k for k in ["<80", "80-120", "120-160", ">=160"] if k in bins]
    mae = [bins[k]["point"]["mae_cm"] for k in order]
    mape = [bins[k]["point"]["mape_percent"] for k in order]
    n = [bins[k]["point"]["n"] for k in order]
    lo = [mae[i] - bins[k]["mae_cm"]["ci95_low"] for i, k in enumerate(order)]
    hi = [bins[k]["mae_cm"]["ci95_high"] - mae[i] for i, k in enumerate(order)]
    fig, ax1 = plt.subplots(figsize=(12.8, 6.2), dpi=220)
    bars = ax1.bar(range(len(order)), mae, yerr=[lo, hi], color=[f"#{ORANGE}" if k == "<80" else f"#{GREEN}" for k in order], edgecolor="#2F3A45", capsize=4)
    ax1.set_ylabel("MAE (cm)")
    ax1.set_xticks(range(len(order)))
    ax1.set_xticklabels([f"{k}\n(n={n[i]})" for i, k in enumerate(order)])
    ax1.grid(axis="y", color="#D8DEE6", lw=0.6)
    ax2 = ax1.twinx()
    ax2.plot(range(len(order)), mape, marker="o", color=f"#{BLUE}", lw=2.5)
    ax2.set_ylabel("MAPE (%)")
    ax1.set_title("Height-bin error concentrates in early-stage DATA325 plants", fontsize=15, weight="bold")
    for ax in (ax1, ax2):
        ax.spines["top"].set_visible(False)
    save_matplotlib(fig, upload)


def build_fig11_roi_contamination(rev: dict) -> None:
    upload = FIG_NAME_BY_STEM["fig11_roi_contamination"]
    rows = rev["roi_metrics"]
    fg = np.array([to_float(r, "foreground_fraction") for r in rows])
    bg = np.array([to_float(r, "background_fraction") for r in rows])
    err = np.array([to_float(r, "abs_error_cm") for r in rows])
    std = np.array([to_float(r, "pred_std_cm") for r in rows])
    fig = plt.figure(figsize=(13.2, 8.0), dpi=220, constrained_layout=True)
    gs = fig.add_gridspec(2, 2, width_ratios=[1, 1], height_ratios=[1, 1])
    ax1 = fig.add_subplot(gs[0, 0])
    ax2 = fig.add_subplot(gs[0, 1])
    ax3 = fig.add_subplot(gs[1, 0])
    ax4 = fig.add_subplot(gs[1, 1])
    ax1.scatter(fg * 100, err, c=f"#{GREEN}", edgecolors="white", s=46, alpha=0.9)
    ax1.set_xlabel("Foreground fraction (%)")
    ax1.set_ylabel("Absolute error (cm)")
    ax1.set_title("(a) Foreground vs error")
    ax2.scatter(bg * 100, err, c=f"#{ORANGE}", edgecolors="white", s=46, alpha=0.9)
    ax2.set_xlabel("Background fraction (%)")
    ax2.set_ylabel("Absolute error (cm)")
    ax2.set_title("(b) Background vs error")
    ax3.scatter(std, err, c=f"#{BLUE}", edgecolors="white", s=46, alpha=0.9)
    ax3.set_xlabel("TTA prediction std. (cm)")
    ax3.set_ylabel("Absolute error (cm)")
    ax3.set_title("(c) TTA uncertainty vs error")
    for ax in (ax1, ax2, ax3):
        ax.grid(color="#D8DEE6", lw=0.6)
        ax.spines[["top", "right"]].set_visible(False)
    source_models = rev.get("source_morphometric", {}).get("models", {})
    morph_labels, morph_mae, morph_colors = [], [], []
    for key, label, color in [
        ("ridge_cv", "Source\nRidge", "#8D99AE"),
        ("random_forest", "Source\nRF", "#7E57C2"),
    ]:
        model = source_models.get(key, {})
        if model:
            morph_labels.append(label)
            morph_mae.append(float(model["summary"]["mae_cm"]))
            morph_colors.append(color)
    target_diag = rev.get("morphometric", {}).get("summary", {})
    if target_diag:
        morph_labels.append("Target-label\nLOIO")
        morph_mae.append(float(target_diag.get("mae_cm", float("nan"))))
        morph_colors.append("#D85A30")
    ax4.bar(range(len(morph_labels)), morph_mae, color=morph_colors, edgecolor="#2F3A45", linewidth=0.8)
    ax4.set_title("(d) Morphometric baselines")
    ax4.set_ylabel("DATA325 MAE (cm)")
    ax4.set_xticks(range(len(morph_labels)))
    ax4.set_xticklabels(morph_labels)
    ax4.grid(axis="y", color="#D8DEE6", lw=0.6)
    ax4.spines[["top", "right"]].set_visible(False)
    fig.suptitle("ROI contamination and morphometric diagnostics", fontsize=15, weight="bold")
    save_matplotlib(fig, upload)


def build_fig13_attention_error_clean(records: list[dict], rev: dict) -> None:
    upload = FIG_NAME_BY_STEM["fig8"]
    over = max(records, key=lambda r: float(r["pred_height_cm"]) - float(r["true_height_cm"]))
    under = max(records, key=lambda r: float(r["true_height_cm"]) - float(r["pred_height_cm"]))
    success = min(records, key=lambda r: float(r["abs_error_cm"]))
    sparse = max([r for r in records if float(r["true_height_cm"]) < 80], key=lambda r: float(r["abs_error_cm"]))
    uncertain = max(records, key=lambda r: float(r.get("pred_std_cm", 0.0)))
    cases = [
        ("Successful", success, (29, 158, 117)),
        ("Over-estimated", over, (216, 90, 48)),
        ("Under-estimated", under, (68, 114, 196)),
        ("Sparse early plant", sparse, (216, 90, 48)),
        ("High TTA uncertainty", uncertain, (120, 86, 170)),
    ]
    canvas = Image.new("RGB", (2400, 1550), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(50, True)
    head_font = pil_font(29, True)
    body_font = pil_font(22)
    small_font = pil_font(19)
    black = (32, 33, 36)
    gray = (110, 114, 120)
    draw.text((70, 42), "Error-overlay gallery for real DATA325 ROIs", font=title_font, fill=black)
    draw.text((70, 108), "Each panel uses the real ROI crop, deterministic plant-focus overlay, and the recorded Attn+aug+TTA8 prediction.", font=body_font, fill=gray)
    card_boxes = [
        (90, 210, 825, 610),
        (875, 210, 1610, 610),
        (1660, 210, 2310, 610),
        (275, 745, 1010, 1145),
        (1180, 745, 1915, 1145),
    ]
    for (label, rec, color), box in zip(cases, card_boxes):
        x0, y0, x1, y1 = box
        draw.rectangle((x0, y0, x1, y1), fill=(247, 249, 251), outline=color, width=4)
        draw.text((x0 + 20, y0 + 18), label, font=head_font, fill=color)
        roi = crop_roi(DATA325_IMAGE_DIR / rec["file_name"], rec["bbox"])
        overlay = mask_overlay(roi)
        paste_fit(canvas, roi, (x0 + 24, y0 + 70, x0 + 325, y0 + 315), (250, 250, 250))
        paste_fit(canvas, overlay, (x0 + 355, y0 + 70, x0 + 656, y0 + 315), (250, 250, 250))
        draw.text((x0 + 35, y0 + 324), "ROI crop", font=small_font, fill=gray)
        draw.text((x0 + 365, y0 + 324), "Plant-focus overlay", font=small_font, fill=gray)
        metric_line = (
            f"GT {rec['true_height_cm']:.1f} cm | Pred {rec['pred_height_cm']:.1f} cm | "
            f"AE {rec['abs_error_cm']:.1f} cm | TTA std. {float(rec.get('pred_std_cm', 0.0)):.2f}"
        )
        draw_wrapped(draw, (x0 + 24, y0 + 350), metric_line, body_font, black, x1 - x0 - 48, 1)
    draw.rectangle((90, 1245, 2310, 1390), fill=(248, 238, 233), outline=(216, 90, 48), width=3)
    draw_wrapped(
        draw,
        (125, 1275),
        "The green overlays are deterministic color-index plant-focus diagnostics used for visual QA and ROI contamination analysis. They are not generated images and are not used as ground-truth segmentation labels.",
        body_font,
        black,
        2130,
        2,
    )
    save_png_pdf(canvas, upload)


def build_fig11_qualitative_combined(records: list[dict], rev: dict) -> None:
    stage_path = FIG_OUT / f"{FIG_NAME_BY_STEM['fig_real_stage_error_gallery']}.png"
    error_path = FIG_OUT / f"{FIG_NAME_BY_STEM['fig8']}.png"
    if not stage_path.exists():
        build_real_stage_error_gallery(records)
    if not error_path.exists():
        build_fig13_attention_error_clean(records, rev)
    upload = FIG_NAME_BY_STEM["fig11_qualitative_combined"]
    canvas = Image.new("RGB", (2400, 1650), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(50, True)
    head_font = pil_font(28, True)
    body_font = pil_font(22)
    black = (32, 33, 36)
    gray = (110, 114, 120)
    green = (29, 158, 117)
    orange = (216, 90, 48)
    draw.text((70, 38), "Qualitative stage-wise and attention/error gallery", font=title_font, fill=black)
    draw.text((70, 105), "Panels are assembled from real DATA325 images, manual boxes, recorded predictions, and deterministic overlays.", font=body_font, fill=gray)
    draw.text((70, 160), "(a) Stage-wise DATA325 examples", font=head_font, fill=green)
    paste_fit(canvas, Image.open(stage_path).convert("RGB"), (70, 205, 2330, 855), (255, 255, 255))
    draw.rectangle((70, 205, 2330, 855), outline=(218, 224, 230), width=2)
    draw.text((70, 905), "(b) Error and uncertainty examples", font=head_font, fill=orange)
    paste_fit(canvas, Image.open(error_path).convert("RGB"), (70, 950, 2330, 1510), (255, 255, 255))
    draw.rectangle((70, 950, 2330, 1510), outline=(218, 224, 230), width=2)
    draw_wrapped(
        draw,
        (90, 1548),
        "The combined gallery keeps the main manuscript focused while preserving the real-image evidence behind stage imbalance, over-estimation, under-estimation, sparse plants, and high-uncertainty cases.",
        body_font,
        gray,
        2200,
        2,
    )
    save_png_pdf(canvas, upload)


def build_fig15_release_map(rev: dict) -> None:
    upload = FIG_NAME_BY_STEM["fig15_release_map"]
    canvas = Image.new("RGB", (2400, 1500), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = pil_font(52, True)
    head_font = pil_font(31, True)
    body_font = pil_font(23)
    small_font = pil_font(20)
    blue = (68, 114, 196)
    green = (29, 158, 117)
    orange = (216, 90, 48)
    gray = (110, 114, 120)
    black = (32, 33, 36)
    draw.text((70, 45), "Open release and future deployment map", font=title_font, fill=black)
    draw_wrapped(draw, (70, 115), REPOSITORY_URL, body_font, blue, 1600, 2)

    release_items = [
        ("DATA325 images", "75 raw greenhouse images"),
        ("Annotations", "82 boxes, heights, camera metadata"),
        ("Predictions", "model JSON, bootstrap CI, paired tests"),
        ("Diagnostics", "ROI masks, taxonomy, uncertainty"),
        ("Code", "training, evaluation, figure generation"),
        ("Checkpoints", "DCF heads, no upstream DINOv3 weights"),
    ]
    x0, y0 = 80, 240
    for i, (title, body) in enumerate(release_items):
        col, row = i % 3, i // 3
        x = x0 + col * 760
        y = y0 + row * 245
        color = [green, blue, orange, green, blue, orange][i]
        draw.rectangle((x, y, x + 660, y + 160), fill=(247, 249, 251), outline=color, width=4)
        draw.text((x + 30, y + 25), title, font=head_font, fill=color)
        draw_wrapped(draw, (x + 30, y + 75), body, body_font, black, 590, 2)

    draw.rectangle((120, 825, 2280, 1275), fill=(248, 238, 233), outline=orange, width=4)
    draw.text((155, 855), "Future deployment path", font=head_font, fill=orange)
    stages = [
        ("Automatic ROI detector", "replace manual boxes"),
        ("Segmentation-guided normalization", "SAM-style or crop-specific masks"),
        ("Plant-mask pooling", "reduce background leakage"),
        ("Multi-greenhouse adaptation", "validate across sites/seasons"),
        ("Decision support", "height trend and growth-stage outputs"),
    ]
    sx = 165
    sy = 945
    for i, (a, b) in enumerate(stages):
        color = [blue, green, orange, green, blue][i]
        draw.rectangle((sx, sy, sx + 360, sy + 165), fill="white", outline=color, width=3)
        draw_wrapped(draw, (sx + 22, sy + 24), a, body_font, color, 315, 2)
        draw_wrapped(draw, (sx + 22, sy + 82), b, small_font, black, 315, 2)
        if i < len(stages) - 1:
            draw.line((sx + 365, sy + 82, sx + 425, sy + 82), fill=gray, width=3)
            draw.polygon([(sx + 420, sy + 72), (sx + 440, sy + 82), (sx + 420, sy + 92)], fill=gray)
        sx += 425
    draw_wrapped(
        draw,
        (155, 1335),
        "All evidence panels in the manuscript are traceable to released real images, annotations, prediction JSON, CSV diagnostics, or deterministic figure-generation scripts. Conceptual deployment graphics are code-drawn and not used as experimental evidence.",
        body_font,
        gray,
        2070,
        2,
    )
    save_png_pdf(canvas, upload)


def build_real_image_figures() -> None:
    records = load_data325_records()
    rev = load_revision_results()
    build_real_image_protocol_figure(records)
    build_fig2_benchmark_comparison(records, rev)
    build_fig3_distribution(records, rev)
    build_fig4_mindmap_workflow()
    build_fig4_preprocessing(records, rev)
    build_fig6_attention_roi(records, rev)
    build_fig7_domain_shift_thumbnails(records, rev)
    build_fig8_ablation_ci(metrics(), rev)
    build_fig9_resampling_stats(rev)
    build_fig10_height_bin_ci(rev)
    build_real_photo_prediction_matrix(records, rev)
    build_fig11_roi_contamination(rev)
    build_real_stage_error_gallery(records)
    build_fig13_attention_error_clean(records, rev)
    build_fig11_qualitative_combined(records, rev)
    build_extended_real_photo_matrix(records, rev)
    build_fig15_release_map(rev)
    move_supplementary_figures()


def move_supplementary_figures() -> None:
    SUPP_FIG_OUT.mkdir(parents=True, exist_ok=True)
    for _, upload, _, _ in SUPPLEMENTARY_FIGURES:
        for ext in ("png", "pdf"):
            src = FIG_OUT / f"{upload}.{ext}"
            if not src.exists():
                continue
            dst = SUPP_FIG_OUT / src.name
            shutil.copy2(src, dst)
            src.unlink()


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph(style="Els-Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(TITLE)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13.5)
    run.font.color.rgb = RGBColor(0, 0, 0)
    for line in [
        f"{AUTHOR_NAME} a, {COAUTHOR_NAME} a,*",
        f"a {AFFILIATION}",
        f"First author email: {AUTHOR_EMAIL}",
        f"* Corresponding author. Email: {CORRESPONDING_EMAIL}",
    ]:
        p2 = doc.add_paragraph(style="Els-Affiliation")
        p2.paragraph_format.line_spacing = 2.0
        p2.paragraph_format.space_after = Pt(0)
        r = p2.add_run(line)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(9.5)


def validate_manuscript_title_page(path: Path) -> None:
    """Keep submission-system metadata out of the manuscript source title page."""
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if "Abstract" not in paragraphs:
        raise RuntimeError(f"{path.name}: missing Abstract heading for title-page QA")
    front_matter = "\n".join(paragraphs[: paragraphs.index("Abstract")])
    full_text = "\n".join(paragraphs)
    required = [
        TITLE,
        f"{AUTHOR_NAME} a, {COAUTHOR_NAME} a,*",
        f"a {AFFILIATION}",
        f"First author email: {AUTHOR_EMAIL}",
        f"* Corresponding author. Email: {CORRESPONDING_EMAIL}",
    ]
    forbidden_front_matter = [
        "Article type",
        "Original research paper",
        "Target journal",
        "Highlights file",
        "Graphical abstract",
        "Figure files",
        "Repository",
        "Planned archive",
        "Submission package",
    ]
    forbidden_full_text = [
        "Article type",
        "Target journal",
        "Highlights file",
        "Figure files",
        "Planned archive",
        "Submission package",
    ]
    missing = [item for item in required if item not in front_matter]
    if missing:
        raise RuntimeError(f"{path.name}: title page missing required fields: {missing}")
    found_front = [item for item in forbidden_front_matter if item in front_matter]
    if found_front:
        raise RuntimeError(f"{path.name}: title page contains submission metadata: {found_front}")
    found_full = [item for item in forbidden_full_text if item in full_text]
    if found_full:
        raise RuntimeError(f"{path.name}: manuscript contains submission metadata: {found_full}")


def validate_manuscript_equations(path: Path) -> None:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    text = "\n".join(parts)
    missing = [f"Eq. ({idx})" for idx in range(1, 17) if f"Eq. ({idx})" not in text]
    if missing:
        raise RuntimeError(f"{path.name}: missing display equation labels: {', '.join(missing)}")
    if "Mathematical formulation" in text:
        raise RuntimeError(f"{path.name}: standalone mathematical formulation section still present")
    if "2.4. Diagnostics and statistics" in text:
        raise RuntimeError(f"{path.name}: diagnostics section was not renumbered to 2.3")
    if "native Word equation objects, not screenshots" in text:
        raise RuntimeError(f"{path.name}: internal equation-rendering note still present in manuscript")
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
    math_count = xml.count("<m:oMath")
    if math_count < 16:
        raise RuntimeError(f"{path.name}: expected native Word math objects, found {math_count}")


def revision_refs() -> list[str]:
    return [
        "Ariza-Sentis, M., et al., 2024. Object detection and tracking in precision farming: a systematic review. Computers and Electronics in Agriculture 219, 108757. doi:10.1016/j.compag.2024.108757.",
        "Caron, M., Touvron, H., Misra, I., et al., 2021. Emerging properties in self-supervised vision transformers. Proceedings of the IEEE/CVF International Conference on Computer Vision, 9650-9660.",
        "Chang, A., Jung, J., Maeda, M.M., Landivar, J., 2017. Crop height monitoring with digital imagery from Unmanned Aerial System (UAS). Computers and Electronics in Agriculture 141, 232-237. doi:10.1016/j.compag.2017.07.008.",
        "Che, Y., Gu, Y., Bai, D., Li, D., Li, J., Zhao, C., et al., 2024. Accurately estimate soybean growth stages from UAV imagery by accounting for spatial heterogeneity and climate factors across multiple environments. Computers and Electronics in Agriculture 225, 109313. doi:10.1016/j.compag.2024.109313.",
        "Dosovitskiy, A., Beyer, L., Kolesnikov, A., et al., 2021. An image is worth 16x16 words: transformers for image recognition at scale. International Conference on Learning Representations.",
        "Ganin, Y., Ustinova, E., Ajakan, H., et al., 2016. Domain-adversarial training of neural networks. Journal of Machine Learning Research 17, 1-35.",
        "Gulrajani, I., Lopez-Paz, D., 2021. In search of lost domain generalization. International Conference on Learning Representations.",
        "Jayasuriya, N., Guo, Y., Hu, W., Ghannoum, O., 2024. Machine vision based plant height estimation for protected crop facilities. Computers and Electronics in Agriculture 218, 108669. doi:10.1016/j.compag.2024.108669.",
        "Kamilaris, A., Prenafeta-Boldu, F.X., 2018. Deep learning in agriculture: a survey. Computers and Electronics in Agriculture 147, 70-90. doi:10.1016/j.compag.2018.02.016.",
        "Kim, W., Lee, D., Kim, Y., Kim, T., Lee, W., Choi, C., 2021. Stereo-vision-based crop height estimation for agricultural robots. Computers and Electronics in Agriculture 181, 105937. doi:10.1016/j.compag.2020.105937.",
        "Kirillov, A., Mintun, E., Ravi, N., Mao, H., Rolland, C., Gustafson, L., et al., 2023. Segment anything. IEEE/CVF International Conference on Computer Vision, 4015-4026.",
        "Koh, P.W., Sagawa, S., Marklund, H., Xie, S.M., Zhang, M., Balsubramani, A., et al., 2021. WILDS: a benchmark of in-the-wild distribution shifts. Proceedings of Machine Learning Research 139, 5637-5664.",
        "Li, L., Zhang, Q., Huang, D., 2014. A review of imaging techniques for plant phenotyping. Sensors 14, 20078-20111.",
        "Li, M., Sui, R., Meng, Y., Yan, H., 2019. A real-time fuzzy decision support system for alfalfa irrigation. Computers and Electronics in Agriculture 163, 104870. doi:10.1016/j.compag.2019.104870.",
        "Li, Z., Guo, R., Li, M., Chen, Y., Li, G., 2020. A review of computer vision technologies for plant phenotyping. Computers and Electronics in Agriculture 176, 105672. doi:10.1016/j.compag.2020.105672.",
        "Liu, H., Bruning, B., Garnett, T., Berger, B., 2020. Hyperspectral imaging and 3D technologies for plant phenotyping: from satellite to close-range sensing. Computers and Electronics in Agriculture 175, 105621. doi:10.1016/j.compag.2020.105621.",
        "Oquab, M., Darcet, T., Moutakanni, T., et al., 2024. DINOv2: learning robust visual features without supervision. Transactions on Machine Learning Research.",
        "Patricio, D.I., Rieder, R., 2018. Computer vision and artificial intelligence in precision agriculture for grain crops: a systematic review. Computers and Electronics in Agriculture 153, 69-81. doi:10.1016/j.compag.2018.08.001.",
        "Reena, Doonan, J.H., Williams, K., Corke, F.M.K., Zhang, H., Batke, S., Liu, Y., 2025. Wheat3D PartNet: annotated dataset for 3D wheat part segmentation. Computers and Electronics in Agriculture 238, 110697. doi:10.1016/j.compag.2025.110697.",
        "Simeoni, O., Vo, H.V., Seitzer, M., et al., 2025. DINOv3. arXiv:2508.10104. doi:10.48550/arXiv.2508.10104.",
        "Sun, B., Saenko, K., 2016. Deep CORAL: correlation alignment for deep domain adaptation. European Conference on Computer Vision Workshops, 443-450. doi:10.1007/978-3-319-49409-8_35.",
        "Veramendi, W., Cruvinel, P., 2024. Method for maize plants counting and crop evaluation based on multispectral images analysis. Computers and Electronics in Agriculture 216, 108470. doi:10.1016/j.compag.2023.108470.",
        "Xie, T., Li, J., Yang, C., Jiang, Z., Chen, Y., Guo, L., Zhang, J., 2021. Crop height estimation based on UAV images: methods, errors, and strategies. Computers and Electronics in Agriculture 185, 106155. doi:10.1016/j.compag.2021.106155.",
        "Xing, Z., Zhang, Z., Shi, R., Guo, Q., Zeng, C., 2023. Filament-necking localization method via combining improved PSO with rotated rectangle algorithm for safflower-picking robots. Computers and Electronics in Agriculture 215, 108464. doi:10.1016/j.compag.2023.108464.",
    ]


def build_manuscript(m: dict) -> Path:
    rev = load_revision_results()
    model_summary = rev["summary"].get("model_summary", {})
    best = rev["summary"].get("best_model_metrics", {})
    roi_summary = rev["roi_summary"]
    morph = rev["morphometric"].get("summary", {})
    source_morph = rev.get("source_morphometric", {})
    source_models = source_morph.get("models", {})
    ridge_summary = source_models.get("ridge_cv", {}).get("summary", {})
    rf_summary = source_models.get("random_forest", {}).get("summary", {})
    tax = rev["taxonomy_summary"]
    unc = rev["uncertainty"].get("overall", {})
    old_mae = model_summary.get("old", {}).get("mae_cm", 41.757666425007145)
    attn_mae = model_summary.get("attn", {}).get("mae_cm", 30.407656506794254)
    corrected_mae = model_summary.get("corrected_camheight", {}).get("mae_cm", 30.374317355272247)
    patch_mae = model_summary.get("patch_mean", {}).get("mae_cm", 33.53448021121142)
    cls_mae = model_summary.get("cls", {}).get("mae_cm", 45.86054881026105)
    attn_aug_mae = model_summary.get("attn_aug", {}).get("mae_cm", 29.894323581602514)
    best_mae = best.get("mae_cm", 29.573550898854325)
    best_rmse = best.get("rmse_cm", 38.97683219161265)
    best_median = best.get("median_abs_error_cm", 23.64422082901001)
    best_mape = best.get("mape_percent", 36.13797231033507)

    doc = setup_doc()
    add_title_page(doc)
    doc.add_heading("Abstract", level=1)
    abstract = (
        "External-greenhouse crop-height estimation is difficult because model training images and deployment images differ in plant stage, background, imaging geometry, and ROI quality. "
        "We present DATA325 as a reproducible diagnostic benchmark for external maize height evaluation and test DINOv3-DiffCorn-Fusion (DINOv3-DCF), an ROI-level pipeline that connects frozen DINOv3 features to a DCF regression head. "
        "The source hand-box set contains 156 maize ROIs, and DATA325-v0.1 contains 75 raw greenhouse photographs with an evaluated subset of 82 boxes from 25 completed independent images. "
        f"Compared with the original CLS baseline ({fmt(old_mae)} cm MAE), attention-weighted patch pooling reduced zero-shot DATA325 MAE to {fmt(attn_mae)} cm, and the best attention plus augmentation plus TTA8 variant reached {fmt(best_mae)} cm MAE, {fmt(best_rmse)} cm RMSE, and {fmt(best_mape)}% MAPE. "
        f"A source-trained random-forest morphometric baseline reached {fmt(rf_summary.get('mae_cm', 27.10317886989101))} cm MAE, showing that manual bbox geometry is a strong cue in this benchmark. "
        "Bootstrap confidence intervals, paired per-box comparisons, ROI foreground diagnostics, uncertainty analysis, and error taxonomy show that the remaining gap concentrates in early-stage plants below 80 cm and is not explained by camera-height correction, feature-statistic alignment, or the tested DANN setup. "
        "The open release provides images, annotations, predictions, diagnostics, scripts, and selected DCF checkpoints for future external-greenhouse phenotyping research."
    )
    add_p(doc, abstract)
    add_p(doc, "Keywords: maize phenotyping; plant height; foundation model; DINOv3; domain shift; attention pooling; benchmark")

    doc.add_heading("1. Introduction", level=1)
    for para in [
        "Plant height is a common phenotyping trait because it reflects growth stage, biomass accumulation, lodging risk, and crop-management status. Automated image-based height estimation has been studied through UAS crop-height models, UAV error analysis, stereo vision for agricultural robots, and protected-crop machine vision (Chang et al., 2017; Xie et al., 2021; Kim et al., 2021; Jayasuriya et al., 2024). These systems show that computer and electronic imaging can reduce manual measurement burden, but they also show that camera geometry, environment, and data-processing assumptions strongly affect the final trait estimate.",
        "Agricultural computer vision has moved from hand-engineered features to deep learning, object detection, and phenotyping-specific pipelines (Patricio and Rieder, 2018; Kamilaris and Prenafeta-Boldu, 2018; Li et al., 2020; Ariza-Sentis et al., 2024). CEA studies on alfalfa decision support, maize plant counting, safflower localization, and soybean growth-stage estimation illustrate that agricultural vision systems must deal with crop-specific structure, complex backgrounds, and multi-environment variation (Li et al., 2019; Veramendi and Cruvinel, 2024; Xing et al., 2023; Che et al., 2024).",
        "A central risk is data leakage and weak external validation. Random image splits can put visually near-identical plants, dates, benches, or backgrounds in both training and testing. Broader distribution-shift work shows that source-domain accuracy is not a reliable substitute for deployment evaluation (Koh et al., 2021; Gulrajani and Lopez-Paz, 2021). Greenhouse maize height estimation is especially sensitive because early plants occupy small portions of the ROI and share the crop with pots, labels, substrate, benches, and shadows.",
        "Foundation-model features are attractive because self-supervised ViTs learn reusable visual representations and attention maps with emergent localization behavior (Dosovitskiy et al., 2021; Caron et al., 2021; Oquab et al., 2024; Simeoni et al., 2025). However, a global CLS token can summarize background context together with plant tissue. For ROI-level phenotyping, attention-weighted patch pooling may better match the biological object than a single global token, while keeping the visual backbone frozen and reproducible.",
        "For plant-height estimation, the external-validation problem is not only a matter of model accuracy. A credible agricultural computer vision paper must specify where the images came from, how plant instances were isolated, how height labels and camera metadata were linked to each ROI, and whether the test images are genuinely independent of the training images. Without this protocol detail, a strong MAE can be caused by background similarity, leakage across adjacent frames, or a hidden dependence on one camera setup rather than by transferable plant representation.",
        "DATA325-v0.1 was therefore designed as a small but auditable diagnostic target set. Its value is not the number of images alone; it is the combination of real greenhouse photographs, manually inspectable boxes, measured height labels, camera-height metadata, per-box predictions, and released sidecar files that make each reported error traceable. The v0.1 label is used deliberately because the current release contains 75 raw photographs but only a completed evaluated subset of 25 annotated images and 82 plant boxes. This framing follows the spirit of dataset-oriented CEA papers: the benchmark must be understandable before the model result can be trusted.",
        "This study presents DATA325-v0.1 as an external-greenhouse benchmark and DINOv3-DCF as the evaluation pipeline. Following the logic of Wheat3D PartNet, which first establishes data resource quality before presenting model results (Reena et al., 2025), we first document DATA325-v0.1 acquisition, annotation, distributions, and preprocessing. We then evaluate feature aggregation, statistical robustness, ROI contamination, uncertainty, and failure categories. The goal is not to claim a complete deployed height system, but to identify what frozen foundation features can and cannot solve under zero-shot external-greenhouse transfer.",
    ]:
        add_p(doc, para)

    doc.add_heading("2. Materials and methods", level=1)
    doc.add_heading("2.1. DATA325-v0.1 acquisition, annotation, and open release", level=2)
    for para in [
        "Model development used a source-domain hand-bounding-box maize dataset; external testing used the independent DATA325-v0.1 target greenhouse. DATA325-v0.1 target labels were used only for final evaluation, statistical diagnostics, and visualization, not for training, hyperparameter selection, or early stopping. This defines the reported results as zero-shot external-greenhouse evaluation.",
        "DATA325-v0.1 currently contains 75 raw greenhouse photographs. The evaluated subset contains 25 completed images and 82 manually annotated plant boxes with measured plant heights and camera-height metadata; the remaining raw photographs are released as collection context rather than scored test instances. The source hand-box set contains 156 ROIs and spans 64-270 cm, while the evaluated DATA325-v0.1 subset spans 30-178 cm and contains more early-stage plants. Manual ROIs were used as an experimental control because automatic detection, tracking, and segmentation errors can confound trait-regression evaluation under agricultural clutter (Ariza-Sentis et al., 2024; Xing et al., 2023; Kirillov et al., 2023).",
        "The annotation unit is one maize plant bounding box. Each DATA325 record links an image identifier, file name, box identifier, bounding-box coordinates, measured plant height, camera-height value, predicted height, absolute error, and TTA standard deviation. This record-level design makes the benchmark inspectable: a reader can move from a table row to the raw photograph, verify the box placement, crop the ROI, and reproduce the same error calculation.",
        "The source and target sets differ in both biology and imaging context. The source hand-box data contain taller plants and two camera-height regimes, whereas DATA325 includes many shorter, sparse plants whose leaves occupy a smaller fraction of the ROI. The target images also include greenhouse-specific clutter such as pots, labels, substrate, bench structure, and shadow patterns. These differences are exactly the conditions under which random image splits are likely to overestimate deployment performance.",
        "The source split used for DCF-head checkpoint selection is a random ROI-level 80/20 split of the 156 source ROIs. It is reported for source-domain fitting transparency, not as an independent greenhouse validation result. The paper's external claim rests on DATA325, whose images and labels are held outside training and checkpoint selection.",
        "Quality control was deliberately conservative. Boxes were kept manual to isolate representation and height-regression behavior before adding detector variability. ROI crops were inspected visually in protocol figures, and deterministic color-index foreground estimates were used to summarize whether a crop was plant-dominated or background-dominated. These foreground estimates are not treated as segmentation labels; they are a diagnostic proxy for ROI contamination.",
        f"The release repository ({REPOSITORY_URL}) is organized to expose raw DATA325 images, cleaned annotations, prediction JSON files, statistical diagnostics, figure scripts, and selected DCF checkpoints. Upstream DINOv3 weights are not redistributed and must be obtained from the upstream source under its license.",
    ]:
        add_p(doc, para)
    add_named_figure(doc, "fig_real_protocol")
    add_named_figure(doc, "fig2_benchmark_comparison")
    add_named_figure(doc, "fig3_distribution")

    doc.add_heading("2.2. DINOv3-DCF model and feature aggregation", level=2)
    add_p(doc, "DINOv3-DiffCorn-Fusion (DINOv3-DCF) receives a manually annotated plant ROI and camera-height context. Each crop is resized to 224 x 224 pixels and processed by a frozen DINOv3 ViT-L backbone. The ROI-to-token step is written as:")
    add_display_equation(doc, "r_b = R(crop(I,b)),   {z_cls,z_1,...,z_N}, A = F_DINO(r_b),   z_i ∈ ℝ^1024", 1)
    add_p(doc, "Here b is the manual ROI box, r_b is the resized crop, z_cls and z_i are the final DINOv3 CLS and patch tokens, and A is the final-layer attention tensor. The DINOv3 backbone remains frozen in all reported experiments; only the DCF head is trained for the main source-domain models, except for the DANN diagnostic where an additional domain classifier is added for adversarial training.")
    add_p(doc, "Three frozen-backbone aggregation modes were compared. CLS pooling uses the final CLS token, patch-mean pooling averages all patch tokens, and attention-weighted pooling uses final-layer CLS-to-patch attention averaged across heads:")
    add_display_equation(doc, "f_CLS = z_cls,      f_mean = (1/N) ∑_(i=1)^N z_i", 2)
    add_display_equation(doc, "a_i = (1/H) ∑_(j=1)^H A_(j,cls→i),      α_i = a_i / ∑_(k=1)^N a_k", 3)
    add_display_equation(doc, "f_attn = ∑_(i=1)^N α_i z_i", 4)
    add_p(doc, "The descriptor dimension is unchanged across pooling modes, so the experiment isolates the aggregation choice rather than redesigning the regressor. The selected visual descriptor is then concatenated with camera height normalized by 200 cm:")
    add_display_equation(doc, "x = [f ; h_cam/200] ∈ ℝ^1025", 5)
    add_p(doc, "The camera-conditioned DCF head is a DiffCorn-Fusion MLP with Linear(1025, 512), BatchNorm, ReLU, Linear(512, 256), BatchNorm, ReLU, and Linear(256, 64). Its raw output is mapped to fixed latent ranges by a sigmoid scaling step:")
    add_display_equation(doc, "u = W_3 φ(BN_2(W_2 φ(BN_1(W_1 x))))", 6)
    add_display_equation(doc, "p = p_min + (p_max − p_min) ⊙ σ(u),      p ∈ ℝ^64", 7)
    add_p(doc, "The 64-dimensional output is organized as 16 phytomer-inspired groups with four latent variables per group. Only the derived plant height is supervised by measured height labels; the non-height latent variables should not be interpreted as validated leaf angle, leaf length, or inclination measurements:")
    add_display_equation(doc, "p = {p_(k,m)}_(k=1..16,m=1..4),      h_hat = ∑_(k=1)^16 p_(k,1)", 8)
    add_p(doc, "In Eq. (8), p_(k,1) is the internode-length-related latent component used for the height sum. The other three components in each group keep the DCF head in a structured latent space but are not evaluated as annotated leaf-angle, leaf-length, or inclination traits in DATA325-v0.1.")
    add_p(doc, "DCF-head training used Huber loss on derived plant height with delta = 1.0. The released best attention checkpoints record an ADEL prior configuration where available (L2 prior, lambda = 0.01); the three-seed robustness run disables this prior to isolate the effect of feature aggregation over fixed feature bundles.")
    add_display_equation(doc, "ℓ_δ(e) = 0.5 e^2 if |e| ≤ δ;      ℓ_δ(e) = δ(|e| − 0.5δ) otherwise,      δ = 1", 9)
    add_display_equation(doc, "L = (1/B) ∑_(i=1)^B ℓ_1(h_hat_i − h_i) + λ_prior L_prior", 10)
    add_p(doc, "Optimization used Adam with learning rate 5e-4 and weight decay 1e-4, batch size 16, 300 epochs, gradient clipping at norm 1.0, and cosine learning-rate scheduling to 1e-6. Checkpoint selection used source-domain random ROI-level test MAE. DATA325 labels were never used for checkpoint selection, so the external target set remains a held-out diagnostic benchmark.")
    doc.add_page_break()
    add_named_figure(doc, "fig1")
    add_named_figure(doc, "fig6_attention_roi")

    doc.add_heading("2.3. Diagnostics and statistics", level=2)
    add_p(doc, "Visual augmentation and TTA8 were applied as limited robustness interventions. Training-time augmentation perturbed color and peripheral ROI appearance. TTA8 averaged one original prediction and seven color-perturbed predictions at inference, and prediction standard deviation from these T predictions was retained for uncertainty diagnostics:")
    add_display_equation(doc, "h_bar_i = (1/T) ∑_(t=1)^T h_hat_i^(t),      s_i = sqrt((1/(T−1)) ∑_(t=1)^T (h_hat_i^(t) − h_bar_i)^2)", 11)
    add_p(doc, "Evaluation used MAE, RMSE, and MAPE at the DATA325 box level:")
    add_display_equation(doc, "MAE = (1/n) ∑_(i=1)^n |h_i − h_hat_i|,      RMSE = sqrt((1/n) ∑_(i=1)^n (h_i − h_hat_i)^2)", 12)
    add_display_equation(doc, "MAPE = (100/n) ∑_(i=1)^n |(h_i − h_hat_i)/h_i|", 13)
    add_p(doc, "ROI contamination was quantified using non-generative color-index masks. The diagnostic mask estimates foreground fraction, background fraction, bbox fill ratio, bbox area fraction, aspect ratio, brightness, and edge contact. It is used only for error analysis and figure QA, not as a training label or a replacement for plant segmentation:")
    add_display_equation(doc, "ExG = 2G − R − B,      M = 1{ExG > q_0.60(ExG) or G > 1.04R and G > 1.04B}", 14)
    add_p(doc, "Two morphometric checks were included. First, a source-trained baseline fitted RidgeCV and a small deterministic RandomForestRegressor on source hand-bbox geometry, camera height, and ExG mask features, then evaluated those models directly on DATA325 without using DATA325 labels for training. The source-trained feature vector was:")
    add_display_equation(doc, "g = [w/W, h_b/H, wh_b/(WH), w/h_b, h_cam, h_b/h_cam, ρ_fg, v_green, c_y]", 15)
    add_p(doc, "Second, a leave-one-image-out ridge regression using DATA325 morphometric and mask features was retained as a target-label diagnostic lower bound, not as a zero-shot comparator. Bootstrap 95% confidence intervals were computed for MAE, RMSE, and MAPE with 5000 resamples. Paired bootstrap differences and exact sign tests compared per-box absolute errors against the Attn+aug+TTA8 model:")
    add_display_equation(doc, "CI_95(Q) = [q_0.025({Q*}), q_0.975({Q*})],      d_i = |e_i^A| − |e_i^B|", 16)
    add_p(doc, "Independent DCF-head re-training was also run for CLS, patch-mean, attention, and attention+augmentation feature modes using three random seeds (11, 42, and 73), followed by DATA325 zero-shot evaluation with TTA1. The existing negative controls were retained: per-image camera-height correction, bbox geometry, attention-map geometry priors, feature-statistic alignment inspired by Deep CORAL, and a DANN-style adversarial model (Sun and Saenko, 2016; Ganin et al., 2016).")
    add_p(doc, "The feature-statistic alignment and DANN results are interpreted as domain-gap diagnostics rather than as the main deployment protocol. Their purpose is to test whether a simple marginal feature correction or adversarial source-target confusion is sufficient to explain the remaining error. All confidence intervals and paired tests operate at the DATA325 box level, so every resampled unit corresponds to a released annotated plant instance.")

    doc.add_heading("3. Results", level=1)
    doc.add_heading("3.1. DATA325 exposes feature-space domain shift", level=2)
    add_p(doc, f"DINOv3 feature analysis showed a source-target centroid distance of {fmt(m['centroid_distance'])}, relative mean shift of {fmt(m['mean_shift'])}%, and relative standard-deviation shift of {fmt(m['std_shift'])}%. The thumbnails in Fig. 6 show that the embedding separation corresponds to visible greenhouse and growth-stage differences rather than an abstract numerical artifact.")
    add_named_figure(doc, "fig7_domain_shift_thumb")

    doc.add_heading("3.2. Attention pooling gives the largest zero-shot gain", level=2)
    add_p(doc, f"The original CLS baseline reached {fmt(old_mae)} cm DATA325 MAE. CLS retraining reached {fmt(cls_mae)} cm MAE, patch-mean pooling reached {fmt(patch_mae)} cm, and attention-weighted pooling reached {fmt(attn_mae)} cm. Corrected camera-height metadata produced {fmt(corrected_mae)} cm MAE, which is numerically close to attention pooling but represents a metadata diagnostic rather than the main attention result. The best Attn+aug+TTA8 model reached {fmt(best_mae)} cm MAE, {fmt(best_rmse)} cm RMSE, {fmt(best_median)} cm median AE, and {fmt(best_mape)}% MAPE.")
    add_p(doc, f"The gain pattern is important. Moving from CLS to patch-mean pooling reduced error by {fmt(cls_mae - patch_mae)} cm relative to the retrained CLS model, showing that the token interface matters even when the regression head and source split are held fixed. Moving from patch-mean to attention-weighted pooling reduced error by another {fmt(patch_mae - attn_mae)} cm, suggesting that plant-focused patch weighting is more useful than treating all ROI patches equally. Augmentation and TTA8 then provided smaller additional gains, from {fmt(attn_aug_mae)} to {fmt(best_mae)} cm.")
    paired = rev.get("paired", {})
    attn_pair = paired.get("attn", {})
    attn_aug_pair = paired.get("attn_aug", {})
    add_p(doc, f"Paired tests support a cautious interpretation. The best Attn+aug+TTA8 setting is the lowest observed DINOv3-DCF error, but its advantage over plain attention is not statistically stable: the attention-minus-best paired difference has a 95% interval from {fmt(attn_pair.get('ci95_low', -1.403))} to {fmt(attn_pair.get('ci95_high', 3.035))} cm and sign-test p={fmt(attn_pair.get('sign_test_p', 0.74065), 3)}. The attention+augmentation comparison is also marginal (p={fmt(attn_aug_pair.get('sign_test_p', 0.097), 3)}). The robust conclusion is therefore that attention pooling is the main gain; augmentation and TTA8 are best observed settings rather than the central contribution.")
    seed = rev.get("seed_retraining", {})
    if seed.get("available"):
        agg = seed["aggregate"]
        add_p(doc, f"Three-seed DCF-head re-training supports the same feature-aggregation conclusion under TTA1 evaluation. Mean DATA325 MAE across seeds was {fmt(agg['cls']['data325_mae_mean_cm'])} cm for CLS, {fmt(agg['patch_mean']['data325_mae_mean_cm'])} cm for patch-mean, {fmt(agg['attn']['data325_mae_mean_cm'])} cm for attention, and {fmt(agg['attn_aug']['data325_mae_mean_cm'])} cm for attention+augmentation. Source test MAE remained low for all four modes (approximately 1.05-1.52 cm), but this comes from the source random ROI-level split and is treated as evidence that the head fits the source domain, not as independent deployment validation.")
    else:
        add_p(doc, "Bootstrap intervals and paired differences confirm that the improvement is distributed across DATA325 boxes. Independent multi-seed re-training was not available for this build.")
    add_named_figure(doc, "fig8_ablation_ci")
    add_named_figure(doc, "fig9_resampling_stats")
    add_ablation_table(doc, m)

    doc.add_heading("3.3. Early-stage plants remain the main failure mode", level=2)
    early = rev["height_bins"].get("<80", {}).get("point", {})
    add_p(doc, f"Height-bin analysis identifies early-stage plants as the main bottleneck. For DATA325 boxes below 80 cm, MAE was {fmt(early.get('mae_cm', 31.80))} cm and MAPE was {fmt(early.get('mape_percent', 55.11))}%. Taller bins had lower relative error, consistent with stronger plant-structure evidence and less background-dominated ROI content.")
    add_named_figure(doc, "fig10_height_bin_ci")
    add_height_bin_table(doc, m)

    doc.add_heading("3.4. ROI contamination and uncertainty contribute but do not explain the gap", level=2)
    fg_corr = roi_summary.get("correlations_with_abs_error", {}).get("foreground_fraction", {})
    bg_corr = roi_summary.get("correlations_with_abs_error", {}).get("background_fraction", {})
    std_corr = unc.get("pred_std_vs_abs_error", {})
    add_p(doc, f"The deterministic foreground mask estimated mean foreground fraction at {fmt(roi_summary.get('foreground_fraction_mean', 0) * 100)}%. Correlation with absolute error was weak for foreground fraction (Pearson r={fmt(fg_corr.get('pearson_r', 0), 3)}) and background fraction (r={fmt(bg_corr.get('pearson_r', 0), 3)}). TTA prediction standard deviation was also weakly correlated with absolute error (r={fmt(std_corr.get('pearson_r', 0), 3)}).")
    source_models = source_morph.get("models", {})
    ridge_summary = source_models.get("ridge_cv", {}).get("summary", {})
    rf_summary = source_models.get("random_forest", {}).get("summary", {})
    add_p(doc, f"The source-trained morphometric baseline provides a direct agricultural-geometry comparator that does not use DATA325 labels for fitting. RidgeCV reached {fmt(ridge_summary.get('mae_cm'))} cm MAE and the small RandomForestRegressor reached {fmt(rf_summary.get('mae_cm'))} cm MAE on DATA325. The random-forest result shows that manual bbox geometry and green-mask extent can be very strong cues under this benchmark, so DINOv3-DCF should be interpreted as a feature-transfer diagnostic rather than as the sole best possible height estimator.")
    add_p(doc, f"The target-label morphometric baseline reached {fmt(morph.get('mae_cm', 0))} cm MAE under leave-one-image-out evaluation. Because this baseline uses DATA325 labels, it is a diagnostic lower-bound check, not a deployable zero-shot result. Together with the source-trained baseline, it shows that bbox geometry is important but does not remove the need to diagnose early-stage and domain-shift failures.")
    add_p(doc, "Detailed ROI-contamination scatter plots and morphometric-baseline panels are provided as Supplementary Fig. 5 to preserve the audit trail for geometry and foreground-mask diagnostics.")

    doc.add_heading("3.5. Real-image galleries and negative controls localize the residual problem", level=2)
    add_p(doc, "The qualitative galleries show that high-error examples are often early, sparse, background-heavy, or TTA-unstable. The rule-based error taxonomy assigned 32 boxes to early-stage sparse structure, 15 to unstable TTA prediction, 10 to bbox ambiguity, and 25 to residual cross-domain shift. These categories are intended as diagnostic labels for future benchmark development rather than as ground-truth biological classes.")
    add_p(doc, "Figure 10 uses a real-photo matrix format to show low-, median-, and high-error cases within early, mid, and tall height ranges. The layout follows a ground-truth/prediction/difference logic, but the difference panel is explicitly a height-error diagnostic that combines absolute error, signed error, TTA uncertainty, foreground mask, and taxonomy label. It is not a pixel-level segmentation difference or 3D reconstruction error.")
    add_named_figure(doc, "fig12_real_photo_matrix")
    add_p(doc, "The full stage-wise and attention/error gallery is provided as Supplementary Fig. 6, and an extended real-photo matrix is provided as Supplementary Fig. 8. These supplementary panels preserve traceable real-image evidence for cases that are not shown in the main benchmark figures.")
    add_p(doc, f"Negative controls further narrow the interpretation. Camera-height correction changed MAE by only about 0.03 cm. Bbox geometry reached Pearson r={fmt(m['bbox_r'], 3)} and MAE {fmt(m['bbox_mae'])} cm. Feature-statistic alignment worsened MAE to 45.98 cm, and the tested DANN model reached 34.20 cm MAE while the domain classifier remained near {fmt(m['dann_final_acc'] * 100, 0)}% accuracy.")
    add_p(doc, "Diagnostic negative controls are summarized in Supplementary Fig. 7. They are discussed as explanatory checks rather than as competing deployment models.")

    doc.add_heading("4. Discussion", level=1)
    for para in [
        "The benchmark component matters because DATA325 is an external target greenhouse with real image clutter, stage imbalance, manual annotation, and released evaluation records. This improves the credibility of the result and aligns the manuscript with CEA data-resource papers such as Wheat3D PartNet, while remaining focused on a 2D crop-height evaluation problem (Reena et al., 2025).",
        "The released supplementary material extends the main resource description with ROI-contamination plots, qualitative galleries, negative controls, preprocessing examples, and an open-release map. This structure keeps the benchmark narrative focused while preserving traceable evidence for reanalysis.",
        "The method contribution is deliberately scoped. DINOv3-DCF does not rebuild a complete agricultural robot, UAS, or stereo geometry system; instead it tests frozen foundation features under a strict ROI-level external-greenhouse setting. This distinguishes it from UAS/SfM crop-height models, stereo crop-height estimation, and protected-facility geometry pipelines (Chang et al., 2017; Xie et al., 2021; Kim et al., 2021; Jayasuriya et al., 2024).",
        "The practical lesson is two-sided. Within DINOv3-DCF, plant-focused attention pooling matters more than CLS or uniform patch averaging. Across baseline families, the source-trained random-forest morphometric baseline shows that manual bbox geometry and green-mask extent are powerful cues in this small benchmark. This does not invalidate the foundation-feature diagnostic; it clarifies that the current setting is an ROI-level transfer benchmark rather than a complete automatic height system.",
        "The corrected camera-height result deserves cautious interpretation. Its MAE is close to the attention result, but it changes the metadata assumption rather than the learned representation. It therefore does not replace attention pooling as the primary computational result. Instead, it shows that camera metadata is important enough to check carefully, while the broader diagnostic set shows that metadata correction alone does not explain the remaining early-stage errors.",
        "The failed or partial negative controls are also informative. Feature-statistic alignment worsened external performance, which argues against a simple mean-and-covariance mismatch explanation. The tested DANN configuration did not eliminate the gap even though adversarial training encouraged source-target confusion. These outcomes support the interpretation that the residual problem is structured: early maize morphology, background contamination, ROI ambiguity, and greenhouse-specific context interact with the frozen representation.",
        "The main limitation is scale. DATA325-v0.1 is intentionally presented as a diagnostic benchmark, not a complete deployment dataset. It contains 75 raw photographs, but only 82 evaluated boxes from 25 completed annotated images are used for scored external testing in this paper. The source/target split covers one external greenhouse rather than multiple locations, seasons, cultivars, cameras, and management regimes. The source-domain split used for DCF checkpoint selection is random at the ROI level, so its low MAE is evidence of source fitting rather than independent validation. The three-seed robustness experiment retrains only the DCF head on pre-extracted features and evaluates DATA325-v0.1 with TTA1; it does not represent full foundation-backbone fine-tuning or a multi-site deployment trial.",
        "A second limitation is that manual ROI annotation removes detector error from the main result. This is a necessary experimental control for the current question, but a fieldable greenhouse system will need automatic instance detection, temporal association when videos are used, box-quality rejection, and segmentation-aware normalization. The current release is designed to support those next steps by making the existing boxes, crops, predictions, and errors auditable.",
        "Future work should replace manual boxes with an automatic detector evaluated under the same leakage controls, add segmentation-guided ROI normalization using crop-specific masks or SAM-style prompts, develop plant-mask pooling and structural priors inspired by 3D phenotyping resources, and expand DATA325 across additional greenhouses and growth stages (Kirillov et al., 2023; Liu et al., 2020; Reena et al., 2025).",
    ]:
        add_p(doc, para)

    doc.add_heading("5. Conclusions", level=1)
    add_p(doc, f"DATA325 provides a reproducible external-greenhouse maize height benchmark for diagnosing foundation-feature transfer. Attention-weighted DINOv3 patch pooling reduced zero-shot external DINOv3-DCF MAE from {fmt(old_mae)} to {fmt(attn_mae)} cm, corrected camera-height metadata reached {fmt(corrected_mae)} cm as a diagnostic control, and Attn+aug+TTA8 reached {fmt(best_mae)} cm MAE. A source-trained random-forest morphometric baseline reached {fmt(rf_summary.get('mae_cm'))} cm MAE, emphasizing that manual ROI geometry is a strong cue and that DATA325 should be used to diagnose transfer behavior rather than to claim a complete automatic height system. The residual error is concentrated in early-stage plants and is not solved by simple camera-height correction, feature-statistic alignment, or the tested DANN configuration. The release package makes the current benchmark, predictions, diagnostics, and figure-generation workflow available for future agricultural computer vision work.")

    doc.add_heading("CRediT authorship contribution statement", level=1)
    add_p(doc, f"{AUTHOR_NAME}: Conceptualization, Methodology, Software, Validation, Formal analysis, Investigation, Data curation, Visualization, Writing - original draft, Writing - review & editing. {COAUTHOR_NAME}: Supervision, Methodology guidance, Project administration, Writing - review & editing, Correspondence.")
    doc.add_heading("Declaration of competing interest", level=1)
    add_p(doc, "The authors declare no competing interests.")
    doc.add_heading("Funding", level=1)
    add_p(doc, "No external funding was reported for this work.")

    doc.add_heading("Data availability", level=1)
    add_p(doc, f"DATA325 images, annotations, prediction outputs, diagnostic outputs, selected DCF checkpoints, and scripts are organized for public release at {REPOSITORY_URL}. The companion open-source image-acquisition utility used to support reproducible maize data collection is available at {CAPTURE_TOOL_URL}. Correspondence for the manuscript and data package should be addressed to {COAUTHOR_NAME} ({CORRESPONDING_EMAIL}). Upstream DINOv3 weights are not redistributed.")

    doc.add_heading("Declaration of generative AI and AI-assisted technologies", level=1)
    add_p(doc, "AI-assisted writing and coding tools were used to help draft and organize text and scripts. The authors reviewed and edited all content and are responsible for the final manuscript. Generative AI was not used to create or modify DATA325/source experimental images, attention heatmaps, statistical plots, or evidence figures in this submission package. Conceptual and workflow graphics were generated by deterministic Python drawing scripts.")

    doc.add_heading("References", level=1)
    for ref in revision_refs():
        add_p(doc, ref)

    out = OUT / "manuscript_cea.docx"
    elsevier = OUT / "manuscript_cea_elsevier_style.docx"
    doc.save(out)
    doc.save(elsevier)
    build_up_word_equations(out)
    build_up_word_equations(elsevier)
    return out


def build_highlights() -> Path:
    highlights = [
        "DATA325 benchmarks zero-shot external-greenhouse maize height transfer.",
        "Attention pooling reduced DINOv3-DCF MAE from 41.76 to 30.41 cm.",
        "A source-trained RF morphometric baseline reached 27.10 cm MAE.",
        "Bootstrap diagnostics localize errors in plants below 80 cm.",
        "Open assets include images, boxes, predictions, diagnostics, and code.",
    ]
    for item in highlights:
        if len(item) > 85:
            raise ValueError(f"Highlight exceeds Elsevier guidance: {item}")
    doc = setup_doc()
    doc.add_heading("Highlights", level=1)
    for item in highlights:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.add_run(f"- {item}")
    out = OUT / "highlights.docx"
    doc.save(out)
    (OUT / "highlights.txt").write_text("\n".join(f"- {item}" for item in highlights) + "\n", encoding="utf-8")
    return out


def build_cover_letter() -> Path:
    doc = setup_doc()
    doc.add_heading("Cover letter", level=1)
    rev = load_revision_results()
    summary = rev["summary"].get("model_summary", {})
    old_mae = summary.get("old", {}).get("mae_cm", 41.757666425007145)
    attn_mae = summary.get("attn", {}).get("mae_cm", 30.407656506794254)
    corrected_mae = summary.get("corrected_camheight", {}).get("mae_cm", 30.374317355272247)
    best = rev["summary"].get("best_model_metrics", {})
    rf_mae = rev.get("source_morphometric", {}).get("models", {}).get("random_forest", {}).get("summary", {}).get("mae_cm", 27.10317886989101)
    paragraphs = [
        "Dear Editors,",
        f"I submit the manuscript entitled \"{TITLE}\" for consideration as an Original research paper in Computers and Electronics in Agriculture.",
        "This manuscript presents DATA325-v0.1 as a reproducible external-greenhouse maize height benchmark and evaluates DINOv3-DiffCorn-Fusion (DINOv3-DCF) as a frozen-foundation-feature diagnostic pipeline. This framing matches the journal's emphasis on computational innovation in agricultural imaging.",
        "The manuscript documents DATA325-v0.1 acquisition, manual ROI annotation, distributions, preprocessing, and open-release assets, then reports zero-shot model results, bootstrap confidence intervals, paired tests, ROI contamination diagnostics, source-trained morphometric baselines, uncertainty analysis, error taxonomy, and negative controls.",
        f"Attention-weighted pooling reduced DATA325 MAE from {fmt(old_mae)} to {fmt(attn_mae)} cm, while corrected camera-height metadata reached {fmt(corrected_mae)} cm as a diagnostic control. The best attention plus augmentation plus TTA8 variant reached {fmt(best.get('mae_cm', 29.573550898854325))} cm MAE and {fmt(best.get('rmse_cm', 38.97683219161265))} cm RMSE on 82 boxes. A source-trained random-forest morphometric baseline reached {fmt(rf_mae)} cm MAE, showing that manual ROI geometry is a strong cue and that the paper is a benchmark and transfer diagnostic rather than a complete automatic height system.",
        "The figure set contains 10 main figures plus 8 supplementary figures using real DATA325/source images, true annotations, true prediction outputs, deterministic statistical plots, and code-drawn conceptual diagrams. Generative AI was not used to create or modify scientific evidence figures or the graphical abstract.",
        f"Data, code, predictions, diagnostics, and selected DCF checkpoints are organized for public release at {REPOSITORY_URL}. The companion data-acquisition utility is available at {CAPTURE_TOOL_URL}.",
        "This manuscript is original and is not under consideration elsewhere.",
        "Sincerely,",
        f"{COAUTHOR_NAME}",
        f"Corresponding author, on behalf of all authors",
        f"{AFFILIATION}",
        CORRESPONDING_EMAIL,
    ]
    for para in paragraphs:
        add_p(doc, para)
    out = OUT / "cover_letter_cea.docx"
    doc.save(out)
    return out


def build_supplement(m: dict) -> Path:
    rev = load_revision_results()
    doc = setup_doc()
    doc.add_heading("Supplementary material", level=1)
    add_p(doc, f"Supplementary material for: {TITLE}")
    doc.add_heading("S1. Diagnostic experiment outputs", level=2)
    add_p(doc, f"The companion open-source maize image-acquisition utility is available at {CAPTURE_TOOL_URL}. It is referenced as collection-support software and is not a model-training dependency.")
    for item in [
        "bootstrap_ci.json: 5000-resample MAE, RMSE, and MAPE intervals for each available model output.",
        "paired_tests.json: paired per-box absolute-error differences against Attn+aug+TTA8.",
        "roi_quality_metrics.csv: foreground fraction, background fraction, bbox geometry, brightness, and TTA std for each DATA325 box.",
        "morphometric_baseline.json: leave-one-image-out target-label diagnostic baseline using bbox and mask features.",
        "source_morphometric_baseline.json/csv: source-trained RidgeCV and RandomForestRegressor bbox/mask baselines evaluated zero-shot on DATA325.",
        "uncertainty_diagnostic.json: TTA prediction-std analysis by error and height bin.",
        "error_taxonomy.csv: rule-based failure categories for all DATA325 boxes.",
        "seed_retraining_summary.json/csv: three-seed DCF-head retraining and DATA325 TTA1 evaluation for CLS, patch-mean, attention, and attention+augmentation modes.",
    ]:
        p = doc.add_paragraph()
        p.add_run(f"- {item}")
    doc.add_heading("S2. Diagnostic summaries", level=2)
    add_p(doc, f"Foreground fraction mean: {fmt(rev['roi_summary'].get('foreground_fraction_mean', 0) * 100)}%. Background fraction mean: {fmt(rev['roi_summary'].get('background_fraction_mean', 0) * 100)}%.")
    morph = rev["morphometric"].get("summary", {})
    source_models = rev.get("source_morphometric", {}).get("models", {})
    ridge = source_models.get("ridge_cv", {}).get("summary", {})
    rf = source_models.get("random_forest", {}).get("summary", {})
    add_p(doc, f"Source-trained morphometric baselines: RidgeCV MAE {fmt(ridge.get('mae_cm'))} cm; RandomForestRegressor MAE {fmt(rf.get('mae_cm'))} cm. These models use source labels for fitting and DATA325 labels only for metric calculation.")
    add_p(doc, f"Leave-one-image-out target-label morphometric diagnostic baseline: MAE {fmt(morph.get('mae_cm', 0))} cm, RMSE {fmt(morph.get('rmse_cm', 0))} cm, median AE {fmt(morph.get('median_abs_error_cm', 0))} cm. This uses target labels and is not a zero-shot deployment result.")
    doc.add_heading("S3. Mask QA examples", level=2)
    for path in sorted(MASK_EXAMPLE_DIR.glob("*.jpg"))[:6]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(4.25))
        add_caption(doc, f"Supplementary mask QA example: {path.name}.")
    doc.add_heading("S4. Main figure list", level=2)
    for _, upload, caption, _ in FIGURES:
        add_p(doc, f"{upload}: {caption}")
    doc.add_heading("S5. Supplementary figure list", level=2)
    for _, upload, caption, _ in SUPPLEMENTARY_FIGURES:
        add_p(doc, f"{upload}: {caption}")
        img = SUPP_FIG_OUT / f"{upload}.png"
        if img.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(img), width=Inches(5.7))
            add_caption(doc, caption)
    out = OUT / "supplementary_material.docx"
    doc.save(out)
    return out


def write_sidecars(m: dict) -> None:
    (OUT / "data_availability_statement.md").write_text(
        textwrap.dedent(
            f"""\
            # Data availability statement

            DATA325 raw images, manual bounding boxes, plant-height labels, camera-height
            metadata, model-prediction JSON files, diagnostic outputs, selected DCF
            checkpoints, and reproducibility scripts are organized for public release at:

            {REPOSITORY_URL}

            Companion data-acquisition utility:
            {CAPTURE_TOOL_URL}

            Correspondence for the manuscript and data package: Jian Chen,
            jchen@cau.edu.cn.

            The archive file prepared for release is {OPEN_RELEASE_ARCHIVE}. Upstream
            DINOv3 foundation-model weights are not redistributed and must be obtained
            from their upstream source under the applicable license.
            """
        ),
        encoding="utf-8",
    )
    (OUT / "submission_checklist_cea.md").write_text(
        textwrap.dedent(
            """\
            # Computers and Electronics in Agriculture submission checklist

            ## Included
            - manuscript_cea.docx
            - manuscript_cea_elsevier_style.docx
            - highlights.docx and highlights.txt
            - cover_letter_cea.docx
            - supplementary_material.docx
            - graphical_abstract_non_ai.png/tif/pdf
            - figures/Figure_1...Figure_10 as PNG/PDF
            - supplementary_figures/Supplementary_Figure_1...Supplementary_Figure_8 as PNG/PDF
            - tables/table1_ablation.csv and table2_height_bins.csv
            - tables/roi_quality_metrics.csv, error_taxonomy.csv, seed_retraining_summary.csv, and source_morphometric_baseline.csv
            - reproducibility_json/*.json
            - data_availability_statement.md
            - submission_asset_manifest.md

            ## Notes
            - Article type: Original research paper.
            - Abstract is below 250 words.
            - Each highlight is below 85 characters.
            - Evidence figures use real images, real annotations, model outputs, or deterministic plotting.
            - No generated DATA325/source experimental image is included.
            - Seed robustness is reported for DCF-head retraining with three seeds and DATA325 TTA1 evaluation.
            - Source-trained morphometric baselines do not use DATA325 labels for model fitting.
            - Companion data-acquisition utility: https://github.com/sanyueuy/corn-capture
            """
        ),
        encoding="utf-8",
    )
    manifest = [
        "# CEA submission asset manifest",
        "",
        "## Main files",
        "- manuscript_cea.docx",
        "- manuscript_cea_elsevier_style.docx",
        "- highlights.docx",
        "- highlights.txt",
        "- cover_letter_cea.docx",
        "- supplementary_material.docx",
        "- graphical_abstract_non_ai.png/tif/pdf",
        "- tables/roi_quality_metrics.csv",
        "- tables/error_taxonomy.csv",
        "- tables/seed_retraining_summary.csv",
        "- tables/source_morphometric_baseline.csv",
        "",
        "## Main figures",
    ]
    for _, upload_stem, caption, _ in FIGURES:
        manifest.append(f"- figures/{upload_stem}.png and .pdf: {caption}")
    manifest.extend(["", "## Supplementary figures"])
    for _, upload_stem, caption, _ in SUPPLEMENTARY_FIGURES:
        manifest.append(f"- supplementary_figures/{upload_stem}.png and .pdf: {caption}")
    manifest.extend(
        [
            "",
            "## Diagnostic outputs",
            "- experiments/cea_revision/run_cea_revision_experiments.py generated bootstrap CI, paired tests, ROI quality metrics, morphometric baselines, uncertainty diagnostics, and error taxonomy.",
            "- experiments/cea_revision/seed_retraining contains the 3-seed DCF-head training reports and DATA325 zero-shot evaluation output; summary sidecars are seed_retraining_summary.json/csv.",
            "- source_morphometric_baseline.json/csv reports source-trained RidgeCV and RandomForestRegressor bbox/mask baselines evaluated zero-shot on DATA325.",
            "- Deterministic color-index masks quantify foreground/background only; they are not generated images, labels, or model training inputs.",
            "- Multi-seed robustness is DCF-head retraining over frozen feature bundles; it does not fine-tune DINOv3.",
            "",
            "## Open-source release",
            f"- Repository: {REPOSITORY_URL}",
            f"- Companion data-acquisition utility: {CAPTURE_TOOL_URL}",
            f"- Planned archive: {OPEN_RELEASE_ARCHIVE}",
            "",
            "## Added literature sources",
            "- CEA crop-height/protected crop: Chang 2017; Xie 2021; Kim 2021; Jayasuriya 2024.",
            "- CEA phenotyping/AI reviews and datasets: Patricio 2018; Kamilaris 2018; Li 2020; Liu 2020; Reena 2025.",
            "- CEA crop-specific examples: Li 2019; Xing 2023; Che 2024; Veramendi and Cruvinel 2024; Ariza-Sentis 2024.",
            "- Generalization/foundation models: Sun and Saenko 2016; Koh 2021; Gulrajani and Lopez-Paz 2021; Caron 2021; Kirillov 2023.",
            "",
            "## Reproducibility JSON",
        ]
    )
    for p in sorted(REPRO_OUT.glob("*.json")):
        manifest.append(f"- reproducibility_json/{p.name}")
    (OUT / "submission_asset_manifest.md").write_text("\n".join(manifest) + "\n", encoding="utf-8")


def main() -> None:
    copy_assets()
    build_real_image_figures()
    patch_embedded_figure_titles()
    build_diagnostic_figure()
    build_fig4_mindmap_workflow()
    move_supplementary_figures()
    m = metrics()
    write_csvs(m)
    build_graphical_abstract(m)
    manuscript = build_manuscript(m)
    highlights = build_highlights()
    cover = build_cover_letter()
    supplement = build_supplement(m)
    validate_manuscript_title_page(manuscript)
    validate_manuscript_title_page(OUT / "manuscript_cea_elsevier_style.docx")
    validate_manuscript_equations(manuscript)
    validate_manuscript_equations(OUT / "manuscript_cea_elsevier_style.docx")
    write_sidecars(m)
    package = zip_package()
    print("CEA_DIR", OUT)
    print("MANUSCRIPT", manuscript)
    print("HIGHLIGHTS", highlights)
    print("COVER", cover)
    print("SUPPLEMENT", supplement)
    print("ZIP", package)


if __name__ == "__main__":
    main()
